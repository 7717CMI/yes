from flask import Flask, render_template, request, send_file, redirect, url_for, flash, make_response, jsonify, session
import pandas as pd
import os
import subprocess
import time
import pyautogui
from functools import wraps
from werkzeug.utils import secure_filename
import win32com.client
from docx import Document
from datetime import datetime, timedelta
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from webdriver_manager.chrome import ChromeDriverManager
import threading
import logging
import re
import json
import requests



import os
import pandas as pd
from flask import request, flash, redirect, render_template, jsonify
from werkzeug.utils import secure_filename
from docx import Document
import openai
import time

import random
import os
from datetime import datetime
import pandas as pd
from docx import Document
import openai
import re
from flask import request, flash, redirect, render_template, url_for
from werkzeug.utils import secure_filename  


# NEW NEW
import random
import re
from docx import Document
import win32com.client
import openai
from datetime import datetime
import pandas as pd
import os


app = Flask(__name__)
app.secret_key = 'your_secret_key_change_this_in_production'
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024  # 25MB max file size

# Hardcoded download filename and path
app.config['DOWNLOAD_FILENAME'] = 'ROB.xlsx'
app.config['DOWNLOAD_PATH'] = r'C:\Users\abhijit\Desktop\RPA\\' + app.config['DOWNLOAD_FILENAME']

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Set up logging to capture output
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Add a global variable to track processing status
processing_status = {
    'active': False,
    'message': 'Ready',
    'progress': 0,
    'total': 0,
    'current_file': '',
    'logs': []
}

def allowed_file(filename):
    """Check if file extension is allowed"""
    ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_backend_file(filename):
    """Check if backend file extension is allowed"""
    BACKEND_EXTENSIONS = {'xlsx', 'xls'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in BACKEND_EXTENSIONS

def log_to_status(message):
    """Add a message to the processing status logs"""
    global processing_status
    processing_status['logs'].append(f"{datetime.now().strftime('%H:%M:%S')}: {message}")
    print(f"[LOG] {message}")

# ============================================================================
# HOME ROUTE
# ============================================================================

@app.route('/')
def home():
    return render_template('home.html')
@app.route('/index.html')
def index():
    return render_template('index.html')

# ============================================================================
# DOCUMENT PROCESSING ROUTES
# ============================================================================

@app.route('/document_processing', methods=['GET', 'POST'])
def document_processing():
    global processing_status
    
    if request.method == 'POST':
        try:
            # Get form data - use session data as defaults if available
            article_code = request.form.get('article_code') or request.form.get('open_pr_id') or session.get('open_pr_id', 'D5A-2025-QDFH8C')
            author_name = request.form.get('author_name') or session.get('username', 'abhijit tiwari')
            author_email = request.form.get('author_email') or session.get('email', 'abhijit@coherentmarketinsights.com')
            company_name = request.form.get('company_name', 'Coherent Market Insights')
            phone_number = request.form.get('phone_number') or session.get('mobile', '1234567890')
            image_path = request.form.get('image_path')  # Get image path from form
            
            # Power Automate output folder path
            custom_folder = request.form.get('custom_folder')
            if custom_folder:
                folder_path = custom_folder
            else:
                today = datetime.today()
                folder_path = rf'C:\Users\abhijit\Desktop\RPA\Files\{today.year}\{today.strftime("%m")}\{today.strftime("%d")}'
            
            processing_mode = request.form.get('processing_mode', 'manual')
            
            # Validate paths before processing
            excel_path = r'C:\Users\abhijit\Desktop\RPA\ROB.xlsx'
            
            # Check if required files exist
            validation_errors = []
            if not os.path.exists(excel_path):
                validation_errors.append(f"Excel file not found: {excel_path}")
            if not os.path.exists(folder_path):
                validation_errors.append(f"Folder not found: {folder_path}")
            if image_path and not os.path.exists(image_path):
                validation_errors.append(f"Image file not found: {image_path}")
            if not image_path:
                validation_errors.append("Image path is required")
            
            if validation_errors:
                for error in validation_errors:
                    flash(error)
                return render_template('document_processing.html', 
                                     session_data={
                                         'username': session.get('username', ''),
                                         'email': session.get('email', ''),
                                         'mobile': session.get('mobile', ''),
                                         'open_pr_id': session.get('open_pr_id', ''),
                                         'image_path': image_path or ''
                                     })
            
            # Reset processing status
            processing_status = {
                'active': True,
                'message': 'Starting processing...',
                'progress': 0,
                'total': 0,
                'current_file': '',
                'logs': []
            }
            
            # Start processing in background thread - NOW INCLUDING image_path
            if processing_mode == 'auto':
                threading.Thread(target=process_documents_auto_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number, image_path)).start()
            else:
                threading.Thread(target=process_documents_manual_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number, image_path)).start()
            
            flash('Processing started! Check the status page for updates.')
            return redirect(url_for('processing_status'))
            
        except Exception as e:
            flash(f'Error starting processing: {str(e)}')
            logger.error(f"Error in document_processing: {e}")
            return render_template('document_processing.html', 
                                 session_data={
                                     'username': session.get('username', ''),
                                     'email': session.get('email', ''),
                                     'mobile': session.get('mobile', ''),
                                     'open_pr_id': session.get('open_pr_id', ''),
                                     'image_path': request.form.get('image_path', '')
                                 })
    
    # Pre-populate form with session data if available
    return render_template('document_processing.html', 
                         session_data={
                             'username': session.get('username', ''),
                             'email': session.get('email', ''),
                             'mobile': session.get('mobile', ''),
                             'open_pr_id': session.get('open_pr_id', ''),
                             'image_path': session.get('image_path', '')
                         })

@app.route('/processing_status')
def processing_status_page():
    return render_template('processing_status.html')

@app.route('/api/get_processing_status')
def get_processing_status():
    """API endpoint to get current processing status"""
    global processing_status
    return jsonify(processing_status)

# ============================================================================
# DOCUMENT PROCESSING FUNCTIONS
# ============================================================================
import win32com.client
import re

import win32com.client
import re

def text_of_press_release(doc_path):
    # Open Word application
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Run in the background

    # Open the .doc file (adjust the file path if necessary)
    doc2 = word.Documents.Open(doc_path)

    # Extract the entire text from the document
    doc_text = doc2.Content.Text

    # Remove the first line from the document text
    lines = doc_text.splitlines()
    

    # Define the headings for which you want to add line breaks
    headings = [
        "➤Market Size and Overview",
        "➤Actionable Insights",
        "➤Actionable insights",
        "➤Growth factors",
        "➤Growth Factors",
        "➤Market trends",
        "➤Market Trends",
        "➤Key takeaways ",
        "➤Key Takeaways",
        "➤Market Segment and Regional Coverage ",
        "➤Market segment and regional coverage",
        "➤Key players",
        "➤Key Players",
        "➤Competitive Strategies and Outcomes",
        "❓ Frequently Asked Questions",
        "❓ Frequently asked questions",
        "➤ Frequently Asked Questions",
        "➤ Frequently asked questions"
    ]

    # FIXED: Add a line space AFTER each heading (not before and after)
    for heading in headings:
        doc_text = doc_text.replace(heading, f"{heading}\n")

    # Define the regex pattern for URLs
    url_pattern = re.compile(r"(https?://[^\s]+)")
    
    # Define regex patterns for FAQ questions (numbered questions and roman numerals)
    faq_pattern_numbers = re.compile(r"^\d+\.\s")  # Matches "1. ", "2. ", etc.
    faq_pattern_roman = re.compile(r"^[ivxlcdmIVXLCDM]+\.\s")  # Matches "i. ", "ii. ", "I. ", "II. ", etc.
    
    # Define regex pattern for CTA links (➔)
    cta_pattern = re.compile(r"^➔")  # Matches lines starting with ➔

    # Split the text into lines
    lines = doc_text.splitlines()
    processed_lines = []

    # Iterate over each line
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Skip empty lines in processing, we'll add them strategically
        if not line_stripped:
            continue
            
        # Check if this is a CTA line
        is_cta = cta_pattern.match(line_stripped)
        
        # Check if previous line was a CTA (for adjacent CTA handling)
        prev_was_cta = False
        if processed_lines:
            last_non_empty = None
            for prev_line in reversed(processed_lines):
                if prev_line.strip():
                    last_non_empty = prev_line.strip()
                    break
            if last_non_empty and cta_pattern.match(last_non_empty):
                prev_was_cta = True
        
        # Check if this line is a heading (starts with ➤ or ❓)
        is_heading = line_stripped.startswith('➤') or line_stripped.startswith('❓')
        
        # If a line contains a URL, add space before and after the URL
        if url_pattern.search(line):
            # Add space before (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            processed_lines.append('')  # Add space after
            
        # If a line is an FAQ question (starts with number or roman numeral), add space before it
        elif faq_pattern_numbers.match(line_stripped) or faq_pattern_roman.match(line_stripped):
            # Add space before FAQ question (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this is a CTA line
        elif is_cta:
            # Add space before CTA (unless previous was also CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this line is a heading (starts with ➤ or ❓)
        elif is_heading:
            # Add space before heading (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            # FIXED: Add space AFTER heading
            processed_lines.append('')
            
        else:
            # Regular content line
            processed_lines.append(line)

    # Join the processed lines back into a single string
    chunk = "\n".join(processed_lines)
    
    # Clean up multiple consecutive empty lines (replace with single empty line)
    chunk = re.sub(r'\n\s*\n\s*\n+', '\n\n', chunk)

    # Close the document
    doc2.Close()
    word.Quit()

    # Return the processed content
    return chunk
def run_selenium_automation_single(row_data, category, article_code, author_name, author_email, company_name, phone_number, image_path):
    """Run Selenium automation for a single press release submission"""
    try:
        import random
        AUTHOR_DESCRIPTIONS = [
    """Author of this marketing PR:
Ravina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.
 """    
    """ Author of this marketing PR :
Money Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc. 
"""    
    """ Author of this marketing PR:

Alice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from Openpr her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights.
"""]
        # Extract data from the row
        market_name = row_data['Market Name']
        category = row_data['Category']
        TITLE_PROMPTS = [
    "Is Booming Worldwide 2025-2032",
    "Generated Opportunities, Future Scope 2025-2032",
    "Future Business Opportunities 2025-2032",
    "Growth in Future Scope 2025-2032",
    "Is Booming So Rapidly Growth by 2032",
    "Is Booming So Rapidly with CAGR of 6.9%",
    "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
    "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
    "Set to Witness Significant Growth by 2025-2032",
    "to Witness Massive Growth by 2032",
    "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
    "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
    "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
    "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
    "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
    "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
]

        
        # Extract companies covered (and ensure it handles any missing or malformed data)
        companies = row_data.get('Companies covered', 'No companies listed')
        log_to_status(f"Trying to '{companies}' for market '{market_name}'")

        # Create article title from market name and companies
        # If companies are covered, limit to the first 5 companies, otherwise just use market name
        if companies and isinstance(companies, str) and companies.strip():
            company_list = [c.strip() for c in companies.split(',') if c.strip()]
            first_five_companies = ', '.join(company_list[:5]) if company_list else 'No companies available'
        else:
            first_five_companies = 'No companies available'
        random_prompt = random.choice(TITLE_PROMPTS)
        article_title = f"{market_name} {random_prompt} {first_five_companies}"
        
        # Create multiline text from the row data
        multiline_text= f"""
{market_name} - Market Insights Report

Market Overview:
{row_data.get('Market Size', 'Market analysis and insights')}

Forecast Period: {row_data.get('Forecast Period', 'N/A')}
CAGR: {row_data.get('CAGR', 'N/A')}

Key Market Players:
{row_data.get('Key Players', 'Leading companies in the market')}

For more detailed information, please refer to our comprehensive market research report.
        """
        
        log_to_status(f"Processing: {market_name}")
        log_to_status(f"Using category: {category}")

        log_to_status("Starting Selenium automation for: " + market_name)
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.openpr.com/')
        
        # Handle cookie consent
        try:
            reject = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
            )
            reject.click()
        except:
            pass
        
        # Navigate to submit page
        submit = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
        )
        submit.click()
        
        # Enter article code
        input_box = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
        )
        input_box.clear()
        input_box.send_keys(article_code)
        
        # Submit code
        try:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
            )
            submit2.click()
        except:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
            )
            submit2.click()
        
        # Fill form fields
        name = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
        )
        name.send_keys(author_name)
        
        email = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
        )
        email.clear()
        email.send_keys(author_email)
        
       
        
        number = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
        )
        number.clear()
        number.send_keys(phone_number)
        
        ComName = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
        )
        ComName.clear()
        ComName.send_keys("Coherent Market Insights")
        
        s1 = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
        )
        s1.click()
        
        # Handle category selection with better error handling
        Category_element = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
        )
        
        # Wait a moment for the dropdown to be fully loaded
        time.sleep(1)
        
        # Get all available options from the dropdown
        select_obj = Select(Category_element)
        available_options = [option.text.strip() for option in select_obj.options]
        log_to_status(f"Available dropdown options: {available_options}")
        
        # Use the category directly from the row data
        website_category = category.strip()
        log_to_status(f"Trying to select category: '{website_category}'")
        
        # Check if the exact category exists in the dropdown
        if website_category in available_options:
            log_to_status(f"Category '{website_category}' found in dropdown options")
        else:
            log_to_status(f"Category '{website_category}' NOT found in dropdown options")
            log_to_status(f"Available options are: {available_options}")
        
        try:
            # Try different selection methods
            select_obj.select_by_visible_text(website_category)
            log_to_status(f"Successfully selected category: '{website_category}'")
        except Exception as e:
            log_to_status(f"Error selecting category '{website_category}' by visible text: {e}")
            
            # Try selecting by value if visible text fails
            try:
                for option in select_obj.options:
                    if option.text.strip() == website_category:
                        select_obj.select_by_value(option.get_attribute('value'))
                        log_to_status(f"Successfully selected category by value: '{website_category}'")
                        break
                else:
                    raise Exception(f"Could not find option with text '{website_category}'")
            except Exception as e2:
                log_to_status(f"Error selecting category by value: {e2}")
                
                # Final fallback - try to select "Health & Medicine" directly if we have a health-related category
                try:
                    if "health" in website_category.lower() or "medicine" in website_category.lower():
                        select_obj.select_by_visible_text("Business,Economy,Finance,Banking & Insurance")
                        log_to_status("Selected 'Business,Economy,Finance,Banking & Insurance' as fallback for health-related category")
                    else:
                        select_obj.select_by_index(1)  # Select first real option
                        log_to_status("Selected first available option as final fallback")
                except Exception as e3:
                    log_to_status(f"Final fallback also failed: {e3}")
                    select_obj.select_by_index(0)  # Select any option to continue
        
        title = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
        )
        
        title.clear()
        title.send_keys(article_title,)
        
        text = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
        )
        text.clear()
        text.send_keys(multiline_text)
        
        about = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
        )
        about.clear()
        multi = """Contact Us:

            Mr. Shah
            Coherent Market Insights
            533 Airport Boulevard,
            Suite 400, Burlingame,
            CA 94010, United States
            US: + 12524771362
            UK: +442039578553
            AUS: +61-8-7924-7805
            India: +91-848-285-0837"""
        about.send_keys(multi)
        address = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
        )
        address.clear()
        random_author = random.choice(AUTHOR_DESCRIPTIONS)
        address_content = f"{random_author}\n\nAbout Us:\nCoherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries."

        address.send_keys(address_content)
        
        image = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
        )
        image.clear()
        image.send_keys(image_path)
        
        caption = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
        )
        caption.clear()
        caption.send_keys("This is a test caption for the image.")
        
        notes = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
        )
        notes.clear()
        notes.send_keys("This is a test notes section for the press release submission.")
        
        # Agree to terms
        tick1 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
        )
        tick1.click()
        
        tick2 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
        )
        tick2.click()
        
        # Submit form
        final = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
        )
        final.click()
        
        time.sleep(5)
        log_to_status(f"Selenium automation completed successfully for: {market_name}")
        driver.quit()  # Quit the browser tab after publication
        return True, market_name
        
    except Exception as e:
        market_name = row_data.get('Market Name', 'Unknown') if 'row_data' in locals() else 'Unknown'
        log_to_status(f"Selenium automation error for {market_name}: {e}")
        try:
            driver.quit()
        except:
            pass
        return False, market_name


def run_selenium_automation_all_rows(article_code, author_name, author_email, company_name, phone_number):
    """Run Selenium automation for all rows in the Excel file"""
    try:
        # Read all data from Excel
        import pandas as pd
        excel_path = r'C:\Users\abhijit\Desktop\RPA\ROB.xlsx'
        df = pd.read_excel(excel_path)
        
        log_to_status(f"Found {len(df)} rows in Excel file")
        
        # Results tracking
        successful_submissions = []
        failed_submissions = []
        
        # Process each row
        for index, row in df.iterrows():
            log_to_status(f"\n{'='*50}")
            log_to_status(f"Processing Row {index + 1} of {len(df)}")
            log_to_status(f"{'='*50}")
            
            try:
                category = row['Category'] if 'Category' in row else ''
                success, market_name = run_selenium_automation_single(
                    row_data=row,
                    category=category,
                    article_code=article_code,
                    author_name=author_name,
                    author_email=author_email,
                    company_name=company_name,
                    phone_number=phone_number,
                    image_path=image_path if 'image_path' in locals() else None
                )
                
                if success:
                    successful_submissions.append(market_name)
                    log_to_status(f"✅ Successfully submitted: {market_name}")
                else:
                    failed_submissions.append(market_name)
                    log_to_status(f"❌ Failed to submit: {market_name}")
                
                # Add a delay between submissions to avoid overwhelming the server
                if index < len(df) - 1:  # Don't wait after the last submission
                    log_to_status("Waiting 360 seconds before next submission...")
                    time.sleep(20)
                    
            except Exception as e:
                market_name = row.get('Market Name', f'Row {index + 1}')
                failed_submissions.append(market_name)
                log_to_status(f"❌ Error processing row {index + 1} ({market_name}): {e}")
                continue
        
        # Final summary
        log_to_status(f"\n{'='*50}")
        log_to_status("FINAL SUMMARY")
        log_to_status(f"{'='*50}")
        log_to_status(f"Total rows processed: {len(df)}")
        log_to_status(f"Successful submissions: {len(successful_submissions)}")
        log_to_status(f"Failed submissions: {len(failed_submissions)}")
        
        if successful_submissions:
            log_to_status(f"\n✅ Successfully submitted:")
            for market in successful_submissions:
                log_to_status(f"  - {market}")
        
        if failed_submissions:
            log_to_status(f"\n❌ Failed submissions:")
            for market in failed_submissions:
                log_to_status(f"  - {market}")
        
        return len(successful_submissions), len(failed_submissions)
        
    except Exception as e:
        log_to_status(f"Error in run_selenium_automation_all_rows: {e}")
        return 0, 0

def run_selenium_automation(article_code, article_title, multiline_text, category, author_name, 
                          author_email, company_name, phone_number, image_path):
    """Enhanced run_selenium_automation function that reads category from Excel"""
    try:
        import random
        AUTHOR_DESCRIPTIONS = [
            """Author of this marketing PR:\nRavina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.\n """,
            """ Author of this marketing PR :\nMoney Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc. \n""",
            """ Author of this marketing PR:\n\nAlice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from Openpr her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights.\n"""
        ]
        log_to_status("Starting Selenium automation...")
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.openpr.com/')
        
        # Handle cookie consent
        try:
            reject = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
            )
            reject.click()
        except:
            pass
        
        # Navigate to submit page
        submit = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
        )
        submit.click()
        
        # Enter article code
        input_box = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
        )
        input_box.clear()
        input_box.send_keys(article_code)
        
        # Submit code
        try:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
            )
            submit2.click()
        except:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
            )
            submit2.click()
        
        # Fill form fields
        name = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
        )
        name.send_keys(author_name)
        
        email = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
        )
        email.clear()
        email.send_keys(author_email)
        
        '''pr_agency = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[3]/div/input'))
        )
        pr_agency.clear()
        pr_agency.send_keys(author_name)'''
        
        number = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
        )
        number.clear()
        number.send_keys(phone_number)
        
        ComName = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
        )
        ComName.clear()
        ComName.send_keys("Coherent Market Insights")
        
        s1 = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
        )
        s1.click()
        
        # Handle category selection with provided category argument
        Category_element = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
        )
        time.sleep(1)
        select_obj = Select(Category_element)
        available_options = [option.text.strip() for option in select_obj.options]
        log_to_status(f"Available dropdown options: {available_options}")
        website_category = category.strip()
        log_to_status(f"Trying to select category: '{website_category}'")
        try:
            select_obj.select_by_visible_text(website_category)
            log_to_status(f"Successfully selected category: '{website_category}'")
        except Exception as e:
            log_to_status(f"Error selecting category '{website_category}': {e}")
            try:
                select_obj.select_by_visible_text("Business,Economy,Finance,Banking & Insurance")
                log_to_status("Selected 'Business,Economy,Finance,Banking & Insurance'")
            except:
                select_obj.select_by_index(1)
                log_to_status("Selected first available option as final fallback")
        
        title = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
        )
        title.clear()
        log_to_status(f"Trying to '{article_title}'")
        title.send_keys(article_title)
        
        text = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
        )
        text.clear()
        text.send_keys(multiline_text)
        
        about = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
        )
        about.clear()
        multi = """Contact Us:

Mr. Shah
Coherent Market Insights
533 Airport Boulevard,
Suite 400, Burlingame,
CA 94010, United States
US: + 12524771362
UK: +442039578553
AUS: +61-8-7924-7805
India: +91-848-285-0837
"""
        about.send_keys(multi)
        
        address = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
        )
        address.clear()
        random_author = random.choice(AUTHOR_DESCRIPTIONS)
        address_content = f"{random_author}\n\nAbout Us:\nCoherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries."

        address.send_keys(address_content)
        
        image = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
        )
        image.clear()
        image.send_keys(image_path)
        
        caption = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
        )
        caption.clear()
        caption.send_keys("This is a test caption for the image.")
        
        notes = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
        )
        notes.clear()
        notes.send_keys("This is a test notes section for the press release submission.")
        
        # Agree to terms
        tick1 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
        )
        tick1.click()
        
        tick2 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
        )
        tick2.click()
        
        # Submit form
        final = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
        )
        final.click()
        
       
        
        time.sleep(10)
        log_to_status("Selenium automation completed successfully")
        return True
        
        
    except Exception as e:
        log_to_status(f"Selenium automation error: {e}")
        try:
            driver.quit()
        except:
            pass
        return False

def process_documents_auto_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number,image_path):
    """Process documents automatically with status feedback"""
    global processing_status
    
    try:
        log_to_status(f"Starting auto processing. Folder: {folder_path}")
        
        excel_path = r'C:\Users\abhijit\Desktop\RPA\ROB.xlsx'
        import random
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0
        TITLE_PROMPTS = [
    "Is Booming Worldwide 2025-2032",
    "Generated Opportunities, Future Scope 2025-2032",
    "Future Business Opportunities 2025-2032",
    "Growth in Future Scope 2025-2032",
    "Is Booming So Rapidly Growth by 2032",
    "Is Booming So Rapidly with CAGR of 6.9%",
    "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
    "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
    "Set to Witness Significant Growth by 2025-2032",
    "to Witness Massive Growth by 2032",
    "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
    "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Growth Factors, Trends",
    "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
    "Size, Share Growth Status, Emerging Technology, Key Players, Industry Opportunities, and Forecast till 2032",
    "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
    "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
]

        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Auto-processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(doc_path)
                # Get companies for this market from Excel row
                matching_row = keywords_df[keywords_df['Market Name'] == market_name]
                if not matching_row.empty:
                    companies = matching_row.iloc[0].get('Companies covered', '')
                else:
                    companies = ''
                if companies and isinstance(companies, str) and companies.strip():
                    company_list = [c.strip() for c in companies.split(',') if c.strip()]
                    first_five_companies = ', '.join(company_list[:3]) if company_list else 'No companies available'
                else:
                    first_five_companies = 'No companies available'
                random_prompt = random.choice(TITLE_PROMPTS)
                x = f"{market_name} {random_prompt}|{first_five_companies}"

                # Refine article_title using OpenAI for grammar correction
                def refine_title_with_openai(title):
                    try:
                        import openai
                        client = openai.OpenAI(api_key=OPENAI_CONFIG['API_KEY'])
                        prompt = f"Please improve the grammar, structure, and readability of this press release title to make it more interesting and engaging for readers. Keep all original words intact - only rearrange, or adjust formatting as needed and there should be no parenthesis at end or start of title : '{title}'" 
                        response = client.chat.completions.create(
                            model=OPENAI_CONFIG['MODEL'],
                            messages=[
                                {"role": "system", "content": "You are an expert editor for press releases."},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=60,
                            temperature=2.0
                        )
                        return response.choices[0].message.content.strip()
                    except Exception as e:
                        log_to_status(f"OpenAI title refinement error: {e}")
                        return title

                article_title = refine_title_with_openai(x)
                
                category = matching_row.iloc[0].get('Category', '') if not matching_row.empty else ''
                # Run automation
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code,article_title, multiline_text, category, 
                                                author_name, author_email, company_name, phone_number, image_path)
                
                if success:
                    log_to_status(f"SUCCESS: Published {market_name}")
                    processed_count += 1
                else:
                    log_to_status(f"FAILED: Could not publish {market_name}")
                
                time.sleep(10)
                  # Wait 60 seconds between submissions to avoid rate limiting
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Auto-processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Auto processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Auto processing error: {e}")

def process_documents_manual_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number,image_path):
    """Process documents with manual intervention and status feedback"""
    global processing_status
    import random
    
    try:
        log_to_status(f"Starting manual processing. Folder: {folder_path}")
        
        excel_path = r'C:\Users\abhijit\Desktop\RPA\ROB.xlsx'
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0

        TITLE_PROMPTS = [
    "Is Booming Worldwide 2025-2032",
    "Generated Opportunities, Future Scope 2025-2032",
    "Future Business Opportunities 2025-2032",
    "Growth in Future Scope 2025-2032",
    "Is Booming So Rapidly Growth by 2032",
    "Is Booming So Rapidly with CAGR of 6.9%",
    "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
    "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
    "Set to Witness Significant Growth by 2025-2032",
    "to Witness Massive Growth by 2032",
    "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
    "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
    "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
    "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
    "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
    "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
]

        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(doc_path)
                # Get companies for this market from Excel row
                matching_row = keywords_df[keywords_df['Market Name'] == market_name]
                if not matching_row.empty:
                    companies = matching_row.iloc[0].get('Companies covered', '')
                else:
                    companies = ''
                if companies and isinstance(companies, str) and companies.strip():
                    company_list = [c.strip() for c in companies.split(',') if c.strip()]
                    first_five_companies = ', '.join(company_list[:5]) if company_list else 'No companies available'
                else:
                    first_five_companies = 'No companies available'
                random_prompt = random.choice(TITLE_PROMPTS)
                article_title = f"{market_name} {random_prompt} {first_five_companies}"
                
                # Get category for this market from Excel row
                category = matching_row.iloc[0].get('Category', '') if not matching_row.empty else ''
                # Run automation
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code, article_title, multiline_text, category, 
                                                author_name, author_email, company_name, phone_number, image_path)
                
                if success:
                    log_to_status(f"Published {market_name}")
                    processed_count += 1
                
                time.sleep(5)
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Manual processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Manual processing error: {e}")

# ============================================================================
# ROB PROCESSING ROUTES
# ============================================================================

@app.route('/rob', methods=['GET', 'POST'])
def rob():
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        open_pr_id = request.form.get('open_pr_id')
        mobile = request.form.get('mobile')
        extract_count = int(request.form.get('extract_count', 200))

        # Validate required fields
        if not all([username, email, open_pr_id, mobile, extract_count]):
            flash('All fields are required!')
            return redirect(request.url)

        file = request.files.get('file')
        if not file or file.filename == '':
            flash('Excel file is required!')
            return redirect(request.url)

        if not allowed_file(file.filename):
            flash('Only Excel files (.xlsx, .xls) and CSV files are allowed!')
            return redirect(request.url)

        # Use secure_filename to avoid path issues
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        # Store user data in session for later use
        session['username'] = username
        session['email'] = email
        session['open_pr_id'] = open_pr_id
        session['mobile'] = mobile
        
        return redirect(url_for('process_rob', file_path=input_path,
                                username=username, email=email,
                                open_pr_id=open_pr_id, mobile=mobile,
                                extract_count=extract_count))
    return render_template('rob.html')

@app.route('/process_rob')
def process_rob():
    file_path = request.args.get('file_path')
    username = request.args.get('username')
    email = request.args.get('email')
    open_pr_id = request.args.get('open_pr_id')
    mobile = request.args.get('mobile')
    extract_count = int(request.args.get('extract_count', 200))

    if not file_path or not os.path.exists(file_path):
        flash('Missing or invalid file path')
        return redirect(url_for('rob'))

    try:
        # Read the cleaned ROB file
        if file_path.endswith('.csv'):
            df_original = pd.read_csv(file_path,encoding='utf-8', engine='python')
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')

        total_rows = len(df_original)
        
        if total_rows < extract_count:
            flash(f'⚠️ File only has {total_rows} rows, but you requested {extract_count} rows!')
            extract_count = total_rows

        # Step 1: Extract top N rows
        extracted_rows = df_original.head(extract_count).copy()
        
        # Step 2: Get remaining rows (original minus extracted)
        remaining_rows = df_original.iloc[extract_count:].copy()

        # Step 3: Create timestamp for remaining file
        today = datetime.today()
        timestamp = f"{today.year}_{today.month:02d}_{today.day:02d}"
        
        # Step 4: Save extracted rows as ROB.xlsx to Desktop/RPA
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        rob_output_path = os.path.join(rpa_folder, "ROB.xlsx")
        extracted_rows.to_excel(rob_output_path, index=False)
        
      # Step 5: Save remaining rows with timestamp in date-wise folder under Weekly_RID
        weekly_rid_folder = os.path.join(r"C:\Users\abhijit\Desktop\RPA\Weekly_RID", str(today.year), f"{today.month:02d}", f"{today.day:02d}")
        os.makedirs(weekly_rid_folder, exist_ok=True)
        remaining_filename = f"cleaned_weekly_rid_{timestamp}.xlsx"
        remaining_output_path = os.path.join(weekly_rid_folder, remaining_filename)
        remaining_rows.to_excel(remaining_output_path, index=False)
        
        # Step 6: Store info in session for the result page
        session['rob_file_path'] = rob_output_path
        session['remaining_file_path'] = remaining_output_path
        session['remaining_filename'] = remaining_filename
        session['extracted_count'] = extract_count
        session['remaining_count'] = len(remaining_rows)
        session['total_count'] = total_rows
        
        flash(f'✅ Successfully processed {total_rows} rows!')
        flash(f'📁 Remaining {len(remaining_rows)} rows saved in: {weekly_rid_folder} as {remaining_filename} (ready for download)')
        
        # Use render_template instead of redirect
        return render_template('rob_result.html',
                             extracted_count=extract_count,
                             remaining_count=len(remaining_rows),
                             total_count=total_rows,
                             username=username,
                             records_processed=total_rows,
                             weekly_rid_folder=weekly_rid_folder,
                             remaining_filename=remaining_filename)

    except Exception as e:
        flash(f'❌ Error processing ROB file: {str(e)}')
        return redirect(url_for('rob'))

@app.route('/download_remaining_rob')
def download_remaining_rob():
    """Download the remaining ROB file (original minus extracted rows)"""
    try:
        remaining_file_path = session.get('remaining_file_path')
        remaining_filename = session.get('remaining_filename', 'cleaned_rob_remaining.xlsx')
        
        if remaining_file_path and os.path.exists(remaining_file_path):
            return send_file(remaining_file_path, as_attachment=True, download_name=remaining_filename)
        else:
            flash('❌ Remaining ROB file not found. Please process a file first.')
            return redirect(url_for('rob'))
    except Exception as e:
        flash(f'❌ Error downloading remaining file: {str(e)}')
        return redirect(url_for('rob'))

@app.route('/download_extracted_rob')
def download_extracted_rob():
    """Download the extracted ROB.xlsx file and trigger OpenAI content generation"""
    try:
        rob_file_path = session.get('rob_file_path')
        
        if rob_file_path and os.path.exists(rob_file_path):
            
            # Start OpenAI content generation in background thread with 5-second delay
            print("🔄 Starting background OpenAI content generation...")
            threading.Thread(target=delayed_openai_content_generation, args=(5,)).start()
            
            return send_file(rob_file_path, as_attachment=True, download_name='ROB.xlsx')
        else:
            flash('❌ ROB.xlsx file not found. Please process a file first.')
            return redirect(url_for('rob'))
            
    except Exception as e:
        flash(f'❌ Error downloading ROB file: {str(e)}')
        return redirect(url_for('rob'))


def delayed_openai_content_generation(delay_seconds=5):
    """Generate content using OpenAI after delay (replaces Power Automate)"""
    try:
        print(f"⏳ Waiting {delay_seconds} seconds before starting content generation...")
        time.sleep(delay_seconds)
        
        print("🤖 Starting OpenAI content generation...")
        
        # ROB file path
        rob_file_path = r"C:\Users\abhijit\Desktop\RPA\ROB.xlsx"
        
        if not os.path.exists(rob_file_path):
            print("❌ ROB.xlsx file not found!")
            return
        
        # Check if API key is configured
        if OPENAI_CONFIG['API_KEY'] == 'your-openai-api-key-here':
            print("❌ OpenAI API key not configured! Using fallback content.")
        
        # Read ROB file
        df = pd.read_excel(rob_file_path)
        print(f"📊 Found {len(df)} markets in ROB file")
        
        # Create output directory with current date
        today = datetime.today()
        output_dir = os.path.join(
            r"C:\Users\abhijit\Desktop\RPA\Files",
            str(today.year),
            f"{today.month:02d}",
            f"{today.day:02d}"
        )
        os.makedirs(output_dir, exist_ok=True)
        print(f"📁 Output directory: {output_dir}")
        
        successful = 0
        failed = 0
        
        # Process each row
        for index, row in df.iterrows():
            try:
                # Extract market data
                
                #print(f"[{index+1}/{len(df)}] Processing: {data['market_name']}")
                
                # Generate content
                content = generate_blog_from_row(row)
                
                # Save document
                success, filepath = save_market_document(row.get('Market Name'), content, output_dir)
                
                if success:
                    print(f"✅ Generated: {os.path.basename(filepath)}")
                    successful += 1
                else:
                    print(f"❌ Failed to save: (row.get('Market Name')")
                    failed += 1
                
                # Rate limiting
                time.sleep(2)
                
            except Exception as e:
                print(f"❌ Error processing  {e}")
                failed += 1
                continue
        
        print(f"✅ Content generation completed!")
        print(f"📊 Successful: {successful}, Failed: {failed}")
        print(f"📁 Files saved in: {output_dir}")
        
    except Exception as e:
        print(f"❌ Error in content generation: {e}")


# ============================================================================
# ADD THESE SIMPLE HELPER FUNCTIONS
# ============================================================================

import re
import openai

# OpenAI config
OPENAI_CONFIG = {
    'API_KEY': 'lP',
    'MODEL': 'gpt-4.1-mini',  # Or 'gpt-4o' or 'gpt-3.5-turbo'
    'TEMPERATURE': 0.7
}

def generate_blog_from_row(row):
    """Extract data from row, format prompt, and generate blog via OpenAI."""
    try:
        # Extract values
        market_name = row.get('Market Name', 'Unknown Market')
        forecast_period = row.get('Forecast Period', 'Not specified')
        market_size_year = row.get('Market Size Year', 'Not specified')
        market_size = row.get('Market Size', '')
        cagr = row.get('CAGR', 'Not specified')
        key_players = row.get('Key Players', 'Not specified')

        # Extract size values from combined string
        match = re.search(r'USD ([\d.]+ Bn); Market Size in 2032: USD ([\d.]+ Bn)', market_size)
        if match:
            market_size_2025 = f"USD {match.group(1)}"
            market_size_2032 = f"USD {match.group(2)}"
        else:
            market_size_2025 = ""
            market_size_2032 = ""

        # Prepare prompt string from extracted values
        output = ( f"""`
We are coherent market insights and going to publish this data on third part PR website, thus do not mention any competitors name (market intelligence firm, market research, consulting and data and analytics companies), and take ownership for market share, size, growth rate as this research done by our own research team however any real instances and stats you can use reference such as company or organisation or others as mentioned in detail explanation. Write the blog covering the below pointers: Start content with market size and overview paragraph (do not add heading and do not use bullet for this section), ➔ add first CTA link here, ➤Strategic Actionable Insights for the Market, ➤Market Taxonomy and Regional coverage of Report, ➔ add Second CTA link here, ➤Leading Companies of the Market, ➤Key Growth Drivers Fueling Market Expansion, ➔add Third CTA link here, ➤Key Reasons for Buying the (insert market name here) Report ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ➤ Frequently Asked Questions, All these pointers should act as a heading to respective paragraphs, do not miss any bullet for above given. Consider the following points while generating content: Flow of information, all given secondary keywords must be covered, and there should be a heading for each paragraph or bullet pointers. I need this content to be very presentable format thus make sure there should be space after and before of each heading or CTA links (first CTA link, Strategic Actionable Insights for the Market, Market Taxonomy and Regional coverage of Report, Second CTA link, Leading Companies of the Market, Key Growth Drivers Fueling Market Expansion, Third CTA link, Key Reasons for Buying the (insert market name here) Report, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, and Frequently Asked Questions), this will increase the readability. Cover content in in bullet pointers whenever possible each paragraph should be short. Our readers are already experts in the field, so always try to generate content that provides unique insights and value addition for experts. Thus, while generating each piece of content, it should be data-backed with actual instances and stats from recent years 2025 and 2024, covering maximum stats that increase the authenticity and are enough to support or make decisions based upon reading this blog. Do not add generic content that is already known to readers, such as definitions, advantages, disadvantages, or other generic content. Flow of information should be as below. Start the blog with writing 5 to 7 sentence ling paragraph start content with talking about Industry in 2 to 3 sentences and should be matches with heading of the blog. followed by 2 to 4 sentence on market size and CAGR formatted as The Global (Insert Market Name) Market size is estimated to be valued at USD (Insert Market Value for 2025) in 2025 (market size) and is expected to reach USD (Insert Market Value for 2032) by 2032 (market forecast), exhibiting a compound annual growth rate (CAGR) of (Insert CAGR)% from 2025 to 2032. Do not add heading and do not use bullet for this section. (Make sure all values such as market size, CAGR, Leading Companies of the Market should reflect exactly same in output as provided in input). Then First CTA link. Then Strategic Actionable Insights for the Market: In Strategic Actionable Insights for the Market, cover 3 to 4 Strategic Actionable Insights for the Market in 10 to 12 sentences each insights comprises of 2 t 3 sentences covering supporting data for each insights. Insights could be essential quantitative indicators that lead to market size. For example, supply-side indicators (ex. Production Capacity, Pricing, and Exports. but not necessary to add if not matches to given market), demand-side indicators (ex. Pricing, Imports, Various Use Cases across industries. but not necessary to add if not matches to given market), micro-indicators, nano-size indicators—focus more on the quantitative aspects. Each Strategic Actionable Insights for the Market must have two sentence stats or actual instance examples from the recent year to support each point given in Strategic Actionable Insights for the Market, so that each given point look complete and meaningful. Next part is Market Taxonomy and Regional coverage of Report where enlist the all subsegment under each segment categories and fragment region into given format. Comprehensive Segmentation and Classification of the Report: • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. Regional and Country Analysis: • North America: U.S. and Canada • Latin America: Brazil, Argentina, Mexico, and Rest of Latin America • Europe: Germany, U.K., Spain, France, Italy, Benelux, Denmark, Norway, Sweden, Russia, and Rest of Europe • Asia Pacific: China, Taiwan, India, Japan, South Korea, Indonesia, Malaysia, Philippines, Singapore, Australia, and Rest of Asia Pacific • Middle East & Africa: Bahrain, Kuwait, Oman, Qatar, Saudi Arabia, United Arab Emirates, Israel, South Africa, North Africa, Central Africa, and Rest of MEA. Then Second CTA link. Then Leading Companies of the Market: List 12 to 20 highly relevant Leading Companies of the Market for the given market. Furthermore, add 2-3 statements on competitive strategies adopted by a few Leading Companies of the Market, mentioning actual strategies and entities involved along with the actual outcome. Key Growth Drivers Fueling Market Expansion: Growth factor heading and short paragraph (3-4 Key Growth Drivers Fueling Market Expansion covered under 10 to 12 sentences) with supporting stats or examples from the recent year in the content, each factors should be covered in two to three sentences thus entire Key Growth Drivers Fueling Market Expansion content will be covered in 10 to 12 sentences long. No sub bullet is needed in Growth Factor. Then Add Third CTA link. Key Reasons for Buying the (insert market name here) Report, and its exact content as shared in data. Then Emerging Trends and Market Shift: Market Trend heading and short paragraphs with supporting stats or examples from the recent year in the content (No bullet needed for as opportunity are written in paragraph format). Then High-Impact Market Opportunities by Segment and Region: Provide 3 to 4 High-Impact Market Opportunities by Segment and Region, 2-3 opportunities based upon segment and one opportunity based upon region in a paragraph format. Each opportunity described in two to three sentences and supported by actual data. For each opportunity, identify a single segment and provide a short description of the opportunity within that segment. Similarly, highlight one region, along with a brief description of the opportunity in each regional market (No bullet needed for as opportunity are written in paragraph format). Make sure to not to mention our organization name or relevant terms anywhere in the output such as coherent market insights or our analyst team or our research team Given Market Name and Data:
Input of different market
            f"Market Name- {market_name}; "
            f"CAGR:{cagr}; "
            f"Forecast period is: {forecast_period}; "
            f"Market Size for {market_size_year} is {market_size_2025}; "
            f"Market Size in 2032: {market_size_2032}; "
            f"Key players: {key_players}"
 ➤Key Reasons for Buying the (insert market name here) Report: • Comprehensive analysis of the changing competitive landscape • Assists in decision-making processes for the businesses along with detailed strategic planning methodologies • The report offers forecast data and an assessment of the (insert market name here) • Helps in understanding the key product segments and their estimated growth rate • In-depth analysis of market drivers, restraints, trends, and opportunities • Comprehensive regional analysis of the (insert market name here) • Extensive profiling of the key stakeholders of the business sphere • Detailed analysis of the factors influencing the growth of the (insert market name here). From an SEO perspective, we need to cover all given keywords from the list below. However, they should appear naturally so that the content flow looks natural for the reader. Keyword List: market share, market size, market research, market insights, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, market challenges, market growth, market forecast, market companies, market players, market analysis, market drivers, market restraints, market scope, market dynamics, market segments, market report, market growth strategies, market revenue, industry size, industry share, industry trends, and business growth, furthermore - Market size and market report, market revenue, market share, trends keywords are mandatory to be added twice in content. In addition to the above requirement, in 5 places, add the actual market name along with the above keywords so that long-tail keywords will be generated. These long-tail keywords are market name + size, market name + report, market name + revenue, market name + share, market name + trends. Make sure all given keywords are naturally fit, do not try to infuse forcefully, flow of information should be natural and meaningful, furthermore make sure spelling and structure of sentences from generated output are grammatically correct. Furthermore, based on the market name, create a set of Frequently Asked Questions that are highly relevant and customized to the specific market. The sample Frequently Asked Questions below are for understanding purposes only. For the given market, questions can be completely replaced. However, please tailor the actual questions to the market name and the insights provided in the report: 1. Who are the dominant players in the (Market Name) market? 2. What will be the size of the (Market Name) market in the coming years? 3. Which end users industry has the largest growth opportunity? 4. How will market development trends evolve over the next five years? 5. What is the nature of the competitive landscape and challenges in the (Market Name) market? 6. What go-to-market strategies are commonly adopted in the (Market Name) market? Make sure to answer to all FAQs. In the case of country-level markets, please exclude the word 'Global' and High-Impact Market Opportunities by Segment and Region where other regions are mentioned. Make sure to add catchy bullet in generated output. I have shared the reference bullet with you. Make sure to add this bullet. For heading use these bullet- ), ➔ add first CTA link here, ➤Strategic Actionable Insights for the Market, ➤Market Taxonomy and Regional coverage of Report, ➔ Inserted Second CTA link, ➤Leading Companies of the Market, ➤Key Growth Drivers Fueling Market Expansion, ➔ Inserted Third CTA link, ➤Key Reasons for Buying the (insert market name here) Report, ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ➤ Frequently Asked Questions. Make sure do not miss any bullet including CTA bullet which is ➔. For subpointers under main headings use bullets which is in reference as provided- Strategic Actionable Insights for the Market ●, Market Taxonomy and Regional coverage of Report●, Leading Companies of the Market●. Make sure to use these bullets for given subpointers. Ensure proper bullet formatting so that each point is marked with a single bullet only, and avoid placing two bullets adjacent to each other.
""")

        # Send to OpenAI
        client = openai.OpenAI(api_key=OPENAI_CONFIG['API_KEY'])

        response = client.chat.completions.create(
            model=OPENAI_CONFIG['MODEL'],
            messages=[
                {"role": "user", "content": output}
            ],
            temperature=OPENAI_CONFIG.get('TEMPERATURE', 0.7)
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"OpenAI error: {e}")
        return "Error generating content."

def save_market_document(market_name, content, output_folder):
    """Save content as Word document"""
    try:
        doc = Document()
        doc.add_heading(f"{market_name} - Market Research Report", level=1)
        
        # Add content paragraphs
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Add contact info
      
        
        # Save file
        safe_name = "".join(c for c in market_name if c.isalnum() or c in (' ', '_')).strip()
        filename = f"ROB_{safe_name}.doc"
        filepath = os.path.join(output_folder, filename)
        doc.save(filepath)
        
        return True, filepath
        
    except Exception as e:
        print(f"Error saving document: {e}")
        return False, None


'''app.route('/api/auto_trigger_power_automate', methods=['POST'])
def auto_trigger_power_automate():
    """API endpoint for auto-triggering Power Automate"""
    try:
        # Check if we should trigger (based on recent download)
        if session.get('trigger_power_automate'):
            # Clear the flag
            session['trigger_power_automate'] = False
            
            # Trigger in background
            threading.Thread(target=delayed_power_automate_trigger, args=(0,)).start()
            
            return jsonify({
                'status': 'success', 
                'message': 'Power Automate triggered automatically after ROB download'
            })
        else:
            return jsonify({
                'status': 'error', 
                'message': 'No recent ROB download detected'
            })
    except Exception as e:
        return jsonify({
            'status': 'error', 
            'message': f'Error: {str(e)}'
        })'''

# ============================================================================
# WEEKLY REPORT ROUTES
# ============================================================================

@app.route('/weekly_report', methods=['GET', 'POST'])
def weekly_report():
    if request.method == 'POST':
        form_type = request.form.get('form_type')
        
        if form_type == 'backend_processing':
            return handle_backend_processing()
        else:
            return handle_rid_analysis()
    
    # GET request - show form (no data to display)
    return render_template('weekly_report.html', qualified_rids=None, filter_summary=None, backend_result=None)

def handle_rid_analysis():
    """Handle RID analysis with dual file input - ranking sheet + cleaned ROB file"""
    try:
        print("RID Analysis POST request received!")
        
        # Get filter parameters from form
        min_search_volume = int(request.form.get('min_search_volume', 5000))
        competition_level = request.form.get('competition_level', 'Low')
        analyze_trends = request.form.get('analyze_trends') == 'on'
        
        print(f"User Filters: Search >= {min_search_volume}, Competition = {competition_level}")
        print(f"Google Trends: {'Enabled' if analyze_trends else 'Disabled'}")
        
        # Validate form inputs
        if not min_search_volume or min_search_volume < 0:
            flash('❌ Please enter a valid minimum search volume!')
            return redirect(request.url)
            
        if not competition_level:
            flash('❌ Please select a competition level!')
            return redirect(request.url)
        
        # Handle RANKING SHEET upload
        ranking_file = request.files.get('ranking_file')
        if not ranking_file or ranking_file.filename == '':
            flash('❌ Please select a ranking Excel file!')
            return redirect(request.url)

        if not allowed_file(ranking_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ranking sheet!')
            return redirect(request.url)

        # Handle CLEANED ROB FILE upload
        rob_file = request.files.get('cleaned_rob_file')
        if not rob_file or rob_file.filename == '':
            flash('❌ Please select a cleaned ROB Excel file!')
            return redirect(request.url)

        if not allowed_file(rob_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ROB file!')
            return redirect(request.url)

        # Save both uploaded files
        ranking_filename = secure_filename(ranking_file.filename)
        ranking_path = os.path.join(app.config['UPLOAD_FOLDER'], ranking_filename)
        ranking_file.save(ranking_path)
        print(f"Ranking file saved: {ranking_path}")
        
        rob_filename = secure_filename(rob_file.filename)
        rob_path = os.path.join(app.config['UPLOAD_FOLDER'], rob_filename)
        rob_file.save(rob_path)
        print(f"ROB file saved: {rob_path}")
        
        # Process both files and get qualified ROB data
        result_summary = process_dual_files_and_extract_rob(
            ranking_path, rob_path, min_search_volume, competition_level, analyze_trends
        )
        
        # Format success/warning messages based on results
        if result_summary['success']:
            flash(f'✅ Success! Found {result_summary["qualified_rids_count"]} qualified RIDs')
            flash(f'✅ Extracted {result_summary["matched_rob_rows"]} matching ROB rows')
            flash(f'📁 Weekly ROB.xlsx saved to Desktop/RPA folder!')
            print(f"Dual file processing completed: {result_summary}")
        else:
            flash(f'❌ Error: {result_summary.get("error", "Unknown error")}')
            result_summary = None
        
        # Clean up uploaded files after processing
        try:
            os.remove(ranking_path)
            os.remove(rob_path)
            print(f"Cleaned up uploaded files")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up files: {cleanup_error}")
        
        # Render template with results
        return render_template('weekly_report.html', 
                              qualified_rids=result_summary.get('qualified_rids', []) if result_summary else [],
                              filter_summary=result_summary.get('filter_summary', {}) if result_summary else {},
                              backend_result=None,
                              rob_extraction_result=result_summary)
        
    except ValueError as ve:
        print(f"Value Error: {ve}")
        flash('❌ Invalid input values. Please check your filters.')
        return redirect(request.url)
    except Exception as e:
        print(f"Error: {e}")
        flash(f'❌ Error processing files: {str(e)}')
        return redirect(request.url)

def process_dual_files_and_extract_rob(ranking_path, rob_path, min_search_volume, competition_level, analyze_trends):
    """Process ranking sheet and ROB file together, extract matching rows"""
    try:
        print(f"\n=== PROCESSING DUAL FILES ===")
        print(f"Ranking file: {ranking_path}")
        print(f"ROB file: {rob_path}")
        
        # STEP 1: Process ranking sheet to get qualified RIDs
        print("\n📊 STEP 1: Processing ranking sheet...")
        qualified_rids, filter_summary, updated_ranking_path = get_qualified_rids_and_remove_trending(
            ranking_path, min_search_volume, competition_level, analyze_trends
        )
        
        if not qualified_rids:
            return {
                'success': False,
                'error': 'No qualified RIDs found in ranking sheet with your filter criteria'
            }
        
        print(f"✅ Found {len(qualified_rids)} qualified RIDs: {qualified_rids}")
        
        # STEP 2: Process ROB file and extract matching rows
        print(f"\n📋 STEP 2: Processing ROB file and extracting matching rows...")
        
        # Read the cleaned ROB file
        if rob_path.endswith('.csv'):
            rob_df = pd.read_csv(rob_path)
        else:
            rob_df = pd.read_excel(rob_path, engine='openpyxl')
        
        total_rob_rows = len(rob_df)
        print(f"ROB file loaded: {total_rob_rows} rows")
        print(f"ROB file columns: {list(rob_df.columns)}")
        
        # Find Report ID column
        report_id_column = None
        possible_columns = ['Report ID', 'ReportID', 'report_id', 'ID', 'Report_ID', 'Market Name']
        
        for col in possible_columns:
            if col in rob_df.columns:
                report_id_column = col
                break
        
        if not report_id_column:
            return {
                'success': False,
                'error': f'Report ID column not found in ROB file. Available columns: {list(rob_df.columns)}'
            }
        
        print(f"Using Report ID column: {report_id_column}")
        
        # Convert qualified_rids to same type as Report ID column
        rob_df[report_id_column] = rob_df[report_id_column].astype(str).str.strip()
        qualified_rids_str = [str(rid).strip() for rid in qualified_rids]
        
        print(f"Sample Report IDs in ROB file: {rob_df[report_id_column].head().tolist()}")
        print(f"Looking for RIDs: {qualified_rids_str}")
        
        # Filter ROB rows that match qualified RIDs
        matching_rob_rows = rob_df[rob_df[report_id_column].isin(qualified_rids_str)].copy()
        matched_count = len(matching_rob_rows)
        
        print(f"Found {matched_count} matching ROB rows")
        
        if matched_count == 0:
            return {
                'success': False,
                'error': f'No matching Report IDs found in ROB file. Check if Report IDs {qualified_rids} exist in the ROB file.'
            }
        
        # Show which RIDs were found and missing
        found_rids = matching_rob_rows[report_id_column].tolist()
        missing_rids = [rid for rid in qualified_rids_str if rid not in found_rids]
        
        print(f"Found Report IDs: {found_rids}")
        if missing_rids:
            print(f"Missing Report IDs: {missing_rids}")
        
        # STEP 3: Save to Desktop/RPA folder as ROB.xlsx
        print(f"\n💾 STEP 3: Saving to Desktop...")
        
        # Create RPA folder on Desktop if it doesn't exist
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
            print(f"Created RPA folder: {rpa_folder}")
        
        # Save as ROB.xlsx in Desktop/RPA folder
        output_path = os.path.join(rpa_folder, "weekly_RID.xlsx")
        
        # Use xlsxwriter for better performance
        with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
            matching_rob_rows.to_excel(writer, index=False, sheet_name='ROB_Data')
        
        print(f"✅ weekly_RID.xlsx saved to: {output_path}")
        
        # Display sample of extracted data
        print("\nSample of extracted ROB data:")
        print(matching_rob_rows.head(2))
        
        # Create comprehensive summary
        summary = {
            'success': True,
            'qualified_rids': qualified_rids,
            'qualified_rids_count': len(qualified_rids),
            'total_rob_rows': total_rob_rows,
            'matched_rob_rows': matched_count,
            'found_rids': found_rids,
            'missing_rids': missing_rids,
            'output_path': output_path,
            'report_id_column': report_id_column,
            'filter_summary': filter_summary
        }
        
        return summary
        
    except Exception as e:
        print(f"Error in dual file processing: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def handle_backend_processing():
    """Handle backend file processing form submission"""
    try:
        print("Backend Processing POST request received!")
        
        # Get processing options
        auto_detect_header = request.form.get('auto_detect_header') == 'on'
        clean_columns = request.form.get('clean_columns') == 'on'
        remove_empty_rows = request.form.get('remove_empty_rows') == 'on'
        
        print(f"Processing options: Header={auto_detect_header}, Clean={clean_columns}, Remove Empty={remove_empty_rows}")
        
        # Handle backend file upload
        backend_file = request.files.get('backend_file')
        if not backend_file or backend_file.filename == '':
            flash('❌ Please select a backend Excel file!')
            return redirect(request.url)

        if not allowed_backend_file(backend_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) are allowed for backend processing!')
            return redirect(request.url)

        # Save uploaded backend file
        backend_filename = secure_filename(backend_file.filename)
        backend_path = os.path.join(app.config['UPLOAD_FOLDER'], backend_filename)
        backend_file.save(backend_path)
        print(f"Backend file saved: {backend_path}")
        
        # Process the backend file
        backend_result = process_backend_file(
            backend_path, 
            auto_detect_header=auto_detect_header, 
            clean_columns=clean_columns, 
            remove_empty_rows=remove_empty_rows
        )
        
        # Clean up uploaded file after processing
        try:
            os.remove(backend_path)
            print(f"Cleaned up backend file: {backend_path}")
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up backend file {backend_path}: {cleanup_error}")
        
        # Format success/error messages
        if backend_result['success']:
            flash(f'✅ Backend file processed successfully!')
            flash(f'📁 Processed {backend_result["final_rows"]} rows from {backend_result["original_rows"]} original rows')
            flash(f'📥 ROB.xlsx ready for download!')
        else:
            flash(f'❌ Backend processing failed: {backend_result["error"]}')
        
        # Render template with backend results
        return render_template('weekly_report.html', 
                             qualified_rids=None,
                             filter_summary=None,
                             backend_result=backend_result)
        
    except Exception as e:
        print(f"Backend processing error: {e}")
        flash(f'❌ Error processing backend file: {str(e)}')
        return redirect(request.url)

def get_qualified_rids_and_remove_trending(file_path, min_search_volume, competition_level, analyze_trends=False):
    """Apply custom filters, get qualified RIDs with priority system and 1500 keyword limit"""
    try:
        print(f"Processing file: {file_path}")
        
        # Read the file
        if file_path.endswith('.csv'):
            df = pd.read_csv(file_path)
        else:
            df = pd.read_excel(file_path, engine='openpyxl')
        
        original_count = len(df)
        print(f"Original data loaded: {original_count} rows")
        
        # Validate required columns exist
        required_columns = ['AVG. Search', 'Competition', 'RID']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            raise ValueError(f"Missing required columns: {missing_columns}")
        
        # STEP 1: Apply search volume filter first
        search_filtered_df = df[df['AVG. Search'] >= min_search_volume].copy()
        print(f"After search volume filter (>= {min_search_volume:,}): {len(search_filtered_df)} rows")
        
        # STEP 2: Apply priority-based competition filter
        print(f"Applying priority-based competition filter...")
        
        if competition_level == 'All':
            filtered_df = search_filtered_df
            print(f"No competition filter applied")
            competition_display = "All levels"
        else:
            # Priority order: Low -> Medium -> High
            priority_order = ['Low', 'Medium', 'High']
            
            # Get the selected level and all levels up to that priority
            if competition_level in priority_order:
                selected_index = priority_order.index(competition_level)
                allowed_levels = priority_order[:selected_index + 1]  # Include all levels up to selected
                
                print(f"Competition priority: {' → '.join(allowed_levels)} (up to {competition_level})")
                
                # Filter for allowed competition levels
                filtered_df = search_filtered_df[
                    search_filtered_df['Competition'].isin(allowed_levels)
                ].copy()
                
                # Sort by priority order (Low first, then Medium, then High) and search volume
                priority_map = {'Low': 1, 'Medium': 2, 'High': 3}
                filtered_df['competition_priority'] = filtered_df['Competition'].map(priority_map)
                filtered_df = filtered_df.sort_values(['competition_priority', 'AVG. Search'], ascending=[True, False])
                filtered_df = filtered_df.drop('competition_priority', axis=1)
                
                competition_display = f"Priority: {' → '.join(allowed_levels)}"
            else:
                # Fallback to exact match if level not in priority order
                filtered_df = search_filtered_df[
                    search_filtered_df['Competition'] == competition_level
                ].copy()
                competition_display = competition_level
        
        filtered_count = len(filtered_df)
        print(f"After applying priority competition filter: {filtered_count} rows")
        
        # STEP 3: Apply 1500 keyword limit BEFORE Google Trends
        max_keywords_for_trends = 1500
        
        if filtered_count > max_keywords_for_trends:
            print(f"⚠️  Too many keywords ({filtered_count}) for Google Trends analysis!")
            print(f"🔪 Limiting to top {max_keywords_for_trends} keywords (sorted by priority & search volume)")
            
            # Take top 1500 keywords (already sorted by priority and search volume)
            filtered_df = filtered_df.head(max_keywords_for_trends).copy()
            filtered_count = len(filtered_df)
            
            print(f"✅ Limited to {filtered_count} keywords for processing")
        
        # Create filter summary with priority info and limit info
        filter_summary = {
            'min_search': f"{min_search_volume:,}",
            'competition': competition_display,
            'original_count': original_count,
            'filtered_count': filtered_count,
            'trends_enabled': analyze_trends,
            'keyword_limit_applied': filtered_count == max_keywords_for_trends,
            'max_keywords_limit': max_keywords_for_trends
        }
        
        updated_file_path = None
        
        if filtered_count == 0:
            print("No records match the filter criteria")
            return [], filter_summary, updated_file_path
        
        if analyze_trends:
            # Run Google Trends analysis on filtered data (max 1500 keywords)
            print("🔥 Running Google Trends analysis on filtered keywords...")
            print(f"📊 Processing {filtered_count} keywords (within 1500 limit)")
            
            # Check if API key is configured
            if not GOOGLE_TRENDS_CONFIG.get('API_KEY') or GOOGLE_TRENDS_CONFIG['API_KEY'] == 'YOUR_API_KEY_HERE':
                print("⚠️ No API key configured - returning all filtered RIDs")
                qualified_rids = filtered_df['RID'].tolist()
                return qualified_rids, filter_summary, updated_file_path
            
            # Run actual Google Trends analysis
            keywords_data = filtered_df.to_dict('records')
            trending_data = analyze_keywords_with_google_trends(keywords_data)
            qualified_rids = [item['RID'] for item in trending_data if 'RID' in item]
            
            print(f"Google Trends analysis complete: {len(qualified_rids)} trending RIDs out of {filtered_count} filtered")
            filter_summary['trends_qualified'] = len(qualified_rids)
            filter_summary['trends_message'] = f"After Google Trends analysis: {len(qualified_rids)} out of {filtered_count} keywords are trending"
            
            # Remove trending RIDs from original dataframe
            if qualified_rids:
                print(f"🗑️ Removing {len(qualified_rids)} trending RIDs from ranking sheet...")
                
                # Create a copy of original dataframe
                df_updated = df.copy()
                
                # Remove rows where RID is in the qualified_rids list
                df_updated = df_updated[~df_updated['RID'].isin(qualified_rids)]
                
                rows_removed = len(df) - len(df_updated)
                print(f"✅ Removed {rows_removed} trending rows from ranking sheet")
                
                # Save the updated ranking sheet in C:\Users\abhijit\Desktop\RPA\Ranking sheet\<year>\<month>\updated_ranking_sheet_<year>_<month>_<day>.xlsx
                today = datetime.today()
                ranking_base_folder = os.path.join(r"C:\Users\abhijit\Desktop\RPA\Ranking sheet", str(today.year), f"{today.month:02d}")
                os.makedirs(ranking_base_folder, exist_ok=True)
                updated_filename = f"updated_ranking_sheet_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
                updated_file_path = os.path.join(ranking_base_folder, updated_filename)
                # Use xlsxwriter for better performance
                with pd.ExcelWriter(updated_file_path, engine='xlsxwriter') as writer:
                    df_updated.to_excel(writer, index=False, sheet_name='Sheet1')
                print(f"💾 Updated ranking sheet saved as: {updated_file_path}")
                flash(f'💾 Updated ranking sheet saved as: {updated_file_path}')
                
                # Update filter summary with removal info
                filter_summary['rows_removed'] = rows_removed
                filter_summary['final_sheet_rows'] = len(df_updated)
                filter_summary['removal_message'] = f"Removed {rows_removed} trending rows. Updated sheet has {len(df_updated)} rows."
            else:
                print("ℹ️ No trending RIDs found - ranking sheet unchanged")
                filter_summary['removal_message'] = "No trending RIDs found - ranking sheet unchanged"
            
        else:
            # No Google Trends - return all filtered RIDs (max 1500)
            qualified_rids = filtered_df['RID'].tolist()
            print(f"Returning all filtered RIDs: {len(qualified_rids)} RIDs")
            filter_summary['trends_message'] = "Google Trends analysis disabled - showing all filtered results"
            
            if filter_summary['keyword_limit_applied']:
                filter_summary['trends_message'] += f" (limited to top {max_keywords_for_trends})"
        
        return qualified_rids, filter_summary, updated_file_path
        
    except Exception as e:
        print(f"Error in get_qualified_rids_and_remove_trending: {e}")
        raise e

def process_backend_file(file_path, auto_detect_header=True, clean_columns=True, remove_empty_rows=True):
    """Process large backend file directly to ROB format with optimization"""
    try:
        print(f"\n=== PROCESSING LARGE BACKEND FILE TO ROB FORMAT ===")
        print(f"Processing file: {file_path}")
        
        # Step 1: Read the file with optimization for large files
        try:
            # Try reading with openpyxl engine for better large file handling
            df_raw = pd.read_excel(file_path, header=None, engine='openpyxl')
        except Exception as e:
            print(f"Error with openpyxl, trying alternative: {e}")
            # Fallback to default engine
            df_raw = pd.read_excel(file_path, header=None)
        
        original_rows = df_raw.shape[0]
        print(f"Initial raw data shape: {df_raw.shape}")
        
        # Step 2: Find the actual header row if auto-detect is enabled
        if auto_detect_header:
            header_row_index = find_header_row(df_raw)
        else:
            header_row_index = 0  # Assume first row is header
        
        if header_row_index is not None:
            # Set the header
            header = df_raw.iloc[header_row_index]
            # Drop rows before the header (inclusive)
            df_data = df_raw[header_row_index + 1:].copy()
            # Assign the correct header
            df_data.columns = header
            
            # Reset index
            df_data.reset_index(drop=True, inplace=True)
            
            print(f"Data extracted with header found at index {header_row_index}. New shape: {df_data.shape}")
            
            if clean_columns:
                # Clean column names (remove leading/trailing spaces, handle duplicates)
                df_data.columns = df_data.columns.str.strip()
                
                # Handle duplicate columns
                cols = pd.Series(df_data.columns)
                for dup in cols[cols.duplicated()].unique():
                    cols[cols[cols == dup].index.values.tolist()] = [dup + '_' + str(i) if i != 0 else dup for i in range(sum(cols == dup))]
                df_data.columns = cols
                
                print("Columns cleaned.")
            
            if remove_empty_rows:
                # Drop rows that are entirely null after extraction
                initial_rows = df_data.shape[0]
                df_data.dropna(how='all', inplace=True)
                rows_dropped = initial_rows - df_data.shape[0]
                print(f"Dropped {rows_dropped} empty rows.")
            
            # Save the processed file as "ROB.xlsx" directly
            output_path = os.path.join(app.config['UPLOAD_FOLDER'], 'ROB.xlsx')
            
            # Use xlsxwriter engine for better performance with large files
            with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
                df_data.to_excel(writer, index=False, sheet_name='Sheet1')
            
            print(f"ROB file saved as: {output_path}")
            
            # Display sample of processed data
            print(f"\nProcessed {len(df_data)} records successfully")
            print("Sample of processed data:")
            print(df_data.head(2).to_string())
            
            # Create summary
            summary = {
                'success': True,
                'original_rows': original_rows,
                'final_rows': len(df_data),
                'header_row': header_row_index,
                'final_columns': len(df_data.columns),
                'output_file': 'ROB.xlsx'
            }
            
            return summary
        
        else:
            print("Could not automatically detect header row.")
            return {
                'success': False,
                'error': 'Could not automatically detect header row. Please check your file format.'
            }
    
    except MemoryError:
        print("Memory error - file too large")
        return {
            'success': False,
            'error': 'File too large to process. Please try with a smaller file or contact support.'
        }
    except Exception as e:
        print(f"Error processing backend file: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def find_header_row(df):
    """Heuristic function to find the header row"""
    for index, row in df.iterrows():
        if sum(isinstance(x, str) for x in row) >= 5:
            print(f"Potential header row found at index: {index}")
            return index
    return None

@app.route('/download_backend_file')
def download_backend_file():
    """Download the processed ROB file"""
    try:
        filename = 'ROB.xlsx'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            flash('❌ Processed file not found. Please process a backend file first.')
            return redirect(url_for('weekly_report'))
    except Exception as e:
        flash(f'❌ Error downloading file: {str(e)}')
        return redirect(url_for('weekly_report'))

@app.route('/download_updated_ranking')
def download_updated_ranking():
    """Download the updated ranking sheet (with trending rows removed)"""
    try:
        filename = 'updated_ranking_sheet.xlsx'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name='ranking_sheet_trending_removed.xlsx')
        else:
            flash('❌ Updated ranking sheet not found. Please run Google Trends analysis first.')
            return redirect(url_for('weekly_report'))
    except Exception as e:
        flash(f'❌ Error downloading updated ranking sheet: {str(e)}')
        return redirect(url_for('weekly_report'))

# ============================================================================
# GOOGLE TRENDS FUNCTIONS - UPDATED TO USE NEW EXTRACTOR
# ============================================================================

# Updated Google Trends Config
GOOGLE_TRENDS_CONFIG = {
    'API_KEY': '6891e2b1972e1d06afb85c44',  # ScrapingDog API key
    'INTEREST_THRESHOLD': 50,
    'DAYS_ABOVE_THRESHOLD': 2,
    'TERMS_TO_REMOVE': ['market', 'size', 'analysis', 'report', 'industry', 'global'],
    'REQUEST_DELAY': 2,
    'RETRY_DELAY': 10,   # NEW: Delay when retrying failed requests
    'MAX_RETRIES': 3,    # NEW: Maximum retry attempts
    'BATCH_SIZE': 50,    # NEW: Process in smaller batches
    'BATCH_DELAY': 30
}

class GoogleTrendsExtractor:
    def __init__(self, api_key):
        self.api_key = api_key
        self.base_url = "https://api.scrapingdog.com/google_trends"
        self.consecutive_failures = 0  # Track consecutive failures
    
    def get_values(self, keyword, retry_count=0):
        """Get exact values for last 7 days with enhanced error handling"""
        params = {
            "api_key": self.api_key,
            "query": keyword,
            "geo": "US",
            "tz": "330",
            "date": "now 7-d",
            "data_type": "TIMESERIES"
        }
        
        try:
            # Add timeout to prevent hanging connections
            response = requests.get(self.base_url, params=params, timeout=30)
            
            if response.status_code == 200:
                data = response.json()
                values = self.extract_values(data)
                self.consecutive_failures = 0  # Reset failure counter on success
                return values
                
            elif response.status_code == 429:  # Rate limit error
                print(f"    ⚠️ Rate limit hit (429). Backing off...")
                self.handle_rate_limit(retry_count)
                if retry_count < GOOGLE_TRENDS_CONFIG['MAX_RETRIES']:
                    return self.get_values(keyword, retry_count + 1)
                else:
                    print(f"    ❌ Max retries exceeded for: {keyword}")
                    return []
                    
            else:
                print(f"    API Error: {response.status_code}")
                print(f"    Response: {response.text[:200]}...")
                return []
                
        except requests.exceptions.ConnectionError as e:
            print(f"    🔌 Connection Error: {e}")
            self.consecutive_failures += 1
            
            if retry_count < GOOGLE_TRENDS_CONFIG['MAX_RETRIES']:
                delay = GOOGLE_TRENDS_CONFIG['RETRY_DELAY'] * (retry_count + 1)
                print(f"    ⏳ Retrying in {delay} seconds... (attempt {retry_count + 1})")
                time.sleep(delay)
                return self.get_values(keyword, retry_count + 1)
            else:
                print(f"    ❌ Max retries exceeded for: {keyword}")
                return []
                
        except requests.exceptions.Timeout as e:
            print(f"    ⏰ Timeout Error: {e}")
            if retry_count < GOOGLE_TRENDS_CONFIG['MAX_RETRIES']:
                print(f"    ⏳ Retrying with longer timeout...")
                time.sleep(GOOGLE_TRENDS_CONFIG['RETRY_DELAY'])
                return self.get_values(keyword, retry_count + 1)
            return []
            
        except Exception as e:
            print(f"    ❌ Unexpected Error: {e}")
            return []
    
    def handle_rate_limit(self, retry_count):
        """Handle rate limiting with exponential backoff"""
        base_delay = GOOGLE_TRENDS_CONFIG['RETRY_DELAY']
        # Exponential backoff: 10s, 20s, 40s
        delay = base_delay * (2 ** retry_count)
        print(f"    ⏳ Rate limit backoff: Waiting {delay} seconds...")
        time.sleep(delay)
    
    def extract_values(self, data):
        """Extract values using standard timeline method"""
        values = []
        
        try:
            if 'interest_over_time' in data:
                timeline_data = data['interest_over_time'].get('timeline_data', [])
                
                for entry in timeline_data:
                    if isinstance(entry, dict) and 'values' in entry:
                        for val_item in entry['values']:
                            if isinstance(val_item, dict) and 'value' in val_item:
                                try:
                                    val = int(val_item['value'])
                                    if 0 <= val <= 100:
                                        values.append(val)
                                except (ValueError, TypeError):
                                    pass
        except Exception:
            pass
        
        return values
    
    def filter_keyword(self, keyword, threshold=1):
        """Check if keyword has threshold+ values > 50 in last 7 days"""
        values = self.get_values(keyword)
        
        if values:
            count_above_50 = sum(1 for val in values if val > 50)
            print(f"    Values: {values[:10]}{'...' if len(values) > 10 else ''} | Count >50: {count_above_50} | Threshold: {threshold}+")
            return count_above_50 >= threshold
        
        print(f"    No values retrieved | Threshold: {threshold}+")
        return False

def analyze_keywords_with_google_trends_batched(keywords_data):
    """Analyze keywords with Google Trends using batch processing to prevent rate limiting"""
    
    if not keywords_data:
        return []
    
    api_key = GOOGLE_TRENDS_CONFIG['API_KEY']
    extractor = GoogleTrendsExtractor(api_key)
    
    total_keywords = len(keywords_data)
    batch_size = GOOGLE_TRENDS_CONFIG['BATCH_SIZE']
    qualifying_keywords = []
    
    print(f"🔥 Starting BATCHED Google Trends Analysis")
    print(f"📊 Total keywords: {total_keywords}")
    print(f"📦 Batch size: {batch_size}")
    print(f"⏱️ Delay per request: {GOOGLE_TRENDS_CONFIG['REQUEST_DELAY']}s")
    print(f"⏱️ Delay between batches: {GOOGLE_TRENDS_CONFIG['BATCH_DELAY']}s")
    
    # Process in batches
    for batch_num in range(0, total_keywords, batch_size):
        batch_end = min(batch_num + batch_size, total_keywords)
        batch_keywords = keywords_data[batch_num:batch_end]
        
        print(f"\n📦 BATCH {batch_num//batch_size + 1}: Processing keywords {batch_num + 1}-{batch_end}")
        print("-" * 60)
        
        # Process each keyword in the batch
        for i, keyword_row in enumerate(batch_keywords):
            global_index = batch_num + i + 1
            
            try:
                original_keyword = keyword_row.get('Keywords', '')
                rid = keyword_row.get('RID', '')
                competition = keyword_row.get('Competition', '')
                search_volume = keyword_row.get('AVG. Search', 0)
                
                if not original_keyword or not rid:
                    print(f"[{global_index}/{total_keywords}] Skipping row with missing data")
                    continue
                    
                clean_keyword = clean_keyword_for_trends(original_keyword)
                
                if not clean_keyword:
                    print(f"[{global_index}/{total_keywords}] Skipping empty keyword after cleaning")
                    continue
                
                print(f"[{global_index}/{total_keywords}] Analyzing RID {rid}: '{clean_keyword}' [{competition}, {search_volume:,}]")
                
                # Check if keyword qualifies
                is_trending = extractor.filter_keyword(clean_keyword)
                
                if is_trending:
                    qualifying_keywords.append({
                        'RID': rid, 
                        'keyword': original_keyword,
                        'competition': competition,
                        'search_volume': search_volume
                    })
                    print(f"  ✅ TRENDING: RID {rid}")
                else:
                    print(f"  ❌ Not trending: RID {rid}")
                
                # Standard delay between requests
                time.sleep(GOOGLE_TRENDS_CONFIG['REQUEST_DELAY'])
                
            except Exception as e:
                print(f"  ❌ Error analyzing RID {keyword_row.get('RID', 'unknown')}: {e}")
                continue
        
        # Batch completed - add delay before next batch (except for last batch)
        if batch_end < total_keywords:
            print(f"\n⏸️ Batch {batch_num//batch_size + 1} complete. Resting for {GOOGLE_TRENDS_CONFIG['BATCH_DELAY']} seconds...")
            time.sleep(GOOGLE_TRENDS_CONFIG['BATCH_DELAY'])
    
    print(f"\n🎯 Batched Analysis Complete!")
    print(f"📈 Result: {len(qualifying_keywords)} out of {total_keywords} keywords are trending")
    
    return qualifying_keywords


def analyze_keywords_with_google_trends(keywords_data):
    """Main entry point - now uses batched processing"""
    return analyze_keywords_with_google_trends_batched(keywords_data)

def check_google_trends_api_health():
    """Quick health check for Google Trends API"""
    print("🏥 Checking Google Trends API health...")
    
    api_key = GOOGLE_TRENDS_CONFIG['API_KEY']
    extractor = GoogleTrendsExtractor(api_key)
    
    # Test with a simple keyword
    test_values = extractor.get_values("technology")
    
    if test_values:
        print("✅ API is healthy and responding")
        return True
    else:
        print("❌ API is not responding properly")
        return False


# ADD this function to gracefully handle interruptions
def analyze_keywords_with_checkpoint(keywords_data, checkpoint_file="trends_checkpoint.txt"):
    """Analyze keywords with checkpoint support to resume from interruptions"""
    
    # Load checkpoint if exists
    start_index = 0
    if os.path.exists(checkpoint_file):
        try:
            with open(checkpoint_file, 'r') as f:
                start_index = int(f.read().strip())
            print(f"📍 Resuming from checkpoint: keyword {start_index + 1}")
        except:
            start_index = 0
    
    # Process from checkpoint
    remaining_keywords = keywords_data[start_index:]
    qualifying_keywords = []
    
    try:
        for i, keyword_row in enumerate(remaining_keywords):
            current_index = start_index + i
            
            # Process keyword (your existing logic here)
            # ... processing code ...
            
            # Update checkpoint every 10 keywords
            if (current_index + 1) % 10 == 0:
                with open(checkpoint_file, 'w') as f:
                    f.write(str(current_index + 1))
                print(f"💾 Checkpoint saved at keyword {current_index + 1}")
        
        # Clean up checkpoint file when complete
        if os.path.exists(checkpoint_file):
            os.remove(checkpoint_file)
            print("🧹 Checkpoint file cleaned up")
            
    except KeyboardInterrupt:
        print(f"\n⏸️ Process interrupted. Checkpoint saved at keyword {current_index + 1}")
        print("💡 Run again to resume from this point")
        raise
    
    return qualifying_keywords

def clean_keyword_for_trends(keyword):
    """Clean keyword by removing problematic terms"""
    if not keyword:
        return ""
        
    cleaned = str(keyword)
    
    # Remove terms from config
    for term in GOOGLE_TRENDS_CONFIG['TERMS_TO_REMOVE']:
        cleaned = re.sub(rf'\b{re.escape(term)}\b', '', cleaned, flags=re.IGNORECASE)
    
    # Clean up extra spaces and trim
    cleaned = ' '.join(cleaned.split()).strip()
    return cleaned

# ============================================================================
# POWER AUTOMATE ROUTES
# ============================================================================

@app.route('/wait_power_automate')
def wait_power_automate():
    """Show a waiting page for Power Automate Desktop step."""
    return render_template('wait_power_automate.html')

@app.route('/api/trigger_power_automate', methods=['POST'])
def trigger_power_automate_flow():
    """Triggers a Power Automate Desktop flow"""
    pad_exe_path = r"C:\Program Files (x86)\Power Automate Desktop\PAD.Console.Host.exe"
    flow_name = "Paid PR - Files Downloader"
    
    if not os.path.exists(pad_exe_path):
        print("Power Automate Desktop executable not found!")
        return jsonify({'status': 'error', 'message': 'PAD executable not found'})
    
    command = f'"{pad_exe_path}" -flow "{flow_name}"'
    
    try:
        result = subprocess.run(command, shell=True, check=True, text=True, capture_output=True)
        print(f"Flow triggered successfully. Output: {result.stdout}")

        time.sleep(5)
        
        flow_button_coordinates = (463, 395)
        print(f"Clicking at {flow_button_coordinates}")
        pyautogui.click(flow_button_coordinates)
        print("Flow triggered successfully.")

    except subprocess.CalledProcessError as e:
        print(f"Error triggering flow: {e.stderr}")
        return jsonify({'status': 'error', 'message': f'Flow error: {e.stderr}'})
    
    return jsonify({'status': 'success', 'message': 'Power Automate process completed.'})

# ============================================================================
#  Custom APPLICATION RUNNER
# ============================================================================
@app.route('/custom_index.html')
def custom_index():
    """Render custom index page for application"""
    return render_template('custom_index.html')

@app.route('/custom_weekly_report', methods=['GET', 'POST'])
def custom_weekly_report():
    if request.method == 'POST':
        try:
            print("Custom Weekly RID Analysis POST request received!")
            
            # Get ALL parameters from form (don't hardcode!)
            pr_count = int(request.form.get('pr_count', 200))
            min_search_volume = int(request.form.get('min_search_volume', 5000))
            competition_level = request.form.get('competition_level', 'Medium')
            analyze_trends = request.form.get('analyze_trends') == 'on'  # Use form value
            
            print(f"User wants {pr_count} PRs this week")
            print(f"Filters: Search >= {min_search_volume}, Competition = {competition_level}")
            print(f"Google Trends: {'Enabled' if analyze_trends else 'Disabled'}")
            
            # Validate PR count
            if not pr_count or pr_count <= 0:
                flash('❌ Please enter a valid number of PRs!')
                return redirect(request.url)
            
            # Handle RANKING SHEET upload
            ranking_file = request.files.get('ranking_file')
            if not ranking_file or ranking_file.filename == '':
                flash('❌ Please select a ranking Excel file!')
                return redirect(request.url)

            if not allowed_file(ranking_file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed!')
                return redirect(request.url)

            # Save uploaded file
            ranking_filename = secure_filename(ranking_file.filename)
            ranking_path = os.path.join(app.config['UPLOAD_FOLDER'], ranking_filename)
            ranking_file.save(ranking_path)
            print(f"Ranking file saved: {ranking_path}")
            
            # FIXED: Call correct function name with correct parameters
            result_summary = process_custom_weekly_ranking_file(
                ranking_path, min_search_volume, competition_level, analyze_trends, pr_count
            )
            
            # Format success/warning messages
            if result_summary['success']:
                flash(f'✅ Success! Processed {result_summary["qualified_rids_count"]} qualified RIDs')
                flash(f'📁 Custom ROB saved: {result_summary["output_path"]}')
                flash(f'📁 Updated ranking saved: {result_summary["updated_ranking_path"]}')
                print(f"Custom weekly workflow completed: {result_summary}")
            else:
                flash(f'❌ Error: {result_summary.get("error", "Unknown error")}')
                result_summary = None
            
            # Clean up uploaded file
            try:
                os.remove(ranking_path)
                print(f"Cleaned up uploaded file")
            except Exception as cleanup_error:
                print(f"Warning: Could not clean up file: {cleanup_error}")
            
            return render_template('custom_weekly_report.html', 
                                  qualified_rids=result_summary.get('qualified_rids', []) if result_summary else [],
                                  filter_summary=result_summary.get('filter_summary', {}) if result_summary else {},
                                  custom_weekly_result=result_summary)
            
        except ValueError as ve:
            print(f"Value Error: {ve}")
            flash('❌ Invalid input values. Please check your filters.')
            return redirect(request.url)
        except Exception as e:
            print(f"Error: {e}")
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    # GET request - show custom weekly form
    return render_template('custom_weekly_report.html')
def process_custom_weekly_ranking_file(ranking_path, min_search_volume, competition_level, analyze_trends, pr_count):
    """Standalone function with Keyword + Domain + Sub Domain output"""
    try:
        print(f"\n=== CUSTOM WEEKLY RANKING FILE ===")
        print(f"Target PR Count: {pr_count}")
        print(f"Ranking file: {ranking_path}")
        print(f"Google Trends: {'Enabled' if analyze_trends else 'Disabled'}")
        
        # Step 1: Read the ranking file
        if ranking_path.endswith('.csv'):
            df_original = pd.read_csv(ranking_path)
        else:
            df_original = pd.read_excel(ranking_path, engine='openpyxl')
        
        print(f"✅ Original file loaded: {len(df_original)} rows")
        print(f"Columns: {list(df_original.columns)}")
        
        # ⭐ Updated required columns to include Sub Domain
        required_columns = ['RID', 'Keyword', 'AVG. Search', 'Competition', 'Sub Domain', 'Domain']
        missing_columns = [col for col in required_columns if col not in df_original.columns]
        if missing_columns:
            return {
                'success': False,
                'error': f'Missing required columns: {missing_columns}'
            }
        
        # Step 2: Apply search volume filter
        filtered_df = df_original[df_original['AVG. Search'] >= min_search_volume].copy()
        print(f"✅ After search volume filter (>= {min_search_volume:,}): {len(filtered_df)} rows")
        
        if len(filtered_df) == 0:
            return {
                'success': False,
                'error': f'No keywords found with search volume >= {min_search_volume}'
            }
        
        # Step 3: Apply competition filter with priority
        if competition_level == 'Low':
            allowed_levels = ['Low']
        elif competition_level == 'Medium':
            allowed_levels = ['Low', 'Medium']
        elif competition_level == 'High':
            allowed_levels = ['Low', 'Medium', 'High']
        else:  # 'All'
            allowed_levels = filtered_df['Competition'].unique().tolist()
        
        filtered_df = filtered_df[filtered_df['Competition'].isin(allowed_levels)].copy()
        print(f"✅ After competition filter ({competition_level}): {len(filtered_df)} rows")
        
        if len(filtered_df) == 0:
            return {
                'success': False,
                'error': f'No keywords found with competition level: {competition_level}'
            }
        
        # Step 4: Sort by priority (Low first, then by search volume)
        priority_map = {'Low': 1, 'Medium': 2, 'High': 3}
        filtered_df['priority'] = filtered_df['Competition'].map(priority_map).fillna(4)
        filtered_df = filtered_df.sort_values(['priority', 'AVG. Search'], ascending=[True, False])
        filtered_df = filtered_df.drop('priority', axis=1)
        
        # Step 5: ⭐ Apply Google Trends filtering (if enabled)
        if analyze_trends and len(filtered_df) > 0:
            print(f"🔥 Running Google Trends filtering with proper criteria...")
            print(f"   INTEREST_THRESHOLD: {GOOGLE_TRENDS_CONFIG['INTEREST_THRESHOLD']}")
            print(f"   DAYS_ABOVE_THRESHOLD: {GOOGLE_TRENDS_CONFIG['DAYS_ABOVE_THRESHOLD']}")
            
            # Use EXACT PR count for Google Trends analysis
            max_for_trends = min(len(filtered_df), pr_count)
            trends_input_df = filtered_df.head(max_for_trends).copy()
            print(f"   Analyzing EXACTLY {len(trends_input_df)} keywords with Google Trends (user requested {pr_count})")
            
            # Rename 'Keyword' to 'Keywords' for Google Trends function
            trends_input_df = trends_input_df.rename(columns={'Keyword': 'Keywords'})
            print(f"   Renamed 'Keyword' → 'Keywords' for Google Trends compatibility")
            
            try:
                # Use existing Google Trends logic from laststep.py
                keywords_data = trends_input_df.to_dict('records')
                
                # Debug: Print first few rows to verify data format
                print(f"   Sample data for Google Trends:")
                for i, row in enumerate(keywords_data[:3]):
                    print(f"     Row {i+1}: RID={row.get('RID')}, Keywords='{row.get('Keywords', 'MISSING')}'")
                
                # This function applies the GOOGLE_TRENDS_CONFIG criteria:
                qualifying_keywords = analyze_keywords_with_google_trends(keywords_data)
                
                print(f"📊 Google Trends Results:")
                print(f"   Input keywords: {len(keywords_data)}")
                print(f"   Qualifying keywords: {len(qualifying_keywords)}")
                
                if len(qualifying_keywords) > 0:
                    # Extract RIDs from qualifying keywords
                    trending_rids = [str(item['RID']).strip() for item in qualifying_keywords if 'RID' in item]
                    print(f"   Trending RIDs extracted: {len(trending_rids)}")
                    
                    # Use all trending keywords
                    final_rids = trending_rids
                    print(f"🎯 SUCCESS: Using {len(final_rids)} trending keywords from Google Trends!")
                    
                else:
                    print("❌ No keywords passed Google Trends criteria!")
                    # Fallback: use the same keywords we sent to Google Trends
                    final_rids = trends_input_df['RID'].astype(str).tolist()
                    print(f"🔄 Fallback: Using the {len(final_rids)} keywords we analyzed")
                
            except Exception as e:
                print(f"❌ Google Trends API error: {e}")
                import traceback
                traceback.print_exc()
                # Fallback to the keywords we tried to analyze
                final_rids = trends_input_df['RID'].astype(str).tolist()
                print(f"🔄 API error fallback: Using the {len(final_rids)} keywords we tried to analyze")
        
        else:
            # No Google Trends - just use top filtered keywords
            final_rids = filtered_df['RID'].astype(str).head(pr_count).tolist()
            print(f"🚫 Google Trends disabled - using top {len(final_rids)} filtered keywords")
        
        print(f"✅ Final RIDs selected: {len(final_rids)}")
        
        # Step 6: Get final data for output
        df_original['RID'] = df_original['RID'].astype(str).str.strip()
        final_rids_str = [str(rid).strip() for rid in final_rids]
        
        # Filter original data to get final qualified rows
        qualified_data = df_original[df_original['RID'].isin(final_rids_str)].copy()
        print(f"✅ Final qualified data: {len(qualified_data)} rows")
        
        if len(qualified_data) == 0:
            return {
                'success': False,
                'error': 'No qualified data found after processing'
            }
        
        # ⭐ Step 7: Create output with Keyword + Domain + Sub Domain columns
        output_data = qualified_data[['Keyword', 'Domain', 'Sub Domain']].copy()
        print(f"✅ Output data prepared: {len(output_data)} rows with Keyword + Domain + Sub Domain")
        print(f"Sample output data:\n{output_data.head()}")
        
        # Step 8: Save to Desktop/RPA folder
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        # Create RPA folder if it doesn't exist
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
            print(f"✅ Created RPA folder: {rpa_folder}")
        
        # Save output file
        output_path = os.path.join(rpa_folder, 'custom_weekly_rob.xlsx')
        output_data.to_excel(output_path, index=False, engine='xlsxwriter')
        
        # Verify file was created
        if os.path.exists(output_path):
            file_size = os.path.getsize(output_path)
            print(f"✅ SUCCESS: custom_weekly_rob.xlsx created!")
            print(f"   Path: {output_path}")
            print(f"   Size: {file_size} bytes")
            print(f"   Rows: {len(output_data)}")
        else:
            print(f"❌ ERROR: File was not created at {output_path}")
            return {
                'success': False,
                'error': f'Output file was not created at {output_path}'
            }
        
        # Step 9: Create updated ranking sheet (remove qualified RIDs)
        updated_df = df_original[~df_original['RID'].isin(final_rids_str)].copy()
        removed_count = len(df_original) - len(updated_df)
        print(f"✅ Removed {removed_count} qualified RIDs from ranking sheet")
        
        # Save updated ranking in date-wise folder
        today = datetime.today()
        ranking_folder = os.path.join(
            rpa_folder, "Custom Updated Ranking Sheet", 
            str(today.year), f"{today.month:02d}"
        )
        os.makedirs(ranking_folder, exist_ok=True)
        
        date_filename = f"Remaining_ranking_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
        final_ranking_path = os.path.join(ranking_folder, date_filename)
        updated_df.to_excel(final_ranking_path, index=False, engine='xlsxwriter')
        
        if os.path.exists(final_ranking_path):
            print(f"✅ Updated ranking saved: {final_ranking_path}")
        else:
            print(f"⚠️ Updated ranking file not created")
            final_ranking_path = "Failed to create updated ranking"
        
        # Return success summary
        return {
            'success': True,
            'qualified_rids': final_rids,
            'qualified_rids_count': len(final_rids),
            'output_path': output_path,
            'updated_ranking_path': final_ranking_path,
            'analysis_type': 'custom_weekly',
            'excel_rows': len(output_data),
            'pr_count_applied': pr_count,
            'output_columns': ['Keyword', 'Domain', 'Sub Domain'],  # ⭐ Updated
            'google_trends_used': analyze_trends,
            'original_rows': len(df_original),
            'filtered_rows': len(filtered_df),
            'final_rows': len(output_data)
        }
        
    except Exception as e:
        print(f"❌ ERROR in custom weekly ranking file processing: {e}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e)
        }
      
@app.route('/custom_choice')
def custom_choice():
    """Custom choice page for CMI/WMR selection"""
    return render_template('custom_choice.html')



@app.route('/custom_cmi_cta', methods=['GET', 'POST'])
def custom_cmi_cta():
    if request.method == 'POST':
        try:
            # Get form data (only extract_count now since user info fields are removed)
            extract_count = int(request.form.get('extract_count', 200))

            # Validate required fields (only extract_count and file now)
            if not extract_count or extract_count < 1:
                flash('❌ Please enter a valid number of keywords to extract!')
                return redirect(request.url)

            file = request.files.get('file')
            if not file or file.filename == '':
                flash('❌ Keywords file is required!')
                return redirect(request.url)

            if not allowed_file(file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed!')
                return redirect(request.url)

            # Use secure_filename to avoid path issues
            filename = secure_filename(file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(input_path)

            # Process the file for CMI CTA generation
            result = process_cmi_cta_file(input_path, extract_count)
            
            if result['success']:
                flash(f'✅ Successfully processed {result["extracted_count"]} keywords!')
                flash(f'📁 File saved: {result["filename"]}')
                flash('🤖 CMI automation started!')
                
                # Start CMI automation in background
                threading.Thread(target=run_cmi_automation).start()
            else:
                flash(f'❌ Error: {result["error"]}')

            return render_template('custom_cmi_cta.html')

        except ValueError as ve:
            flash('❌ Please enter a valid number for keywords to extract!')
            return redirect(request.url)
        except Exception as e:
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_cmi_cta.html')

def process_cmi_cta_file(file_path, extract_count):
    """Process Custom Weekly ROB file for CMI CTA generation"""
    try:
        # Read the file
        if file_path.endswith('.csv'):
            df_original = pd.read_csv(file_path)
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')

        total_rows = len(df_original)
        
        if total_rows < extract_count:
            extract_count = total_rows

        # Step 1: Extract top N rows for CTA generation
        extracted_rows = df_original.head(extract_count).copy()
        
        # Step 2: Get remaining rows (original minus extracted)
        remaining_rows = df_original.iloc[extract_count:].copy()
        
        # Create filename with current date for extracted data
        today = datetime.today()
        extracted_filename = f"custom_reports_cmi_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
        
        # Save extracted data to Desktop/RPA folder
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        # Save extracted keywords file
        extracted_output_path = os.path.join(rpa_folder, extracted_filename)
        extracted_rows.to_excel(extracted_output_path, index=False)
        
        # Step 3: Update the original Custom_weekly_ROB.xlsx file with remaining data
        custom_weekly_rob_path = os.path.join(rpa_folder, 'Custom_weekly_ROB.xlsx')
        
        if os.path.exists(custom_weekly_rob_path):
            # Update the original file with remaining rows
            remaining_rows.to_excel(custom_weekly_rob_path, index=False)
            print(f"✅ Updated Custom_weekly_ROB.xlsx - Removed {extract_count} extracted keywords")
            print(f"✅ Custom_weekly_ROB.xlsx now contains {len(remaining_rows)} remaining keywords")
        else:
            print(f"⚠️ Warning: Custom_weekly_ROB.xlsx not found at {custom_weekly_rob_path}")
        
        return {
            'success': True,
            'extracted_count': extract_count,
            'remaining_count': len(remaining_rows),
            'total_count': total_rows,
            'filename': extracted_filename,
            'output_path': extracted_output_path,
            'updated_original': os.path.exists(custom_weekly_rob_path)
        }
        
    except Exception as e:
        print(f"Error in process_cmi_cta_file: {e}")
        return {
            'success': False,
            'error': str(e)
        }


def run_cmi_automation():
    """Run CMI automation using Selenium"""
    try:
        print("Starting CMI automation...")
        
        from webdriver_manager.chrome import ChromeDriverManager
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        import time
        
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        #options.add_argument("--headless")  # Run headlessly for deployment
        options.add_argument("--start-maximized")
        
        # Initialize WebDriver
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.coherentmarketinsights.com/cmisitmanup/index.php')
        
        username_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[1]/input'))
        )
        username_input.send_keys('Auto_Ops_Team')
        
        password_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[2]/input'))
        )
        password_input.send_keys('kDp7%8^03Ib')
        
        signup_click = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH,'//*[@id="adlogin"]/div[3]/div/button'))
        )
        signup_click.click()
        
        custom_insights_click = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH,'/html/body/div/aside/section/ul/li[3]/a/span[1]'))
        )
        custom_insights_click.click()
        
        print("CMI automation completed successfully!")
        
        # Keep browser open for now - you can modify this behavior
        time.sleep(10)
        # driver.quit()  # Uncomment to close browser automatically
        
    except Exception as e:
        print(f"CMI automation error: {e}")


@app.route('/custom_wmr_cta', methods=['GET', 'POST'])
def custom_wmr_cta():
    if request.method == 'POST':
        try:
            # Get form data (only extract_count now since user info fields are removed)
            extract_count = int(request.form.get('extract_count', 200))

            # Validate required fields (only extract_count and file now)
            if not extract_count or extract_count < 1:
                flash('❌ Please enter a valid number of keywords to extract!')
                return redirect(request.url)

            file = request.files.get('file')
            if not file or file.filename == '':
                flash('❌ Keywords file is required!')
                return redirect(request.url)

            if not allowed_file(file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed!')
                return redirect(request.url)

            # Use secure_filename to avoid path issues
            filename = secure_filename(file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(input_path)

            # Process the file for WMR CTA generation
            result = process_wmr_cta_file(input_path, extract_count)
            
            if result['success']:
                flash(f'✅ Successfully processed {result["extracted_count"]} keywords!')
                flash(f'📁 File saved: {result["filename"]}')
                flash('🤖 WMR automation started!')
                
                # Start WMR automation in background
                threading.Thread(target=run_wmr_automation).start()
            else:
                flash(f'❌ Error: {result["error"]}')

            return render_template('custom_wmr_cta.html')

        except ValueError as ve:
            flash('❌ Please enter a valid number for keywords to extract!')
            return redirect(request.url)
        except Exception as e:
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_wmr_cta.html')


def process_wmr_cta_file(file_path, extract_count):
    """Process Custom Weekly ROB file for WMR CTA generation - Fixed UTF-8 support"""
    try:
        # ⭐ FIXED: Read CSV with proper UTF-8 encoding handling
        if file_path.endswith('.csv'):
            try:
                # Try UTF-8 first (most common)
                df_original = pd.read_csv(file_path, encoding='utf-8')
                print(f"✅ CSV file read successfully with UTF-8 encoding")
            except UnicodeDecodeError:
                try:
                    # Fallback to UTF-8 with BOM
                    df_original = pd.read_csv(file_path, encoding='utf-8-sig')
                    print(f"✅ CSV file read successfully with UTF-8-sig encoding")
                except UnicodeDecodeError:
                    try:
                        # Fallback to latin-1 (handles most encodings)
                        df_original = pd.read_csv(file_path, encoding='latin-1')
                        print(f"⚠️ CSV file read with latin-1 encoding (fallback)")
                    except Exception as e:
                        # Final fallback - let pandas auto-detect
                        df_original = pd.read_csv(file_path)
                        print(f"⚠️ CSV file read with auto-detected encoding")
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')
            print(f"✅ Excel file read successfully")

        print(f"📊 File loaded: {len(df_original)} rows, {len(df_original.columns)} columns")
        print(f"Columns: {list(df_original.columns)}")
        
        total_rows = len(df_original)
        
        if total_rows < extract_count:
            extract_count = total_rows
            print(f"⚠️ File only has {total_rows} rows, extracting all of them")

        # Step 1: Extract top N rows for WMR CTA generation
        extracted_rows = df_original.head(extract_count).copy()
        
        # Step 2: Get remaining rows (original minus extracted)
        remaining_rows = df_original.iloc[extract_count:].copy()
        
        # Create filename with current date for extracted data
        today = datetime.today()
        extracted_filename = f"custom_reports_wmr_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
        
        # Save extracted data to Desktop/RPA folder
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        # Save extracted keywords file
        extracted_output_path = os.path.join(rpa_folder, extracted_filename)
        extracted_rows.to_excel(extracted_output_path, index=False)
        print(f"✅ Extracted file saved: {extracted_filename}")
        
        # Step 3: Update the original Custom_weekly_ROB.xlsx file with remaining data
        custom_weekly_rob_path = os.path.join(rpa_folder, 'custom_weekly_rob.xlsx')  # Note: lowercase filename
        
        if os.path.exists(custom_weekly_rob_path):
            # Update the original file with remaining rows
            remaining_rows.to_excel(custom_weekly_rob_path, index=False)
            print(f"✅ Updated custom_weekly_rob.xlsx - Removed {extract_count} extracted keywords for WMR")
            print(f"✅ custom_weekly_rob.xlsx now contains {len(remaining_rows)} remaining keywords")
        else:
            print(f"⚠️ Warning: custom_weekly_rob.xlsx not found at {custom_weekly_rob_path}")
        
        return {
            'success': True,
            'extracted_count': extract_count,
            'remaining_count': len(remaining_rows),
            'total_count': total_rows,
            'filename': extracted_filename,
            'output_path': extracted_output_path,
            'updated_original': os.path.exists(custom_weekly_rob_path)
        }
        
    except Exception as e:
        print(f"❌ Error in process_wmr_cta_file: {e}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e)
        }
def run_wmr_automation():
    """Run WMR automation using Selenium with your provided code"""
    try:
        print("Starting WMR automation...")
        
        from webdriver_manager.chrome import ChromeDriverManager
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        import time

        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        #options.add_argument("--headless")  # Run headlessly for deployment
        options.add_argument("--start-maximized")

        # Initialize WebDriver
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.worldwidemarketreports.com/imanagereports')
                
        username_input = WebDriverWait(driver, 10).until(
             EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[1]/input'))
            )
        username_input.send_keys('Auto_Ops_Team')
                
        password_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[2]/input'))
            )
        password_input.send_keys('M9b@0j9Y28O')
                
        login_click = WebDriverWait(driver, 10).until(
          EC.element_to_be_clickable((By.XPATH,'//*[@id="adlogin"]/div[3]/div/button'))
            )
        login_click.click()
                
        custom_insights_click = WebDriverWait(driver, 10).until(
           EC.element_to_be_clickable((By.XPATH,'/html/body/div/aside/section/ul/li[3]/a/span[1]'))
            )
        custom_insights_click.click()
        
        print("WMR automation completed successfully!")
        
        # Keep browser open for now - you can modify this behavior
        time.sleep(10)
        # driver.quit()  # Uncomment to close browser automatically
        
    except Exception as e:
        print(f"WMR automation error: {e}")


@app.route('/custom_content_generation_choice')
def custom_content_generation_choice():
    """Custom content generation choice page"""
    return render_template('custom_content_generation_choice.html')

# Add these imports at the top if not already present
import openai
from docx import Document
import re

# Configure OpenAI (add your API key)
OPENAI_API_KEY = "spzlP-DmgsXi7l1ivu8oKC8hmo4pDIBblicYkX_cHll6bEA"  # Replace with your actual API key

@app.route('/custom_ai_content', methods=['GET', 'POST'])
def custom_ai_content():
    if request.method == 'POST':
        try:
            # Handle file upload
            cta_file = request.files.get('cta_file')
            if not cta_file or cta_file.filename == '':
                flash('❌ CTA excel file is required!')
                return redirect(request.url)

            if not allowed_file(cta_file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) are allowed!')
                return redirect(request.url)

            # Save uploaded file
            filename = secure_filename(cta_file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            cta_file.save(input_path)

            # Process the file for AI content generation
            result = process_ai_content_generation(input_path)
            
            if result['success']:
                flash(f'✅ Successfully generated {result["articles_created"]} AI articles!')
                flash(f'📁 Articles saved to Desktop/RPA folder')
            else:
                flash(f'❌ Error: {result["error"]}')

            # Clean up uploaded file
            try:
                os.remove(input_path)
            except:
                pass

            return render_template('custom_ai_content.html')

        except Exception as e:
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_ai_content.html')


def generate_article_with_openai(clean_title, promo_link, sample_link):
    """Generate article using OpenAI API"""
    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        
        prompt = f"""       
We are coherent market insights and going to publish this data on third part PR website, thus do not mention any competitors name (market intelligence firm, market research, consulting and data and analytics companies), and take ownership for market share, size, growth rate as this research done by our own research team however any real instances and stats you can use reference such as company or organisation or others as mentioned in detail explanation. 

IMPORTANT LINK FORMATTING RULES:
- Do NOT use brackets for links like [text](url) 
- Use ONLY this format: ➔ Link Text: URL
- Example: ➔ Get the Sample Copy of the Research Report: {sample_link}
- First CTA must use sample_link: ➔ Get the Sample Copy of the Research Report: {sample_link}
- Second CTA must use promo_link: ➔ Get Instant Access! Purchase Research Report and Receive up to 70% Discount: {promo_link}
- The second CTA with promo_link must always be the last link in the article

Write the blog covering the below pointers: Start content with market size and overview paragraph (do not add heading and do not use bullet for this section), ➤ Strategic Actionable Insights for the Market, ➔ add first CTA link here, ➤Market Taxonomy and Regional Coverage of Report, ➤Leading Companies of the Market, ➔ add Second CTA link here, ➤ Key Growth Drivers Fueling Market Expansion, ➤Key Reasons for Buying the (insert market name here) Report ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ➤ Frequently Asked Questions, All these pointers should act as a heading to respective paragraphs, do not miss any bullet for above given. Consider the following points while generating content: Flow of information, all given secondary keywords must be covered, and there should be a heading for each paragraph or bullet pointers. I need this content to be very presentable format thus make sure there should be space after and before of each heading or CTA links (Strategic Actionable Insights for the Market, first CTA link, Market Taxonomy and Regional coverage of Report, Leading Companies of the Market, Second CTA link, Key Growth Drivers Fueling Market Expansion, Key Reasons for Buying the (insert market name here) Report, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, and Frequently Asked Questions), this will increase the readability. Cover content in in bullet pointers whenever possible each paragraph should be short. Our readers are already experts in the field, so always try to generate content that provides unique insights and value addition for experts. Thus, while generating each piece of content, it should be data-backed with actual instances and stats from recent years 2025 and 2024, covering maximum stats that increase the authenticity and are enough to support or make decisions based upon reading this blog. Do not add generic content that is already known to readers, such as definitions, advantages, disadvantages, or other generic content. Flow of information should be as below. Start the blog with writing 5 to 7 sentence ling paragraph start content with talking about Industry in 2 to 3 sentences and should be matches with heading of the blog. followed by 2 to 4 sentence on market size and CAGR formatted as The Global (Insert Market Name) Market size is estimated to be valued at USD (Identify and Insert Market Value for 2025) in 2025 (market size) and is expected to reach USD (Identify and Insert Market Value for 2032) by 2032 (market forecast), exhibiting a compound annual growth rate (CAGR) of (Identify and Insert CAGR)% from 2025 to 2032. Do not add heading and do not use bullet for this section. (Make sure all values such as market size, CAGR, key players should be carefully identified with research approach). Then Strategic Actionable Insights for the Market: In Strategic Actionable Insights for the Market, cover 3 to 4 Strategic Actionable Insights for the Market in 10 to 12 sentences each insights comprises of 2 t 3 sentences covering supporting data for each insights. Insights could be essential quantitative indicators that lead to market size. For example, supply-side indicators (ex. Production Capacity, Pricing, and Exports. but not necessary to add if not matches to given market), demand-side indicators (ex. Pricing, Imports, Various Use Cases across industries. but not necessary to add if not matches to given market), micro-indicators, nano-size indicators—focus more on the quantitative aspects. Each Strategic Actionable Insights for the Market must have two sentence stats or actual instance examples from the recent year to support each point given in Strategic Actionable Insights for the Market, so that each given point look complete and meaningful. Then First CTA link. Next part is Market Taxonomy and Regional coverage of Report where enlist the all subsegment under each segment categories and fragment region into given format. Identify Comprehensive Market Taxonomy of the Report: • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. Regional and Country Analysis: • North America: U.S. and Canada • Latin America: Brazil, Argentina, Mexico, and Rest of Latin America • Europe: Germany, U.K., Spain, France, Italy, Benelux, Denmark, Norway, Sweden, Russia, and Rest of Europe • Asia Pacific: China, Taiwan, India, Japan, South Korea, Indonesia, Malaysia, Philippines, Singapore, Australia, and Rest of Asia Pacific • Middle East & Africa: Bahrain, Kuwait, Oman, Qatar, Saudi Arabia, United Arab Emirates, Israel, South Africa, North Africa, Central Africa, and Rest of MEA. Then Leading Companies of the Market: Identify and Enlist 12 to 20 highly relevant Leading Companies of the Market for the given market. Furthermore, add 2-3 statements on competitive strategies adopted by a few Leading Companies of the Market, mentioning actual strategies and entities involved along with the actual outcome. Then Add Second CTA link. Key Reasons for Buying the (insert market name here) Report, and its exact content as shared in data. Key Growth Drivers Fueling Market Expansion: Growth factor heading and short paragraph (3-4 Key Growth Drivers Fueling Market Expansion covered under 10 to 12 sentences) with supporting stats or examples from the recent year in the content, each factors should be covered in two to three sentences thus entire Key Growth Drivers Fueling Market Expansion content will be covered in 10 to 12 sentences long. No sub bullet is needed in Growth Factor. Then Emerging Trends and Market Shift: Market Trend heading and short paragraphs with supporting stats or examples from the recent year in the content (No bullet needed for as opportunity are written in paragraph format). Then High-Impact Market Opportunities by Segment and Region: Provide 3 to 4 High-Impact Market Opportunities by Segment and Region, 2-3 opportunities based upon segment and one opportunity based upon region in a paragraph format. Each opportunity described in two to three sentences and supported by actual data. For each opportunity, identify a single segment and provide a short description of the opportunity within that segment. Similarly, highlight one region, along with a brief description of the opportunity in each regional market (No bullet needed for as opportunity are written in paragraph format). Make sure to not to mention our organization name or relevant terms anywhere in the output such as coherent market insights or our analyst team or our research team Given Market Name and Data: {clean_title}. First CTA-  ➔{sample_link}and,
IMPORTANT LINK FORMATTING: 
- Do NOT use brackets for links like [text](url)
- Use this exact format: ➔ Get the Sample Copy of the Research Report: {sample_link}
- Use this exact format: ➔ Get Instant Access! Purchase Research Report and Receive up to 70% Discount: {promo_link}
-


 ➤Key Reasons for Buying the (insert market name here) Report: • Comprehensive analysis of the changing competitive landscape • Assists in decision-making processes for the businesses along with detailed strategic planning methodologies • The report offers forecast data and an assessment of the (insert market name here) • Helps in understanding the key product segments and their estimated growth rate • In-depth analysis of market drivers, restraints, trends, and opportunities • Comprehensive regional analysis of the (insert market name here) • Extensive profiling of the key stakeholders of the business sphere • Detailed analysis of the factors influencing the growth of the (insert market name here). From an SEO perspective, we need to cover all given keywords from the list below. However, they should appear naturally so that the content flow looks natural for the reader. Keyword List: market share, market size, market research, market insights, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, market challenges, market growth, market forecast, market companies, market players, market analysis, market drivers, market restraints, market scope, market dynamics, market segments, market report, market growth strategies, market revenue, industry size, industry share, industry trends, and business growth, furthermore - Market size and market report, market revenue, market share, trends keywords are mandatory to be added twice in content. In addition to the above requirement, in 5 places, add the actual market name along with the above keywords so that long-tail keywords will be generated. These long-tail keywords are market name + size, market name + report, market name + revenue, market name + share, market name + trends. Make sure all given keywords are naturally fit, do not try to infuse forcefully, flow of information should be natural and meaningful, furthermore make sure spelling and structure of sentences from generated output are grammatically correct. Furthermore, based on the market name, create a set of Frequently Asked Questions that are highly relevant and customized to the specific market. The sample Frequently Asked Questions below are for understanding purposes only. For the given market, questions can be completely replaced. However, please tailor the actual questions to the market name and the insights provided in the report: 1. Who are the dominant players in the (Market Name) market? 2. What will be the size of the (Market Name) market in the coming years? 3. Which end users industry has the largest growth opportunity? 4. How will market development trends evolve over the next five years? 5. What is the nature of the competitive landscape and challenges in the (Market Name) market? 6. What go-to-market strategies are commonly adopted in the (Market Name) market? Make sure to answer to all FAQs. In the case of country-level markets, please exclude the word 'Global' and High-Impact Market Opportunities by Segment and Region where other regions are mentioned. Make sure to add catchy bullet in generated output. I have shared the reference bullet with you. Make sure to add this bullet. For heading use these bullet- ), ➤Strategic Actionable Insights for the Market, ➔ add first CTA link here ➤ Market Taxonomy and Regional coverage of Report, ➤ Leading Companies of the Market, ➔ Inserted Second CTA link, ➤Key Reasons for Buying the (insert market name here) Report, ➤Key Growth Drivers Fueling Market Expansion, ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ➤ Frequently Asked Questions. Make sure do not miss any bullet including CTA bullet which is ➔. For subpointers under main headings use bullets which is in reference as provided- Strategic Actionable Insights for the Market ●, Market Taxonomy and Regional coverage of Report●, Leading Companies of the Market●. Make sure to use these bullets for given subpointers. Ensure proper bullet formatting so that each point is marked with a single bullet only, and avoid placing two bullets adjacent to each other.
"""
        
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )
        
        article_content = response.choices[0].message.content
        
        return article_content
        
    except Exception as e:
        print(f"OpenAI error: {e}")
        return "Error generating article content. Please try again later."
    
def save_article_as_doc(article_content, clean_title):
    """Save article as .doc file"""
    try:
        # Create Word document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(f"{clean_title} - Market Analysis", level=1)
        
        # Add content paragraphs
        paragraphs = article_content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Create filename
        today = datetime.today()
        safe_title = re.sub(r'[^\w\s-]', '', clean_title.lower())
        safe_title = re.sub(r'[-\s]+', '_', safe_title)
        filename = f"{safe_title}_cmi_{today.year}_{today.month:02d}_{today.day:02d}.doc"
        
        # Save to Desktop/RPA
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        file_path = os.path.join(rpa_folder, filename)
        doc.save(file_path)
        
        return True, filename
        
    except Exception as e:
        print(f"Error saving document: {e}")
        return False, str(e)

def process_ai_content_generation(file_path):
    """Process CTA excel file and generate AI articles"""
    try:
        print(f"\n=== PROCESSING AI CONTENT GENERATION ===")
        print(f"File: {file_path}")
        
        # Read the excel file
        df = pd.read_excel(file_path, engine='openpyxl')
        
        print(f"Excel columns: {list(df.columns)}")
        print(f"Found {len(df)} rows to process")
        
        # Verify required columns exist
        required_columns = ['KEYWORD', 'PROMOBUY', 'SAMPLECOPY']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return {
                'success': False,
                'error': f'Missing required columns: {missing_columns}'
            }
        
        articles_created = 0
        
        # Process each row
        for index, row in df.iterrows():
            try:
                original_title = str(row['KEYWORD'])
                promo_link = str(row['PROMOBUY'])
                sample_link = str(row['SAMPLECOPY'])
    
                print(f"\n[{index+1}/{len(df)}] Processing: {original_title}")
                
                # Clean the title - FIXED: Use different variable name
                #cleaned_title = clean_title(original_title)  # ✅ FIXED!
                print(f"Cleaned title: {original_title}")
                
                # Generate article using OpenAI
                print("Generating article with OpenAI...")
                article_content = generate_article_with_openai(original_title, promo_link, sample_link)
                
                # Save as .doc file
                success, filename = save_article_as_doc(article_content, original_title)
                
                if success:
                    print(f"✅ Article saved: {filename}")
                    articles_created += 1
                else:
                    print(f"❌ Failed to save article: {filename}")
                
                # Small delay to avoid API rate limits
                time.sleep(1)
                
            except Exception as e:
                print(f"❌ Error processing row {index+1}: {e}")
                continue
        
        return {
            'success': True,
            'articles_created': articles_created,
            'total_rows': len(df)
        }
        
    except Exception as e:
        print(f"Error in AI content generation: {e}")
        return {
            'success': False,
            'error': str(e)
        }


'''@app.route('/custom_template_content', methods=['GET', 'POST'])
def custom_template_content():
    if request.method == 'POST':
        # Handle template-based content generation
        # Similar to custom_ai_content but using predefined templates
        pass
    return render_template('custom_template_content.html')'''

@app.route('/custom_wmr_templates', methods=['GET', 'POST'])
def custom_wmr_templates():
    """WMR specific template selection and processing"""
    if request.method == 'POST':
        # Handle WMR template processing
        return handle_wmr_template_processing()
    return render_template('custom_wmr_templates.html')

@app.route('/custom_cmi_templates', methods=['GET', 'POST'])  
def custom_cmi_templates():
    """CMI specific template selection and processing"""
    if request.method == 'POST':
        # Handle CMI template processing
        return handle_cmi_template_processing()
    return render_template('custom_cmi_templates.html')


#NEW #NEW 
@app.route('/custom_cmi_template_processor', methods=['GET', 'POST'])
def custom_cmi_template_processor():
    if request.method == 'POST':
        try:
            # Get the uploaded file and selected template info
            template_file = request.files.get('template_file')
            selected_category = request.form.get('selected_category', '')
            selected_template = request.form.get('selected_template', '')
            
            if not template_file or template_file.filename == '':
                flash('❌ Please upload an Excel file!')
                return redirect(request.url)

            if not allowed_file(template_file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) are allowed!')
                return redirect(request.url)

            # Save uploaded file
            filename = secure_filename(template_file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            template_file.save(input_path)

            # Process the templates
            result = process_cmi_templates_with_excel(input_path, selected_category, selected_template)
            
            if result['success']:
                flash(f'✅ Successfully generated {result["files_created"]} template files!')
                flash(f'📁 Files saved in respective domain folders')
            else:
                flash(f'❌ Error: {result["error"]}')

            # Clean up uploaded file
            try:
                os.remove(input_path)
            except:
                pass

            return render_template('custom_cmi_templates.html')

        except Exception as e:
            flash(f'❌ Error processing templates: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_cmi_templates.html')

def process_cmi_templates_with_excel(excel_path, selected_category='', selected_template=''):
    """Process Excel file and generate templates for each row"""
    try:
        print(f"\n=== PROCESSING CMI TEMPLATES ===")
        print(f"Excel file: {excel_path}")
        
        # Read the Excel file
        df = pd.read_excel(excel_path, engine='openpyxl')
        
        print(f"Excel columns: {list(df.columns)}")
        print(f"Found {len(df)} rows to process")
        
        # Verify required columns exist
        required_columns = ['KEYWORD', 'PROMOBUY', 'SAMPLECOPY', 'Category']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return {
                'success': False,
                'error': f'Missing required columns: {missing_columns}'
            }
        
        files_created = 0
        base_rpa_path = os.path.join(os.path.expanduser("~"), "Desktop", "RPA")
        
        # Process each row
        for index, row in df.iterrows():
            try:
                keyword = str(row['KEYWORD']).strip()
                category = str(row['Category']).strip()
                
                
                print(f"\n[{index+1}/{len(df)}] Processing: {keyword} (Category: {category})")
                
                # Validate category
                if category not in ['HC', 'CMFE', 'ICT']:
                    print(f"❌ Invalid category: {category}. Skipping.")
                    continue
                
                # Get random template from domain folder
                template_path = get_random_template_from_domain(base_rpa_path, category)
                if not template_path:
                    print(f"❌ No templates found in {category} folder")
                    continue
                
                print(f"Selected template: {os.path.basename(template_path)}")
                
                # Process the template
                success, output_path = process_single_template(template_path, row, category, base_rpa_path)
                
                if success:
                    print(f"✅ Template generated: {os.path.basename(output_path)}")
                    files_created += 1
                else:
                    print(f"❌ Failed to generate template for: {keyword}")
                
            except Exception as e:
                print(f"❌ Error processing row {index+1}: {e}")
                continue
        
        return {
            'success': True,
            'files_created': files_created,
            'total_rows': len(df)
        }
        
    except Exception as e:
        print(f"Error in CMI template processing: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def get_random_template_from_domain(base_path, domain):
    """Get a random .doc template from the specified domain folder"""
    try:
        domain_path = os.path.join(base_path, domain)
        
        if not os.path.exists(domain_path):
            print(f"❌ Domain folder not found: {domain_path}")
            return None
        
        # Find all .doc files in the domain folder
        doc_files = [f for f in os.listdir(domain_path) if f.lower().endswith('.doc') or f.lower().endswith('.docx')]
        
        if not doc_files:
            print(f"❌ No .doc files found in {domain_path}")
            return None
        
        # Select random template
        import random
        selected_template = random.choice(doc_files)
        template_path = os.path.join(domain_path, selected_template)
        
        print(f"Random template selected: {selected_template}")
        return template_path
        
    except Exception as e:
        print(f"Error getting random template: {e}")
        return None
    
def replace_placeholders_in_docx(excel_row, doc):
    # Define the placeholders and their corresponding Excel columns
    placeholder_mapping = {
        'KEYWORD': 'KEYWORD', 
        'PROMOBUY': 'PROMOBUY',
        'SAMPLECOPY': 'SAMPLECOPY'
    }

    # Iterate over paragraphs in the Word document
    for para in doc.paragraphs:
        for placeholder, column_name in placeholder_mapping.items():
            if placeholder in para.text:
                # Get the value from the Excel row for the placeholder
                value = str(excel_row[column_name]).strip()
                # Check if it's the 'KEYWORD' placeholder and remove the word "Market"
                if placeholder == 'KEYWORD' and "Market" in value:
                    value = value.replace("Market", "").strip()  # Remove "Market" from KEYWORD
                
                # Replace the placeholder with the actual value
                para.text = para.text.replace(placeholder, value)

    # Iterate over tables in the Word document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, column_name in placeholder_mapping.items():
                    if placeholder in cell.text:
                        # Get the value from the Excel row for the placeholder
                        value = str(excel_row[column_name]).strip()
                        # Check if it's the 'KEYWORD' placeholder and remove the word "Market"
                        if placeholder == 'KEYWORD' and "Market" in value:
                            value = value.replace("Market", "").strip()  # Remove "Market" from KEYWORD
                        
                        # Replace the placeholder with the actual value
                        cell.text = cell.text.replace(placeholder, value)


def process_single_template(template_path, row_data, category, base_rpa_path):
    try:
        today = datetime.today()
        
        # Generate output filename using the 'KEYWORD' and category
        keyword = str(row_data['KEYWORD']).strip()
        # Remove "Market" from the KEYWORD value
        keyword = keyword.replace("Market", "").strip()  # Remove "Market" from KEYWORD
        safe_keyword = "".join(c for c in keyword if c.isalnum() or c in (' ', '_')).strip()
        
        filename = f"{category}_{safe_keyword}_{today.year}_{today.month:02d}_{today.day:02d}.doc"
        
        # Save generated files in a separate "Generated" subfolder
        domain_folder = os.path.join(base_rpa_path, category, "Generated")
        if not os.path.exists(domain_folder):
            os.makedirs(domain_folder)
        
        output_path = os.path.join(domain_folder, filename)
        
        # Process the template
        success = process_template_with_formatting(template_path, row_data, output_path, keyword)
        
        if success:
            return True, output_path
        else:
            return False, None
            
    except Exception as e:
        print(f"Error processing single template: {e}")
        return False, None

def process_template_with_formatting(template_path, row_data, output_path, keyword):
    try:
        # Open Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Open the template document
        doc = word.Documents.Open(template_path)
        
        # Prepare replacement data
        today = datetime.today()
        replacements = get_replacement_data(row_data, today)
        
        print(f"Applying {len(replacements)} replacements to template")
        
        # Replace placeholders while preserving formatting
        for placeholder, replacement in replacements.items():
            if replacement:  # Only replace if we have a value
                # Check if it's 'KEYWORD' and remove "Market"
                if placeholder == 'KEYWORD':
                    replacement = replacement.replace("Market", "").strip()  # Remove "Market"
                
                replace_text_preserve_formatting(doc, placeholder, str(replacement))
        
        # STEP 2: Replace OpenAI placeholders (ADD THIS SECTION)
        category = str(row_data.get('Category', '')).strip()
        replace_openai_placeholders(doc, keyword, category)

        # Save the document with preserved formatting
        doc.SaveAs2(output_path, FileFormat=0)  # 0 = Word 97-2003 format
        doc.Close()
        word.Quit()
        
        return True
        
    except Exception as e:
        print(f"Error processing template with formatting: {e}")
        try:
            word.Quit()
        except:
            pass
        return False


def get_replacement_data(row_data, today):
    """Get replacement data from Excel - simplified for 3 placeholders only"""
    
    # DEBUG: Print what we're working with
    print(f"DEBUG - Excel row keys: {list(row_data.keys())}")
    print(f"DEBUG - Raw values:")
    print(f"  KEYWORD: '{row_data.get('KEYWORD')}'")
    print(f"  PROMOBUY: '{row_data.get('PROMOBUY')}'") 
    print(f"  SAMPLECOPY: '{row_data.get('SAMPLECOPY')}'")
    
    replacements = {
        'KEYWORD': str(row_data.get('KEYWORD', '')),
        'PROMOBUY': str(row_data.get('PROMOBUY', '')), 
        'SAMPLECOPY': str(row_data.get('SAMPLECOPY', ''))
    }
    
    print(f"DEBUG - Final replacements: {replacements}")
    return replacements

def read_template_content(template_path):
    """Read content from template file"""
    try:
        if template_path.lower().endswith('.docx'):
            # Read .docx file
            doc = Document(template_path)
            content = []
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            return '\n'.join(content)
        
        elif template_path.lower().endswith('.doc'):
            # Read .doc file using win32com
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(template_path)
            content = doc.Content.Text
            doc.Close()
            word.Quit()
            return content
        
        else:
            # Try reading as text file
            with open(template_path, 'r', encoding='utf-8') as f:
                return f.read()
                
    except Exception as e:
        print(f"Error reading template {template_path}: {e}")
        return None

def replace_text_preserve_formatting(doc, find_text, replace_text):
    try:
        print(f"DEBUG - Searching for: '{find_text}' to replace with: '{replace_text}'")
        
        replacements_made = 0
        
        # Use Selection object for better formatting preservation
        selection = doc.Application.Selection
        
        # Start from beginning of document
        selection.HomeKey(Unit=6)  # wdStory - go to start of document
        
        # Loop to find and replace all instances
        while True:
            # Use Find to locate the text
            found = selection.Find.Execute(
                FindText=find_text,
                Forward=True,
                Wrap=0,  # wdFindStop - don't wrap around
                MatchCase=False,
                MatchWholeWord=False
            )
            
            if found:
                print(f"DEBUG - Found '{find_text}' at position {selection.Start}")
                # Replace the selected text (preserves formatting of surrounding text)
                selection.TypeText(replace_text)
                replacements_made += 1
                print(f"DEBUG - Replaced occurrence #{replacements_made}")
            else:
                break  # No more instances found
        
        print(f"DEBUG - Total replacements made: {replacements_made}")
        
        # Verify replacement worked
        final_content = doc.Content.Text
        if find_text in final_content:
            print(f"DEBUG - ❌ '{find_text}' still exists after replacement!")
        else:
            print(f"DEBUG - ✅ '{find_text}' successfully removed!")
            
        if replace_text in final_content:
            print(f"DEBUG - ✅ New text '{replace_text}' found in document!")
        else:
            print(f"DEBUG - ❌ New text '{replace_text}' NOT found!")
        
        print(f"Formatting-preserving replacement completed for '{find_text}'")

    except Exception as e:
        print(f"Error in formatting-preserving replacement '{find_text}': {e}")

def clean_content_for_xml(content):
    """Clean content to make it XML compatible"""
    import re
    
    if not content:
        return ""
    
    # Convert to string if not already
    content = str(content)
    
    # Remove NULL bytes and control characters (except \t, \n, \r)
    content = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', content)
    
    # Remove or replace other problematic characters
    # Remove excessive whitespace
    content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)  # Multiple empty lines to double
    content = re.sub(r'[ \t]+', ' ', content)  # Multiple spaces to single space
    
    # Remove any remaining XML-incompatible characters
    # Keep only printable ASCII + extended characters, newlines, tabs
    content = ''.join(char for char in content if ord(char) >= 32 or char in '\n\r\t')
    
    # Ensure proper encoding
    try:
        # Encode and decode to handle any encoding issues
        content = content.encode('utf-8', errors='ignore').decode('utf-8')
    except:
        # Fallback to ASCII if UTF-8 fails
        content = content.encode('ascii', errors='ignore').decode('ascii')
    
    # Final cleanup - remove any empty lines at start/end
    content = content.strip()
    
    return content

#Open AI api call for template content
def find_openai_placeholders(doc):
    """Find which OpenAI placeholders exist in the document"""
    try:
        # Get document content
        doc_content = doc.Content.Text
        
        # Define OpenAI placeholders to look for
        openai_placeholders = ['KEYPLAYERS', 'SEGMENTS', 'APPLICATIONS', 'TAXONOMY']
        
        # Find which ones exist in the document
        found_placeholders = []
        for placeholder in openai_placeholders:
            if placeholder in doc_content:
                found_placeholders.append(placeholder)
        
        print(f"DEBUG - Found OpenAI placeholders: {found_placeholders}")
        return found_placeholders
        
    except Exception as e:
        print(f"Error finding OpenAI placeholders: {e}")
        return []

def generate_openai_content(keyword, category, placeholders):
    """Generate content for found placeholders using OpenAI"""
    try:
        import openai
        
        # Remove "Market" from keyword for better context
        clean_keyword = keyword.replace("Market", "").strip()
        
        # Create dynamic prompt based on found placeholders
        prompt = f"""
Generate professional market research content for {clean_keyword} in the {category} domain.

Please provide content for these specific sections:
"""
        
        # Add sections based on found placeholders
        if 'KEYPLAYERS' in placeholders:
            prompt += "\n1. KEYPLAYERS: List 10-12 key market players/companies (comma-separated)"
        if 'SEGMENTS' in placeholders:
            prompt += "\n2. SEGMENTS: List market segments by type/product (comma-separated)" 
        if 'APPLICATIONS' in placeholders:
            prompt += "\n3. APPLICATIONS: List market applications/use cases (comma-separated)"
        if 'TAXONOMY' in placeholders:
            prompt += "\n4. TAXONOMY: Provide market segmentation structure (by type, application, region)"
        
        prompt += f"""

Market Context: {clean_keyword}
Industry Domain: {category}

Format your response exactly like this:
KEYPLAYERS: [content here]
SEGMENTS: [content here]  
APPLICATIONS: [content here]
TAXONOMY: [content here]

Only include sections that were requested above. Make content professional and industry-appropriate.
"""
        
        print(f"DEBUG - Making OpenAI call for {len(placeholders)} placeholders")
        
        # Make OpenAI API call
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional market research content generator."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )
        
        ai_content = response.choices[0].message.content
        print(f"DEBUG - OpenAI response received: {len(ai_content)} characters")
        
        return ai_content
        
    except Exception as e:
        print(f"ERROR - OpenAI API call failed: {e}")
        return None

def parse_openai_response(ai_content, placeholders):
    """Parse OpenAI response into individual placeholder content"""
    try:
        replacements = {}
        
        if not ai_content:
            return replacements
        
        # Split response into lines
        lines = ai_content.split('\n')
        
        # Parse each line for placeholder content
        for line in lines:
            if ':' in line:
                parts = line.split(':', 1)
                if len(parts) == 2:
                    placeholder = parts[0].strip()
                    content = parts[1].strip()
                    
                    # Only add if it's one of our requested placeholders
                    if placeholder in placeholders and content:
                        replacements[placeholder] = content
        
        print(f"DEBUG - Parsed {len(replacements)} placeholder replacements")
        for key, value in replacements.items():
            print(f"DEBUG - {key}: {value[:100]}...")
        
        return replacements
        
    except Exception as e:
        print(f"ERROR - Failed to parse OpenAI response: {e}")
        return {}

def replace_openai_placeholders(doc, keyword, category):
    """Main function to handle OpenAI placeholder replacement"""
    try:
        print(f"DEBUG - Starting OpenAI placeholder replacement for: {keyword}")
        
        # Step 1: Find which OpenAI placeholders exist in document
        found_placeholders = find_openai_placeholders(doc)
        
        if not found_placeholders:
            print("DEBUG - No OpenAI placeholders found in document")
            return
        
        # Step 2: Generate content using OpenAI
        ai_content = generate_openai_content(keyword, category, found_placeholders)
        
        if not ai_content:
            print("ERROR - Failed to generate OpenAI content")
            return
        
        # Step 3: Parse OpenAI response
        replacements = parse_openai_response(ai_content, found_placeholders)
        
        if not replacements:
            print("ERROR - Failed to parse OpenAI response into usable content")
            return
        
        # Step 4: Replace each placeholder in document
        for placeholder, content in replacements.items():
            print(f"DEBUG - Replacing {placeholder} with OpenAI content...")
            replace_text_preserve_formatting(doc, placeholder, content)
        
        print(f"DEBUG - OpenAI placeholder replacement completed")
        
    except Exception as e:
        print(f"ERROR - OpenAI placeholder replacement failed: {e}")


#Template code for WMR
@app.route('/custom_wmr_template_processor', methods=['GET', 'POST'])
def custom_wmr_template_processor():
    if request.method == 'POST':
        try:
            # Get the uploaded file
            template_file = request.files.get('template_file')
            
            if not template_file or template_file.filename == '':
                flash('❌ Please upload an Excel file!')
                return redirect(request.url)

            if not allowed_file(template_file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) are allowed!')
                return redirect(request.url)

            # Save uploaded file
            filename = secure_filename(template_file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            template_file.save(input_path)

            # Process the WMR templates
            result = process_wmr_templates_with_excel(input_path)
            
            if result['success']:
                flash(f'✅ Successfully generated {result["files_created"]} WMR template files!')
                flash(f'📁 Files saved in WMR/Generated folder')
            else:
                flash(f'❌ Error: {result["error"]}')

            # Clean up uploaded file
            try:
                os.remove(input_path)
            except:
                pass

            return render_template('custom_wmr_templates.html')  # You'll need to create this template

        except Exception as e:
            flash(f'❌ Error processing WMR templates: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_wmr_templates.html')


def process_wmr_templates_with_excel(excel_path):
    """Process Excel file and generate WMR templates for each row"""
    try:
        print(f"\n=== PROCESSING WMR TEMPLATES ===")
        print(f"Excel file: {excel_path}")
        
        # Read the Excel file
        df = pd.read_excel(excel_path, engine='openpyxl')
        
        print(f"Excel columns: {list(df.columns)}")
        print(f"Found {len(df)} rows to process")
        
        # Verify required columns exist (no Category needed for WMR)
        required_columns = ['KEYWORD', 'PROMOBUY', 'SAMPLECOPY']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return {
                'success': False,
                'error': f'Missing required columns: {missing_columns}'
            }
        
        files_created = 0
        base_rpa_path = os.path.join(os.path.expanduser("~"), "Desktop", "RPA")
        
        # Process each row
        for index, row in df.iterrows():
            try:
                keyword = str(row['KEYWORD']).strip()
                
                print(f"\n[{index+1}/{len(df)}] Processing WMR: {keyword}")
                
                # Get random WMR template
                template_path = get_random_wmr_template(base_rpa_path)
                if not template_path:
                    print(f"❌ No WMR templates found")
                    continue
                
                print(f"Selected WMR template: {os.path.basename(template_path)}")
                
                # Process the template
                success, output_path = process_single_wmr_template(template_path, row, base_rpa_path)
                
                if success:
                    print(f"✅ WMR Template generated: {os.path.basename(output_path)}")
                    files_created += 1
                else:
                    print(f"❌ Failed to generate WMR template for: {keyword}")
                
            except Exception as e:
                print(f"❌ Error processing WMR row {index+1}: {e}")
                continue
        
        return {
            'success': True,
            'files_created': files_created,
            'total_rows': len(df)
        }
        
    except Exception as e:
        print(f"Error in WMR template processing: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def get_random_wmr_template(base_path):
    """Get a random WMR template from the WMR folder"""
    try:
        wmr_path = os.path.join(base_path, "WMR")
        
        if not os.path.exists(wmr_path):
            print(f"❌ WMR folder not found: {wmr_path}")
            return None
        
        # Find all .doc files in the WMR folder, excluding lock files
        doc_files = [f for f in os.listdir(wmr_path) 
                    if (f.lower().endswith('.doc') or f.lower().endswith('.docx'))
                    and not f.startswith('~$')]
        
        if not doc_files:
            print(f"❌ No .doc files found in {wmr_path}")
            return None
        
        # Select random template
        import random
        selected_template = random.choice(doc_files)
        template_path = os.path.join(wmr_path, selected_template)
        
        print(f"Random WMR template selected: {selected_template}")
        return template_path
        
    except Exception as e:
        print(f"Error getting random WMR template: {e}")
        return None

def process_single_wmr_template(template_path, row_data, base_rpa_path):
    """Process a single WMR template with row data"""
    try:
        today = datetime.today()
        
        # Generate output filename using KEYWORD
        keyword = str(row_data['KEYWORD']).strip()
        # Remove "Market" from the KEYWORD value
        keyword = keyword.replace("Market", "").strip()
        safe_keyword = "".join(c for c in keyword if c.isalnum() or c in (' ', '_')).strip()
        safe_keyword = re.sub(r'\s+', '_', safe_keyword)
        
        filename = f"WMR_{safe_keyword}_{today.year}_{today.month:02d}_{today.day:02d}.doc"
        
        # Save generated files in WMR/Generated subfolder
        generated_folder = os.path.join(base_rpa_path, "WMR", "Generated")
        if not os.path.exists(generated_folder):
            os.makedirs(generated_folder)
        
        output_path = os.path.join(generated_folder, filename)
        
        # Process the WMR template
        success = process_wmr_template_with_formatting(template_path, row_data, output_path, keyword)
        
        if success:
            return True, output_path
        else:
            return False, None
            
    except Exception as e:
        print(f"Error processing single WMR template: {e}")
        return False, None

def process_wmr_template_with_formatting(template_path, row_data, output_path, keyword):
    """Process WMR template with formatting preservation"""
    try:
        # Open Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Open the template document
        doc = word.Documents.Open(template_path)
        
        # Prepare replacement data
        today = datetime.today()
        replacements = get_wmr_replacement_data(row_data, today)
        
        print(f"Applying {len(replacements)} basic WMR replacements to template")
        
        # STEP 1: Replace basic placeholders while preserving formatting
        for placeholder, replacement in replacements.items():
            if replacement:  # Only replace if we have a value
                # Check if it's 'KEYWORD' and remove "Market"
                if placeholder == 'KEYWORD':
                    replacement = replacement.replace("Market", "").strip()
                
                replace_text_preserve_formatting(doc, placeholder, str(replacement))
        
        # STEP 2: Replace OpenAI placeholders for WMR
        replace_wmr_openai_placeholders(doc, keyword)
        
        # Save the document with preserved formatting
        doc.SaveAs2(output_path, FileFormat=0)  # 0 = Word 97-2003 format
        doc.Close()
        word.Quit()
        
        return True
        
    except Exception as e:
        print(f"Error processing WMR template with formatting: {e}")
        try:
            word.Quit()
        except:
            pass
        return False

def get_wmr_replacement_data(row_data, today):
    """Get replacement data from Excel for WMR - 3 placeholders only"""
    replacements = {
        'KEYWORD': str(row_data.get('KEYWORD', '')),
        'PROMOBUY': str(row_data.get('PROMOBUY', '')), 
        'SAMPLECOPY': str(row_data.get('SAMPLECOPY', ''))
    }
    
    print(f"DEBUG - WMR Excel row keys: {list(row_data.keys())}")
    print(f"DEBUG - WMR Final replacements: {replacements}")
    
    return replacements

def replace_wmr_openai_placeholders(doc, keyword):
    """Handle OpenAI placeholder replacement for WMR"""
    try:
        print(f"DEBUG - Starting WMR OpenAI placeholder replacement for: {keyword}")
        
        # Step 1: Find which OpenAI placeholders exist in document
        found_placeholders = find_openai_placeholders(doc)  # Reuse existing function
        
        if not found_placeholders:
            print("DEBUG - No OpenAI placeholders found in WMR document")
            return
        
        # Step 2: Generate content using OpenAI (WMR context)
        ai_content = generate_wmr_openai_content(keyword, found_placeholders)
        
        if not ai_content:
            print("ERROR - Failed to generate WMR OpenAI content")
            return
        
        # Step 3: Parse OpenAI response
        replacements = parse_openai_response(ai_content, found_placeholders)  # Reuse existing function
        
        if not replacements:
            print("ERROR - Failed to parse WMR OpenAI response")
            return
        
        # Step 4: Replace each placeholder in document
        for placeholder, content in replacements.items():
            print(f"DEBUG - Replacing WMR {placeholder} with OpenAI content...")
            replace_text_preserve_formatting(doc, placeholder, content)
        
        print(f"DEBUG - WMR OpenAI placeholder replacement completed")
        
    except Exception as e:
        print(f"ERROR - WMR OpenAI placeholder replacement failed: {e}")

def generate_wmr_openai_content(keyword, placeholders):
    """Generate content for found placeholders using OpenAI for WMR"""
    try:
        import openai
        
        # Remove "Market" from keyword for better context
        clean_keyword = keyword.replace("Market", "").strip()
        
        # Create dynamic prompt based on found placeholders (WMR domain)
        prompt = f"""
Generate professional market research content for {clean_keyword} for World Market Reports (WMR).

Please provide content for these specific sections:
"""
        
        # Add sections based on found placeholders
        if 'KEYPLAYERS' in placeholders:
            prompt += "\n1. KEYPLAYERS: List 8-12 key market players/companies (comma-separated)"
        if 'SEGMENTS' in placeholders:
            prompt += "\n2. SEGMENTS: List market segments by type/product (comma-separated)" 
        if 'APPLICATIONS' in placeholders:
            prompt += "\n3. APPLICATIONS: List market applications/use cases (comma-separated)"
        if 'TAXONOMY' in placeholders:
            prompt += "\n4. TAXONOMY: Provide market segmentation structure (by type, application, region)"
        
        prompt += f"""

Market Context: {clean_keyword}
Report Type: World Market Reports (WMR)

Format your response exactly like this:
KEYPLAYERS: [content here]
SEGMENTS: [content here]  
APPLICATIONS: [content here]
TAXONOMY: [content here]

Only include sections that were requested above. Make content professional and industry-appropriate.
"""
        
        print(f"DEBUG - Making WMR OpenAI call for {len(placeholders)} placeholders")
        
        # Make OpenAI API call
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional market research content generator for World Market Reports."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )
        
        ai_content = response.choices[0].message.content
        print(f"DEBUG - WMR OpenAI response received: {len(ai_content)} characters")
        
        return ai_content
        
    except Exception as e:
        print(f"ERROR - WMR OpenAI API call failed: {e}")
        return None





# Step 4 Custom OpenPR Processing
@app.route('/custom_document_processing', methods=['GET', 'POST'])
def custom_document_processing():
    if request.method == 'POST':
        try:
            flash('Custom document processing started!')
            return redirect(url_for('custom_document_processing'))
        except Exception as e:
            flash(f'Error: {str(e)}')
            return redirect(url_for('custom_document_processing'))
    
    return render_template('custom_document_processing.html', 
                          session_data={
                              'username': session.get('username', ''),
                              'email': session.get('email', ''),
                              'mobile': session.get('mobile', ''),
                              'open_pr_id': session.get('open_pr_id', ''),
                              'image_path1': session.get('image_path1', ''),
                              'image_path2': session.get('image_path2', ''),
                              'image_path3': session.get('image_path3', ''),
                              'image_path4': session.get('image_path4', ''),
                              'image_path5': session.get('image_path5', ''),
                          })
    
@app.route('/start_custom_wmr_publishing', methods=['POST'])
def start_custom_wmr_publishing():
    """Start custom WMR publishing process with form data"""
    try:
        # Get form data from JSON request
        form_data = request.get_json()
        
        if not form_data:
            return jsonify({
                'status': 'error',
                'message': 'No form data received'
            }), 400
        
        print(f"Received form data: {form_data}")
        
        # Validate required fields
        required_fields = ['author_name', 'author_email', 'company_name', 'phone_number']
        for field in required_fields:
            if not form_data.get(field):
                return jsonify({
                    'status': 'error',
                    'message': f'Missing required field: {field}'
                }), 400
        
        # Validate image paths
        if not form_data.get('image_paths') or len(form_data['image_paths']) == 0:
            return jsonify({
                'status': 'error',
                'message': 'At least one image path is required'
            }), 400
        
        # Run the selenium automation in a background thread with form data
        def run_wmr_publishing():
            result = selenium_publishing_custom_wmr(form_data)
            print(f"WMR Publishing completed: {result}")
        
        # Start in background thread so user gets immediate response
        threading.Thread(target=run_wmr_publishing).start()
        
        return jsonify({
            'status': 'success',
            'message': 'WMR publishing started! Check console for progress updates.'
        })
        
    except Exception as e:
        print(f"Error in start_custom_wmr_publishing: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Error starting WMR publishing: {str(e)}'
        }), 500

def generate_creative_title_with_openai(keyword, title_prompt):
    """Generate creative, professional title using OpenAI API"""
    try:
        print(f"🤖 Generating creative title for '{keyword}' with prompt '{title_prompt}'")
        
        # Create the OpenAI prompt for title generation
        openai_prompt =f"""
        Create a professional, engaging, and creative title for a market research press release.
        
        Market/Keyword: {keyword} Market
        try your title should be  centred around to this nd  striclty  inlcude these  words in title : {title_prompt} i m saying strictly include these words in title
        
        Requirements:
        - Create a  title inspired by both the keyword and inspiration 
        - MUST include 3-4 key industry players/companies relevant to the {keyword} market
        - Focus on market analysis, growth, trends, and business opportunities
        - Keep the professional tone similar to the inspiration theme and if there KEYWORD is there in inspiration theme then remove it from title
        - Ensure the title does not exceed 110 characters (including spaces and punctuation)
        - The title should be catchy, attention-grabbing, and suitable for a press release
        - DO NOT mention "WMR", "CMI", or any company abbreviations in the title striclty take  look at title i ft has any of these words remove it 
        - DO NOT mention "Worldwide Market Reports", "Coherent Market Insights", or similar company names
        - Start the title with the market name, followed by key descriptors such as forecast, strategic growth, market trends, opportunities, etc.
        - Include key players naturally within the title context and place them at the end of the title.
        - Focus on qualitative descriptors, avoiding any numerical figures, percentages, or monetary values.
        - Make the title professional without semantic error.
        - DO NOT include any numbers, dollar amounts, percentages, or numerical figures (e.g., $85B, 25%, etc.).
"""


       
        
        # Make OpenAI API call
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are an expert copywriter specializing in compelling, professional market research titles for press releases. Create titles that grab attention while maintaining professional credibility."},
                {"role": "user", "content": openai_prompt}
            ],
            max_tokens=100,
            temperature=0.8  # Higher creativity for more varied titles
        )
        
        generated_title = response.choices[0].message.content.strip()
        
        # Remove quotes if OpenAI added them
        generated_title = generated_title.strip('"').strip("'")
        
        # Validate word count (40 words max)
        word_count = len(generated_title.split())
        if word_count > 40:
            print(f"⚠️ Generated title too long ({word_count} words), truncating...")
            words = generated_title.split()[:40]
            generated_title = ' '.join(words)
        
        print(f"✅ OpenAI generated title ({word_count} words): {generated_title}")
        return generated_title
        
    except Exception as e:
        print(f"❌ Error in OpenAI title generation: {e}")
        return None

import openai
from difflib import get_close_matches

def get_keyword_category_from_openai(keyword):
    """
    Classify keyword into appropriate category using OpenAI.
    """
    try:
        # Available OpenPR categories (based on the dropdown in screenshot)
        available_categories = [
            "Advertising, Media Consulting, Marketing Research",
            "Aerospace & Defense", 
            "Arts & Culture",
            "Associations & Organizations",
            "Business, Economy, Finance, Banking & Insurance",
            "Chemicals & Materials",
            "Consumer Goods & Retail",
            "Energy & Environment",
            "Fashion, Lifestyle, Trends",
            "Food & Beverage", 
            "Health & Medicine",
            "Industry, Real Estate & Construction",
            "IT, New Media & Software",
            "Leisure, Entertainment, Miscellaneous",
            "Logistics & Transport",
            "Media & Telecommunications", 
            "Politics, Law & Society",
            "Science & Education",
            "Sports",
            "Tourism, Cars, Traffic"
        ]
        
        # Hardcoded API key
        api_key = "sgIYB9nhvbb2G1maC8XOB0QqW4BNpzlP-DmgsXi7l1ivu8oKC8hmo4pDIBblicYkX_cHll6bEA"
        
        # Create the classification prompt
        prompt = f"""
You are a market research expert. Classify the following keyword/market into the most appropriate category from the list below.

Keyword: "{keyword}"

Available Categories:
{chr(10).join([f"- {cat}" for cat in available_categories])}

Instructions:
1. Analyze the keyword and determine which category it best fits into
2. Return ONLY the exact category name from the list above
3. Do not add any explanation or additional text
4. If uncertain, choose the closest matching category

Category:"""

        print(f"🤖 Classifying keyword: '{keyword}' using OpenAI...")

        # Make OpenAI API call
        client = openai.OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a precise market research categorization expert. Return only the exact category name from the provided list."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=50,
            temperature=0.3  # Lower temperature for more consistent results
        )

        predicted_category = response.choices[0].message.content.strip()

        # Validate that the response is in our available categories
        if predicted_category in available_categories:
            print(f"✅ Category classified: '{predicted_category}' for keyword '{keyword}'")
            return predicted_category
        else:
            print(f"⚠️ OpenAI returned invalid category: '{predicted_category}'")
            # Try to find the closest match using difflib
            closest_match = get_close_matches(predicted_category, available_categories, n=1, cutoff=0.7)
            if closest_match:
                print(f"✅ Found closest match: '{closest_match[0]}'")
                return closest_match[0]
            else:
                print("⚠️ No close match found. Returning None.")
                return None

    except Exception as e:
        print(f"❌ An error occurred: {str(e)}")
        return None


def load_title_prompts_from_excel():
    """Load title prompts from Excel file"""
    try:
        import pandas as pd
        
        title_file_path = r"C:\Users\vishwas\Desktop\RPA\WMR\TITLES\TITLE_SHEET.xlsx"
        
        print(f"📖 Loading title prompts from: {title_file_path}")
        
        # Read the Excel file
        df = pd.read_excel(title_file_path, engine='openpyxl')
        
        # Assuming titles are in the first column - adjust column name as needed
        # Common column names to check
        possible_columns = ['TITLES', 'Title', 'title', 'TITLE_PROMPTS', 'Prompts', 'PROMPTS']
        
        title_column = None
        for col in possible_columns:
            if col in df.columns:
                title_column = col
                break
        
        if title_column is None:
            # If no standard column found, use the first column
            title_column = df.columns[0]
            print(f"⚠️ Using first column '{title_column}' for titles")
        else:
            print(f"✅ Found title column: '{title_column}'")
        
        # Extract titles and remove empty/null values
        titles = df[title_column].dropna().astype(str).tolist()
        
        # Remove empty strings and strip whitespace
        titles = [title.strip() for title in titles if title.strip()]
        
        print(f"✅ Loaded {len(titles)} title prompts from Excel")
        print(f"📝 Sample titles: {titles[:3]}...")
        
        return titles
        
    except FileNotFoundError:
        print(f"❌ Title Excel file not found: {title_file_path}")
        return []
    except Exception as e:
        print(f"❌ Error loading title prompts from Excel: {e}")
        return []


def selenium_publishing_custom_wmr(form_data):
    """Fixed Selenium automation to publish all WMR articles using form data"""
    try:
        import random
        import os
        from datetime import datetime
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        from selenium.common.exceptions import StaleElementReferenceException, TimeoutException
        from webdriver_manager.chrome import ChromeDriverManager
        import time
        import re
        
        print(f"\n=== STARTING WMR CUSTOM PUBLISHING WITH FORM DATA ===")
        
        # Extract form data
        author_name = form_data.get('author_name', 'abhijit Tiwari')
        author_email = form_data.get('author_email', 'abhijit@coherentmarketinsights.com')
        company_name = form_data.get('company_name', 'Coherent Market Insights')
        phone_number = form_data.get('phone_number', '1234567890')
        article_code = form_data.get('article_code', 'D5A-2025-QDFH8C')
        processing_mode = form_data.get('processing_mode', 'auto')
        image_paths = form_data.get('image_paths', [])
        
        print(f"Using form data - Author: {author_name}, Email: {author_email}")
        print(f"Company: {company_name}, Phone: {phone_number}")
        print(f"Article Code: {article_code}, Mode: {processing_mode}")
        print(f"Available image paths: {len(image_paths)}")
        
        # Load title prompts from Excel file
        TITLE_PROMPTS = load_title_prompts_from_excel()
        
        if not TITLE_PROMPTS:
            print("❌ No title prompts loaded from Excel file. Cannot proceed without titles.")
            return {'success': False, 'error': 'Failed to load title prompts from Excel file'}
        
        print(f"✅ Using {len(TITLE_PROMPTS)} title prompts from Excel file")
        
        # Hardcoded constants
        AUTHOR_DESCRIPTIONS = [
            """Author of this marketing PR:
Ravina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.""",
            """Author of this marketing PR:
Money Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc.""",
            """Author of this marketing PR:
Alice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights."""
        ]
        
        website_category = "Business,Economy,Finance,Banking & Insurance"
        
        # Get all WMR articles from Generated folder
        base_rpa_custom_path = os.path.join(os.path.expanduser("~"), "Desktop", "RPA")
        generated_folder = os.path.join(base_rpa_custom_path, "WMR", "Generated")
        
        if not os.path.exists(generated_folder):
            print(f"❌ Generated folder not found: {generated_folder}")
            return {'success': False, 'error': 'Generated folder not found'}
        
        # Get all .doc and .docx files
        article_files = []
        for file in os.listdir(generated_folder):
            if file.lower().endswith(('.doc', '.docx')) and not file.startswith('~'):
                file_path = os.path.join(generated_folder, file)
                if os.path.isfile(file_path):
                    mod_time = os.path.getmtime(file_path)
                    article_files.append((file_path, mod_time, file))
        
        if not article_files:
            print(f"❌ No WMR articles found in: {generated_folder}")
            return {'success': False, 'error': 'No articles found in Generated folder'}
        
        # Sort by modification time (latest first)
        article_files.sort(key=lambda x: x[1], reverse=True)
        
        print(f"✅ Found {len(article_files)} WMR articles to publish")
        for i, (path, mod_time, filename) in enumerate(article_files, 1):
            print(f"  {i}. {filename} ({datetime.fromtimestamp(mod_time).strftime('%Y-%m-%d %H:%M')})")
        
        successful_publications = 0
        failed_publications = 0
        
        # Setup Chrome driver
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")  # Help avoid detection
        
        # Process each article
        for i, (article_path, mod_time, filename) in enumerate(article_files, 1):
            driver = None
            try:
                print(f"\n{'='*60}")
                print(f"Processing Article {i}/{len(article_files)}: {filename}")
                print(f"{'='*60}")
                
                # Randomly select image path for this article
                selected_image_path = random.choice(image_paths)
                print(f"🖼️ Randomly selected image path: {selected_image_path}")
                
                # Extract market name from filename
                market_name_from_file = extract_market_name_from_filename(filename)
                
                # Read article content using text_of_press_release
                print(f"📖 Reading article content from: {filename}")
                article_content = text_of_press_release(article_path)
                
                if not article_content or len(article_content.strip()) < 100:
                    print(f"❌ Article content too short or empty, skipping: {filename}")
                    failed_publications += 1
                    continue
                
                # Generate article title using Excel-loaded prompts
                random_prompt = random.choice(TITLE_PROMPTS)
                print(f"🎯 Selected title prompt from Excel: {random_prompt}")
                
                # Try to generate creative title with OpenAI
                creative_title = generate_creative_title_with_openai(market_name_from_file[1:], random_prompt)
                
                if creative_title:
                    article_title = creative_title
                    print(f"✨ Using OpenAI generated title: {article_title}")
                else:
                    # If OpenAI fails, use direct combination
                    article_title = f"{market_name_from_file} {random_prompt}"
                    print(f"📝 Using direct title combination: {article_title}")

                # Start fresh browser session for each article
                cService = Service(executable_path=chromedriver_path)
                driver = webdriver.Chrome(service=cService, options=options)
                
                # Set up longer default wait
                wait = WebDriverWait(driver, 20)
                
                print(f"🚀 Starting Selenium automation for: {market_name_from_file}")
                driver.get('https://www.openpr.com/')
                
                # Wait for page to fully load
                print("⏳ Waiting for page to load...")
                time.sleep(3)
                
                # Handle cookie consent
                try:
                    reject = wait.until(
                        EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
                    )
                    reject.click()
                    print("✅ Cookie consent handled")
                    time.sleep(2)  # Wait after cookie handling
                except TimeoutException:
                    print("⚠️ Cookie consent button not found or already handled")
                
                # Navigate to submit page
                submit = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
                )
                submit.click()
                print("✅ Navigated to submit page")
                
                # Wait for navigation to complete
                time.sleep(3)
                
                # Enter article code - FIXED: Re-find element after navigation
                print("🔍 Finding article code input field...")
                input_box = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="code"]'))
                )
                input_box.clear()
                input_box.send_keys(article_code)
                print(f"✅ Entered article code: {article_code}")
                
                # Submit code - FIXED: Wait and re-find element
                print("🔍 Finding submit button...")
                time.sleep(2)
                
                try:
                    submit2 = wait.until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
                    )
                    submit2.click()
                except TimeoutException:
                    submit2 = wait.until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
                    )
                    submit2.click()
                
                print("✅ Article code submitted")
                
                # 🔥 CRITICAL FIX: Wait longer for page to transition after code submission
                print("⏳ Waiting for page transition after code submission...")
                time.sleep(5)  # Increased wait time
                
                # FIXED: Re-find all elements after page transition - DON'T reuse old references
                print("🔍 Re-finding form elements after page transition...")
                
                # Fill form fields - FIXED: Find fresh elements
                name = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
                )
                name.clear()
                name.send_keys(author_name)
                print("✅ Name field filled")
                
                email = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
                )
                email.clear()
                email.send_keys(author_email)
                print("✅ Email field filled")
                
                number = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
                )
                number.clear()
                number.send_keys(phone_number)
                print("✅ Phone field filled")
                
                # Company name with popup handling
                ComName = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="archivnmfield"]'))
                )
                ComName.clear()
                ComName.send_keys(company_name)
                time.sleep(1)  # Wait for popup to appear
                
                # Handle company popup
                try:
                    s1 = wait.until(
                        EC.element_to_be_clickable((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
                    )
                    s1.click()
                    time.sleep(1)
                    print("✅ Company name selected")
                except TimeoutException:
                    print("⚠️ Company popup not found, continuing...")
                
                print("✅ Basic form fields filled with form data")
                
                # FIXED: Handle category selection with better error handling and waits
                print("🔍 Finding category dropdown...")
                 # Extra wait before category selection
                
                Category_element = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
                )
                
                # Wait for dropdown to be fully loaded
                time.sleep(2)
                
                # Try to get category using OpenAI
                try:
                    openai_category = get_keyword_category_from_openai(market_name_from_file)
                    if openai_category:
                        category_to_use = openai_category
                    else:
                        category_to_use = website_category
                except Exception as e:
                    print(f"⚠️ OpenAI category failed: {e}")
                    category_to_use = website_category
                
                # Select category with retry mechanism
                select_obj = Select(Category_element)
                category_selected = False
                max_retries = 3
                
                for retry in range(max_retries):
                    try:
                        select_obj.select_by_visible_text(category_to_use)
                        print(f"✅ Selected category: {category_to_use}")
                        category_selected = True
                        break
                    except Exception as e:
                        print(f"⚠️ Category selection attempt {retry+1} failed: {e}")
                        if retry < max_retries - 1:
                            time.sleep(2)
                            # Re-find the element for retry
                            Category_element = wait.until(
                                EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
                            )
                            select_obj = Select(Category_element)
                        else:
                            # Final fallback
                            select_obj.select_by_index(1)
                            print("✅ Used fallback category selection")
                            category_selected = True
                
                if not category_selected:
                    print("❌ Failed to select any category")
                    failed_publications += 1
                    continue
                
                # Fill title - FIXED: Fresh element
                title = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
                )
                title.clear()
                title.send_keys(article_title)
                print(f"✅ Entered title: {len(article_title)} characters")
                
                # Fill article content - FIXED: Fresh element
                text_field = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="inhalt"]'))
                )
                text_field.clear()
                text_field.send_keys(article_content)
                print(f"✅ Entered article content: {len(article_content)} characters")
                
                # Fill about section - FIXED: Fresh element
                about = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
                )
                about.clear()
                contact_info = f"""Contact Us:

Mr. Shah
Worldwide Market Reports,
Tel: U.S. +1-415-871-0703
U.K.: +44-203-289-4040
Australia: +61-2-4786-0457
India: +91-848-285-0837
Email: sales@worldwidemarketreports.com
Website: https://www.worldwidemarketreports.com/
"""
                about.send_keys(contact_info)
                
                # Fill address section with random author - FIXED: Fresh element
                address = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
                )
                address.clear()
                random_author = random.choice(AUTHOR_DESCRIPTIONS)
                address_content = f"{random_author}\n\nAbout Us:\n{company_name} Worldwide Market Reports is a global business intelligence firm offering market intelligence reports, databases, and competitive intelligence reports. We offer reports across various industry domains and an exhaustive list of sub-domains through our varied expertise of consultants having more than 15 years of experience in each industry vertical. With more than 300+ analysts and consultants on board, the company offers in-depth market analysis and helps clients take vital decisions impacting their revenues and growth roadmap."
                address.send_keys(address_content)
                print("✅ About and contact information filled")
                
                # Upload image - FIXED: Fresh element with better error handling
                if selected_image_path and os.path.exists(selected_image_path):
                    try:
                        image = wait.until(
                            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
                        )
                        image.clear()
                        image.send_keys(selected_image_path)
                        print(f"✅ Image uploaded: {selected_image_path}")
                    except Exception as e:
                        print(f"⚠️ Image upload failed: {e}")
                        # Try backup images
                        for backup_path in image_paths:
                            if backup_path != selected_image_path and os.path.exists(backup_path):
                                try:
                                    image.send_keys(backup_path)
                                    print(f"✅ Used backup image: {backup_path}")
                                    break
                                except:
                                    continue
                        else:
                            print(f"⚠️ No valid image paths found, continuing without image")
                else:
                    print(f"⚠️ Selected image file not found: {selected_image_path}")
                
                # Fill caption - FIXED: Fresh element
                caption = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
                )
                caption.clear()
                caption.send_keys(f"{market_name_from_file[3:]} Market Analysis")
                
                # Fill notes - FIXED: Fresh element
                notes = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
                )
                notes.clear()
                notes.send_keys(f"Comprehensive market research report on {market_name_from_file} with detailed analysis and forecasts.")
                
                # Agree to terms - FIXED: Fresh elements
                tick1 = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
                )
                tick1.click()
                
                tick2 = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
                )
                tick2.click()
                print("✅ Terms and conditions accepted")
                
                # Submit form - FIXED: Fresh element with extra wait
                print("🔍 Finding final submit button...")
                time.sleep(2)
                
                final = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
                )
                final.click()
                print("🚀 Form submitted!")
                
                # Wait for submission to complete
                time.sleep(8)  # Longer wait for submission
                
                print(f"✅ Successfully published: {filename}")
                successful_publications += 1
                
            except StaleElementReferenceException as se:
                print(f"❌ Stale element error for {filename}: {se}")
                print("💡 This happens when page elements change after finding them")
                failed_publications += 1
            except TimeoutException as te:
                print(f"❌ Timeout error for {filename}: {te}")
                print("💡 Page took too long to load or element not found")
                failed_publications += 1
            except Exception as e:
                print(f"❌ Error publishing {filename}: {e}")
                failed_publications += 1
            finally:
                # ✅ NEVER AUTO-CLOSE: Let user manually close browser
                print("✅ Article processing completed. Browser left open for manual review.")
                
                # Handle processing mode delays
                if processing_mode == 'manual' and i < len(article_files):
                    print(f"⏸️ Manual mode - Please review the submission...")
                    print(f"📋 You can manually close this tab when ready")
                    input("Press Enter to open NEW tab for next article...")
                elif processing_mode == 'auto' and i < len(article_files):
                    print(f"⏳ Auto mode - Article published. Opening new tab in 30 seconds...")
                    print(f"🌐 Previous tab left open for your reference")
                    time.sleep(15)
                
                # Handle processing mode delays
                if processing_mode == 'manual' and i < len(article_files):
                    print(f"⏸️ Manual mode - Waiting for user input...")
                    input("Press Enter to continue to next article...")
                elif processing_mode == 'auto' and i < len(article_files):
                    print(f"⏳ Auto mode - Waiting 30 seconds before next article...")
                    time.sleep(20)
        
        # Final summary
        print(f"\n{'='*60}")
        print("WMR PUBLISHING SUMMARY")
        print(f"{'='*60}")
        print(f"✅ Successfully published: {successful_publications}")
        print(f"❌ Failed publications: {failed_publications}")
        print(f"📊 Total articles processed: {len(article_files)}")
        print(f"👤 Published by: {author_name} ({author_email})")
        print(f"🏢 Company: {company_name}")
        print(f"🔧 Mode: {processing_mode}")
        print(f"📝 Title prompts loaded from Excel: {len(TITLE_PROMPTS)}")
        print(f"{'='*60}")
        
        return {
            'success': True,
            'total_articles': len(article_files),
            'successful_publications': successful_publications,
            'failed_publications': failed_publications
        }
        
    except Exception as e:
        print(f"❌ Error in WMR publishing automation: {e}")
        return {'success': False, 'error': str(e)}  
    
def extract_market_name_from_filename(filename):
    """Extract clean market name from WMR filename, removing WMR/AI prefixes"""
    try:
        print(f"🔍 Extracting market name from filename: {filename}")
        
        # Remove file extension
        name_without_ext = os.path.splitext(filename)[0]
        print(f"Without extension: {name_without_ext}")
        
        # Split by underscore
        parts = name_without_ext.split('_')
        print(f"Split parts: {parts}")
        
        if len(parts) >= 4:
            # Handle different filename patterns:
            # Pattern 1: WMR_Market_Name_2025_08_01
            # Pattern 2: WMR_AI_Market_Name_2025_08_01
            
            start_index = 1  # Skip "WMR"
            
            # If second part is "AI", skip that too
            if len(parts) > 1 and parts[1].upper() == 'AI':
                start_index = 2  # Skip both "WMR" and "AI"
                print("Found AI prefix, skipping both WMR and AI")
            
            # Remove last 3 parts (year, month, day)
            end_index = len(parts) - 3
            
            # Extract market name parts
            if start_index < end_index:
                market_parts = parts[start_index:end_index]
                print(f"Market parts: {market_parts}")
                
                # Join and capitalize properly
                market_name = ' '.join(market_parts)
                
                # Proper capitalization (Title Case)
                market_name = ' '.join(word.capitalize() for word in market_name.split())
                
                print(f"Cleaned market name: {market_name}")
                return market_name
            else:
                print("⚠️ Not enough parts to extract market name")
                # Fallback to using the original filename
                fallback_name = name_without_ext.replace('_', ' ').title()
                print(f"Using fallback: {fallback_name}")
                return fallback_name
        else:
            # Fallback: use filename as is with proper capitalization
            fallback_name = name_without_ext.replace('_', ' ').title()
            print(f"Using fallback (insufficient parts): {fallback_name}")
            return fallback_name
            
    except Exception as e:
        print(f"❌ Error extracting market name from {filename}: {e}")
        # Emergency fallback
        fallback_name = filename.replace('_', ' ').replace('.docx', '').replace('.doc', '').title()
        print(f"Emergency fallback: {fallback_name}")
        return fallback_name

def text_of_press_release(doc_path):
    """Extract and format press release text from Word document"""
    import win32com.client
    import re
    
    # Open Word application
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Run in the background

    # Open the .doc file (adjust the file path if necessary)
    doc2 = word.Documents.Open(doc_path)

    # Extract the entire text from the document
    doc_text = doc2.Content.Text

  

    # Define the headings for which you want to add line breaks
    headings = [
        "➤Market Size and Overview",
        "➤Actionable Insights",
        "➤Actionable insights",
        "➤Growth factors",
        "➤Growth Factors",
        "➤Market trends",
        "➤Market Trends",
        "➤Key takeaways ",
        "➤Key Takeaways",
        "➤Market Segment and Regional Coverage ",
        "➤Market segment and regional coverage",
        "➤Key players",
        "➤Key Players",
        "➤Competitive Strategies and Outcomes",
        "❓ Frequently Asked Questions",
        "❓ Frequently asked questions",
        "➤ Frequently asked questions",
        "➤ Frequently Asked Questions"
    ]

    # FIXED: Add a line space AFTER each heading (not before and after)
    for heading in headings:
        doc_text = doc_text.replace(heading, f"{heading}\n")

    # Define the regex pattern for URLs
    url_pattern = re.compile(r"(https?://[^\s]+)")
    
    # Define regex patterns for FAQ questions (numbered questions and roman numerals)
    faq_pattern_numbers = re.compile(r"^\d+\.\s")  # Matches "1. ", "2. ", etc.
    faq_pattern_roman = re.compile(r"^[ivxlcdmIVXLCDM]+\.\s")  # Matches "i. ", "ii. ", "I. ", "II. ", etc.
    
    # Define regex pattern for CTA links (➔)
    cta_pattern = re.compile(r"^➔")  # Matches lines starting with ➔

    # Split the text into lines
    lines = doc_text.splitlines()
    processed_lines = []

    # Iterate over each line
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Skip empty lines in processing, we'll add them strategically
        if not line_stripped:
            continue
            
        # Check if this is a CTA line
        is_cta = cta_pattern.match(line_stripped)
        
        # Check if previous line was a CTA (for adjacent CTA handling)
        prev_was_cta = False
        if processed_lines:
            last_non_empty = None
            for prev_line in reversed(processed_lines):
                if prev_line.strip():
                    last_non_empty = prev_line.strip()
                    break
            if last_non_empty and cta_pattern.match(last_non_empty):
                prev_was_cta = True
        
        # Check if this line is a heading (starts with ➤ or ❓)
        is_heading = line_stripped.startswith('➤') or line_stripped.startswith('❓')
        
        # If a line contains a URL, add space before and after the URL
        if url_pattern.search(line):
            # Add space before (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            processed_lines.append('')  # Add space after
            
        # If a line is an FAQ question (starts with number or roman numeral), add space before it
        elif faq_pattern_numbers.match(line_stripped) or faq_pattern_roman.match(line_stripped):
            # Add space before FAQ question (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this is a CTA line
        elif is_cta:
            # Add space before CTA (unless previous was also CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this line is a heading (starts with ➤ or ❓)
        elif is_heading:
            # Add space before heading (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            # FIXED: Add space AFTER heading
            processed_lines.append('')
            
        else:
            # Regular content line
            processed_lines.append(line)

    # Join the processed lines back into a single string
    chunk = "\n".join(processed_lines)
    
    # Clean up multiple consecutive empty lines (replace with single empty line)
    chunk = re.sub(r'\n\s*\n\s*\n+', '\n\n', chunk)

    # Close the document
    doc2.Close()
    word.Quit()

    # Return the processed content
    return chunk


@app.route('/start_custom_cmi_publishing', methods=['POST'])

def start_custom_cmi_publishing():
    try:
        # Get form data from JSON request
        form_data = request.get_json()
        if not form_data:
            return jsonify({
                'status': 'error',
                'message': 'No form data received'
            }), 400
        
        print(f"Received CMI form data: {form_data}")
        
        # Validate required fields
        required_fields = ['author_name', 'author_email', 'company_name', 'phone_number']
        for field in required_fields:
            if not form_data.get(field):
                return jsonify({
                    'status': 'error',
                    'message': f'Missing required field: {field}'
                }), 400
        
        # Validate image paths
        if not form_data.get('image_paths') or len(form_data['image_paths']) == 0:
            return jsonify({
                'status': 'error',
                'message': 'At least one image path is required'
            })
        
        # Check CMI domain folders - CORRECTED PATH STRUCTURE
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        base_rpa_path = os.path.join(desktop_path, "RPA")
        custom_cmi_path = os.path.join(base_rpa_path, "Custom CMI")
        
        # CMI domain folders - Desktop/RPA/Custom CMI/[DOMAIN]/Generated
        cmi_folders = [
            os.path.join(custom_cmi_path, "HC", "Generated"),
            os.path.join(custom_cmi_path, "ICT", "Generated"),
            os.path.join(custom_cmi_path, "CMFE", "Generated")
        ]
        
        # Quick check for CMI files
        cmi_files_found = False
        total_files = 0
        found_folders = []
        
        for folder in cmi_folders:
            if os.path.exists(folder):
                folder_files = 0
                for file in os.listdir(folder):
                    if (file.lower().endswith(('.doc', '.docx')) and 
                        not file.startswith('~')):
                        folder_files += 1
                        cmi_files_found = True
                
                if folder_files > 0:
                    total_files += folder_files
                    domain_name = os.path.basename(os.path.dirname(folder))
                    found_folders.append(f"{domain_name}: {folder_files} files")
            else:
                domain_name = os.path.basename(os.path.dirname(folder))
                print(f"⚠️ Domain folder not found: {folder} ({domain_name})")
        
        if not cmi_files_found:
            return jsonify({
                'status': 'error',
                'message': 'No CMI articles found in HC/ICT/CMFE Generated folders. Please generate content first.'
            })
        
        print(f"Found {total_files} CMI files to publish in: {', '.join(found_folders)}")
        
        # Run the selenium automation in background
        def run_cmi_publishing():
            result = selenium_publishing_custom_cmi_enhanced(form_data, cmi_folders)
            print(f"CMI Publishing completed: {result}")
        
        # Start in background thread so user gets immediate response
        threading.Thread(target=run_cmi_publishing).start()
        
        return jsonify({
            'status': 'success',
            'message': f'CMI publishing started! Processing {total_files} articles from {len(found_folders)} folders. Check console for progress updates.'
        })
        
    except Exception as e:
        print(f"Error in start_custom_cmi_publishing: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Error starting CMI publishing: {str(e)}'
        }), 500
    
def load_cmi_title_prompts_from_excel():
    """Load CMI title prompts from Excel file"""
    try:
        title_file_path = r"C:\Users\vishwas\Desktop\RPA\Custom CMI\TITLES\TITLE_SHEET.xlsx"
        
        print(f"📖 Loading CMI title prompts from: {title_file_path}")
        
        if not os.path.exists(title_file_path):
            print(f"⚠️ CMI title Excel file not found: {title_file_path}")
            return []
        
        # Read the Excel file
        df = pd.read_excel(title_file_path, engine='openpyxl')
        
        # Try different possible column names (same as WMR)
        possible_columns = ['TITLES', 'Title', 'title', 'TITLE_PROMPTS', 'Prompts', 'PROMPTS']
        
        title_column = None
        for col in possible_columns:
            if col in df.columns:
                title_column = col
                break
        
        if title_column is None:
            # Use first column if no standard column found
            title_column = df.columns[0]
            print(f"⚠️ Using first column '{title_column}' for CMI titles")
        else:
            print(f"✅ Found CMI title column: '{title_column}'")
        
        # Extract titles and remove empty/null values
        titles = df[title_column].dropna().astype(str).tolist()
        titles = [title.strip() for title in titles if title.strip()]
        
        print(f"✅ Loaded {len(titles)} CMI title prompts from Excel")
        print(f"📋 Sample titles: {titles[:3]}...")
        
        return titles
        
    except Exception as e:
        print(f"❌ Error loading CMI title prompts from Excel: {e}")
        return []

def selenium_publishing_custom_cmi_enhanced(form_data, cmi_folders):
    """Enhanced CMI Publishing with OpenAI category selection and Excel title prompts"""
    try:
        import random
        import os
        from datetime import datetime
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        from selenium.common.exceptions import TimeoutException, StaleElementReferenceException
        from webdriver_manager.chrome import ChromeDriverManager
        import time
        import re
        
        print(f"\n=== STARTING ENHANCED CMI PUBLISHING ===")
        
        # Extract form data
        author_name = form_data.get('author_name', 'Vishwas Tiwari')
        author_email = form_data.get('author_email', 'vishwas@coherentmarketinsights.com')
        company_name = form_data.get('company_name', 'Coherent Market Insights')
        phone_number = form_data.get('phone_number', '1234567890')
        article_code = form_data.get('article_code', 'D5A-2025-QDFH8C')
        processing_mode = form_data.get('processing_mode', 'auto')
        image_paths = form_data.get('image_paths', [])
        
        print(f"Using form data - Author: {author_name}, Email: {author_email}")
        print(f"Company: {company_name}, Phone: {phone_number}")
        print(f"Article Code: {article_code}, Mode: {processing_mode}")
        print(f"Available image paths: {len(image_paths)}")
        
        # Load CMI title prompts from Excel
        TITLE_PROMPTS = load_cmi_title_prompts_from_excel()
        
       
        
        print(f"📝 Using {len(TITLE_PROMPTS)} CMI title prompts")
        
        # Hardcoded author descriptions (same as before)
        AUTHOR_DESCRIPTIONS = [
            """Author of this marketing PR:
Ravina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.""",
            """Author of this marketing PR:
Money Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc.""",
            """Author of this marketing PR:
Alice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights."""
        ]
        
        # Get all CMI articles from domain folders
        article_files = []
        for folder in cmi_folders:
            if not os.path.exists(folder):
                domain_name = os.path.basename(os.path.dirname(folder))
                print(f"⚠️ Domain folder not found: {folder} ({domain_name})")
                continue
                
            print(f"Scanning CMI domain folder: {folder}")
            for file in os.listdir(folder):
                if file.lower().endswith(('.doc', '.docx')) and not file.startswith('~'):
                    file_path = os.path.join(folder, file)
                    if os.path.isfile(file_path):
                        mod_time = os.path.getmtime(file_path)
                        # Extract domain name: Desktop/RPA/Custom CMI/[DOMAIN]/Generated
                        domain = os.path.basename(os.path.dirname(folder))  # HC, ICT, or CMFE
                        article_files.append((file_path, mod_time, file, domain))
        
        if not article_files:
            print(f"❌ No CMI articles found in domain folders")
            return {'success': False, 'error': 'No articles found in domain Generated folders'}
        
        # Sort by modification time (latest first)
        article_files.sort(key=lambda x: x[1], reverse=True)
        
        print(f"✅ Found {len(article_files)} CMI articles to publish")
        for i, (path, mod_time, filename, domain) in enumerate(article_files, 1):
            print(f"  {i}. {filename} (from {domain}) - {datetime.fromtimestamp(mod_time).strftime('%Y-%m-%d %H:%M')}")
        
        successful_publications = 0
        failed_publications = 0
        
        # Setup Chrome driver
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        options.add_argument("--disable-blink-features=AutomationControlled")
        
        print(f"📋 NOTE: Browser tabs will remain open after publishing - please close them manually when done")
        
        # Process each article
        for i, (article_path, mod_time, filename, domain) in enumerate(article_files, 1):
            driver = None
            try:
                print(f"\n{'='*60}")
                print(f"Processing CMI Article {i}/{len(article_files)}: {filename} (Domain: {domain})")
                print(f"{'='*60}")
                
                # Randomly select image path for this article
                selected_image_path = random.choice(image_paths) if image_paths else None
                print(f"🖼️ Randomly selected image path: {selected_image_path}")
                
                # Extract market name from filename
                market_name_from_file = extract_market_name_from_filename(filename)
                print(f"🏷️ Market name extracted: {market_name_from_file}")
                
                # Read article content using text_of_press_release
                print(f"📖 Reading article content from: {filename}")
                article_content = text_of_press_release(article_path)
                
                if not article_content or len(article_content.strip()) < 100:
                    print(f"❌ Article content too short or empty, skipping: {filename}")
                    failed_publications += 1
                    continue
                
                print(f"✅ Article content loaded: {len(article_content)} characters")
                
                # Generate article title using Excel-loaded prompts
                random_prompt = random.choice(TITLE_PROMPTS)
                print(f"🎯 Selected title prompt from Excel: {random_prompt}")
                
                # Try to generate creative title with OpenAI
                creative_title = generate_creative_title_with_openai(market_name_from_file, random_prompt)
                
                if creative_title and len(creative_title.strip()) > 10:
                    article_title = creative_title
                    print(f"✨ Using OpenAI generated title: {article_title}")
                else:
                    # Fallback to direct combination
                    article_title = f"{market_name_from_file} {random_prompt}"
                    print(f"📝 Using fallback title: {article_title}")

                # Start Selenium automation
                print(f"🚀 Starting Selenium automation for: {market_name_from_file}")
                cService = Service(executable_path=chromedriver_path)
                driver = webdriver.Chrome(service=cService, options=options)
                
                # Set up longer default wait
                wait = WebDriverWait(driver, 20)
                
                driver.get('https://www.openpr.com/')
                print("🌐 Navigated to OpenPR.com")
                
                # Wait for page to fully load
                time.sleep(2)  # 2s after page transition
                
                # Handle cookie consent
                try:
                    reject = wait.until(
                        EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
                    )
                    reject.click()
                    print("✅ Cookie consent handled")
                    time.sleep(1)  # 1s for quick operation
                except TimeoutException:
                    print("⚠️ Cookie consent button not found or already handled")
                
                # Navigate to submit page
                submit = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
                )
                submit.click()
                print("✅ Navigated to submit page")
                time.sleep(2)  # 2s after page transition
                
                # Enter article code
                print(f"🔐 Entering article code...")
                input_box = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="code"]'))
                )
                input_box.clear()
                input_box.send_keys(article_code)
                print(f"✅ Entered article code: {article_code}")
                time.sleep(1)  # 1s for quick operation
                
                # Submit code
                print(f"📤 Submitting code...")
                try:
                    submit2 = wait.until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
                    )
                    submit2.click()
                except TimeoutException:
                    submit2 = wait.until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
                    )
                    submit2.click()
                
                print("✅ Article code submitted")
                time.sleep(2)  # 2s after page transition
                
                # Fill form fields with fresh elements (prevent stale element errors)
                print("📝 Filling form fields...")
                
                name = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
                )
                name.clear()
                name.send_keys(author_name)
                print("✅ Name field filled")
                time.sleep(1)  # 1s for quick operation
                
                email = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
                )
                email.clear()
                email.send_keys(author_email)
                print("✅ Email field filled")
                time.sleep(1)  # 1s for quick operation
                
                number = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
                )
                number.clear()
                number.send_keys(phone_number)
                print("✅ Phone field filled")
                time.sleep(1)  # 1s for quick operation
                
                # Company name with popup handling
                ComName = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="archivnmfield"]'))
                )
                ComName.clear()
                ComName.send_keys(company_name)
                time.sleep(1)  # Wait for popup to appear
                
                # Handle company popup
                try:
                    s1 = wait.until(
                        EC.element_to_be_clickable((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
                    )
                    s1.click()
                    print("✅ Company name selected")
                    time.sleep(1)  # 1s for quick operation
                except TimeoutException:
                    print("⚠️ Company popup not found, continuing...")
                
                # 🔥 ENHANCED: Handle category selection with OpenAI
                print("🤖 Getting category using OpenAI...")
                time.sleep(1)  # 1s for quick operation
                
                # Use OpenAI to get the best category for this market
                try:
                    category_to_use = get_keyword_category_from_openai(market_name_from_file)
                    if not category_to_use:
                        # Fallback to business category if OpenAI fails
                        category_to_use = "Business,Economy,Finance,Banking & Insurance"
                        print("⚠️ OpenAI category failed, using business category fallback")
                    else:
                        print(f"🎯 OpenAI suggested category: {category_to_use}")
                except Exception as e:
                    print(f"⚠️ OpenAI category error: {e}")
                    category_to_use = "Business,Economy,Finance,Banking & Insurance"
                    print("🔄 Using business category fallback")
                
                # Select category with retry mechanism to handle stale elements
                Category_element = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
                )
                time.sleep(1)  # 1s for quick operation
                
                select_obj = Select(Category_element)
                category_selected = False
                max_retries = 3
                
                for retry in range(max_retries):
                    try:
                        select_obj.select_by_visible_text(category_to_use)
                        print(f"✅ Selected category: {category_to_use}")
                        category_selected = True
                        break
                    except StaleElementReferenceException:
                        print(f"⚠️ Stale element on category selection attempt {retry+1}, retrying...")
                        time.sleep(2)  # 2s delay for stale element recovery
                        # Re-find the element
                        Category_element = wait.until(
                            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
                        )
                        select_obj = Select(Category_element)
                    except Exception as e:
                        print(f"⚠️ Category selection attempt {retry+1} failed: {e}")
                        if retry < max_retries - 1:
                            time.sleep(2)  # 2s delay before retry
                            # Re-find the element
                            Category_element = wait.until(
                                EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
                            )
                            select_obj = Select(Category_element)
                        else:
                            # Final fallback - select any option
                            select_obj.select_by_index(1)
                            print("✅ Used fallback category selection (index 1)")
                            category_selected = True
                
                if not category_selected:
                    print("❌ Failed to select any category")
                    failed_publications += 1
                    continue
                
                time.sleep(1)  # 1s after category selection
                
                # Fill title
                title = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
                )
                title.clear()
                title.send_keys(article_title)
                print(f"✅ Entered title: {len(article_title)} characters")
                time.sleep(1)  # 1s for quick operation
                
                # Fill article content (main text)
                text = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="inhalt"]'))
                )
                text.clear()
                text.send_keys(article_content)
                print(f"✅ Entered article content: {len(article_content)} characters")
                time.sleep(1)  # 1s for quick operation
                
                # Fill about section
                about = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
                )
                about.clear()
                contact_info = f"""Contact Us:

Mr. Shah
Coherent Market Insights
533 Airport Boulevard,
Suite 400, Burlingame,
CA 94010, United States
US: + 12524771362
UK: +442039578553
AUS: +61-8-7924-7805
India: +91-848-285-0837"""
                about.send_keys(contact_info)
                time.sleep(1)  # 1s for quick operation
                
                # Fill address section with random author
                address = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
                )
                address.clear()
                random_author = random.choice(AUTHOR_DESCRIPTIONS)
                address_content = f"{random_author}\n\nAbout Us:\n{company_name} leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries."
                address.send_keys(address_content)
                print("✅ About and contact information filled")
                time.sleep(1)  # 1s for quick operation
                
                # Upload image with better error handling
                if selected_image_path and os.path.exists(selected_image_path):
                    try:
                        image = wait.until(
                            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
                        )
                        image.clear()
                        image.send_keys(selected_image_path)
                        print(f"✅ Image uploaded: {os.path.basename(selected_image_path)}")
                    except Exception as e:
                        print(f"⚠️ Image upload failed: {e}")
                        # Try backup images
                        for backup_path in image_paths:
                            if backup_path != selected_image_path and os.path.exists(backup_path):
                                try:
                                    image.send_keys(backup_path)
                                    print(f"✅ Used backup image: {os.path.basename(backup_path)}")
                                    break
                                except:
                                    continue
                        else:
                            print(f"⚠️ No valid image paths found, continuing without image")
                else:
                    print(f"⚠️ No valid image selected, continuing without image")
                
                time.sleep(1)  # 1s for quick operation
                
                # Fill caption
                caption = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
                )
                caption.clear()
                caption.send_keys(f"{market_name_from_file[3:]}  Analysis")
                time.sleep(1)  # 1s for quick operation
                
                # Fill notes
                notes = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
                )
                notes.clear()
                notes.send_keys(f"Comprehensive market research report on {market_name_from_file} with detailed analysis and forecasts by Coherent Market Insights.")
                time.sleep(1)  # 1s for quick operation
                
                # Agree to terms
                tick1 = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
                )
                tick1.click()
                time.sleep(1)  # 1s for quick operation
                
                tick2 = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
                )
                tick2.click()
                print("✅ Terms and conditions accepted")
                time.sleep(1)  # 1s for quick operation
                
                # Submit form
                print("🚀 Submitting form...")
                final = wait.until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
                )
                final.click()
                print("🚀 Form submitted!")
                
                # Wait for submission to complete
                time.sleep(8)  # Longer wait for submission
                
                print(f"✅ Successfully published: {filename} from {domain}")
                successful_publications += 1
                
                # Keep browser open - user will close manually
                # driver.quit()  # Commented out - user will close tabs manually
                
            except StaleElementReferenceException as se:
                print(f"❌ Stale element error for {filename}: {se}")
                print("💡 This happens when page elements change after finding them")
                failed_publications += 1
            except TimeoutException as te:
                print(f"❌ Timeout error for {filename}: {te}")
                print("💡 Page took too long to load or element not found")
                failed_publications += 1
            except Exception as e:
                print(f"❌ Error publishing {filename}: {e}")
                failed_publications += 1
            finally:
                # Handle processing mode delays
                if processing_mode == 'manual' and i < len(article_files):
                    print(f"⏸️ Manual mode - Please review the submission...")
                    input("Press Enter to continue to next article...")
                elif processing_mode == 'auto' and i < len(article_files):
                    print(f"⏳ Auto mode - Waiting 30 seconds before next article...")
                    time.sleep(30)
        
        # Final summary
        print(f"\n{'='*60}")
        print("ENHANCED CMI PUBLISHING SUMMARY")
        print(f"{'='*60}")
        print(f"✅ Successfully published: {successful_publications}")
        print(f"❌ Failed publications: {failed_publications}")
        print(f"📊 Total articles processed: {len(article_files)}")
        print(f"👤 Published by: {author_name} ({author_email})")
        print(f"🏢 Company: {company_name}")
        print(f"🔧 Mode: {processing_mode}")
        print(f"📁 Domains processed: {set([domain for _, _, _, domain in article_files])}")
        print(f"📝 Excel title prompts used: {len(TITLE_PROMPTS)}")
        print(f"🤖 OpenAI category selection enabled")
        print(f"⏱️ Stale element prevention delays added")
        print(f"🌐 Browser tabs left open for manual closure")
        print(f"{'='*60}")
        
        return {
            'success': True,
            'total_articles': len(article_files),
            'successful_publications': successful_publications,
            'failed_publications': failed_publications
        }
        
    except Exception as e:
        print(f"❌ Error in enhanced CMI publishing automation: {e}")
        return {'success': False, 'error': str(e)}
'''
# 🔥 UPDATE: Replace the function call in start_custom_cmi_publishing
@app.route('/start_custom_cmi_publishing', methods=['POST'])
def start_custom_cmi_publishing():
    try:
        # Get form data from JSON request
        form_data = request.get_json()
        if not form_data:
            return jsonify({
                'status': 'error',
                'message': 'No form data received'
            }), 400
        
        print(f"Received CMI form data: {form_data}")
        
        # Validate required fields
        required_fields = ['author_name', 'author_email', 'company_name', 'phone_number']
        for field in required_fields:
            if not form_data.get(field):
                return jsonify({
                    'status': 'error',
                    'message': f'Missing required field: {field}'
                }), 400
        
        # Validate image paths
        if not form_data.get('image_paths') or len(form_data['image_paths']) == 0:
            return jsonify({
                'status': 'error',
                'message': 'At least one image path is required'
            })
        
        # Check CMI domain folders
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        base_rpa_path = os.path.join(desktop_path, "RPA")
        custom_cmi_path = os.path.join(base_rpa_path, "Custom CMI")
        
        # CMI domain folders
        cmi_folders = [
            os.path.join(custom_cmi_path, "HC", "Generated"),
            os.path.join(custom_cmi_path, "ICT", "Generated"),
            os.path.join(custom_cmi_path, "CMFE", "Generated")
        ]
        
        # Quick check for CMI files
        cmi_files_found = False
        total_files = 0
        found_folders = []
        
        for folder in cmi_folders:
            if os.path.exists(folder):
                folder_files = 0
                for file in os.listdir(folder):
                    if (file.lower().endswith(('.doc', '.docx')) and 
                        not file.startswith('~')):
                        folder_files += 1
                        cmi_files_found = True
                
                if folder_files > 0:
                    total_files += folder_files
                    domain_name = os.path.basename(os.path.dirname(folder))
                    found_folders.append(f"{domain_name}: {folder_files} files")
            else:
                domain_name = os.path.basename(os.path.dirname(folder))
                print(f"⚠️ Domain folder not found: {folder} ({domain_name})")
        
        if not cmi_files_found:
            return jsonify({
                'status': 'error',
                'message': 'No CMI articles found in HC/ICT/CMFE Generated folders. Please generate content first.'
            })
        
        print(f"Found {total_files} CMI files to publish in: {', '.join(found_folders)}")
        
        # 🔥 CHANGED: Use the enhanced function
        def run_cmi_publishing():
            result = selenium_publishing_custom_cmi_enhanced(form_data, cmi_folders)
            print(f"Enhanced CMI Publishing completed: {result}")
        
        # Start in background thread
        threading.Thread(target=run_cmi_publishing).start()
        
        return jsonify({
            'status': 'success',
            'message': f'Enhanced CMI publishing started! Processing {total_files} articles with OpenAI category selection and Excel title prompts. Check console for progress updates.'
        })
        
    except Exception as e:
        print(f"Error in start_custom_cmi_publishing: {e}")
        return jsonify({
            'status': 'error',
            'message': f'Error starting CMI publishing: {str(e)}'
        }), 500'''
    
def extract_market_name_from_filename(filename):
    """Extract market name from filename by removing common suffixes"""
    # Remove file extension
    name = os.path.splitext(filename)[0]
    
    # Remove common patterns
    patterns_to_remove = [
        r'_cmi_\d{4}_\d{2}_\d{2}$',  # Remove _cmi_YYYY_MM_DD
        r'_wmr_\d{4}_\d{2}_\d{2}$',  # Remove _wmr_YYYY_MM_DD
        r'_\d{4}_\d{2}_\d{2}$',      # Remove _YYYY_MM_DD
        r'_generated$',               # Remove _generated
        r'_final$',                   # Remove _final
        r'_report$'                   # Remove _report
    ]
    
    for pattern in patterns_to_remove:
        name = re.sub(pattern, '', name, flags=re.IGNORECASE)
    
    # Replace underscores with spaces and title case
    name = name.replace('_', ' ').strip().title()
    
    # If name is empty or too short, use filename
    if not name or len(name) < 3:
        name = os.path.splitext(filename)[0].replace('_', ' ').title()
    
    return name


#CUSTOM NEW NEW

openai.api_key = "GU7KXl6ONWndCnJztZaRTxRFT3BlbkFJjgIYB9nhvbb2G1maC8XOB0QqW4BNpzlP-DmgsXi7l1ivu8oKC8hmo4pDIBblicYkX_cHll6bEA"

@app.route('/custom_wmr_gen_articles', methods=['POST'])
def custom_wmr_gen_articles():
    """Generate WMR articles with AI and Template split based on user percentages"""
    try:
        print("=== WMR ARTICLE GENERATION STARTED ===")
        
        # Get form parameters
        ai_percentage = int(request.form.get('ai_percentage', 50))
        template_percentage = int(request.form.get('template_percentage', 50))
        
        # Validate percentages
        if ai_percentage + template_percentage != 100:
            flash('❌ Percentages must add up to 100%!')
            return redirect('/custom_template_content')
        
        print(f"Generation Split: {ai_percentage}% AI, {template_percentage}% Template")
        
        # Handle file upload
        cta_file = request.files.get('cta_file')
        if not cta_file or cta_file.filename == '':
            flash('❌ Please select a CTA Excel file!')
            return redirect('/custom_template_content')
        
        # Save uploaded file
        filename = secure_filename(cta_file.filename)
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        cta_file.save(upload_path)
        
        # Read Excel file
        try:
            if filename.endswith('.csv'):
                df = pd.read_csv(upload_path,encoding='utf-8')
            else:
                df = pd.read_excel(upload_path, engine='openpyxl')
            print(f"Excel loaded: {len(df)} rows")
        except Exception as e:
            flash(f'❌ Error reading Excel file: {str(e)}')
            return redirect('/custom_template_content')
        
        # Validate required columns
        required_columns = ['KEYWORD', 'PROMOBUY', 'SAMPLECOPY']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            flash(f'❌ Missing required columns: {", ".join(missing_columns)}')
            return redirect('/custom_template_content')
        
        # Calculate split counts
        total_rows = len(df)
        template_count = int((template_percentage / 100) * total_rows)
        ai_count = total_rows - template_count
        
        print(f"Processing: {template_count} template articles, {ai_count} AI articles")
        
        # Setup output directories
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        wmr_path = os.path.join(desktop_path, "RPA", "WMR")
        template_path = wmr_path
        output_path = os.path.join(wmr_path, "Generated")
        
        # Create output directory if it doesn't exist
        if not os.path.exists(output_path):
            os.makedirs(output_path)
            print(f"Created output directory: {output_path}")
        
        # Check if template directory exists
        if not os.path.exists(template_path):
            flash('❌ Template directory not found: Desktop/RPA/WMR')
            return redirect('/custom_template_content')
        
        # Get available template files
        template_files = []
        for file in os.listdir(template_path):
            if file.endswith(('.doc', '.docx')) and not file.startswith('~'):
                template_files.append(os.path.join(template_path, file))
        
        if not template_files:
            flash('❌ No template files found in Desktop/RPA/WMR directory')
            return redirect('/custom_template_content')
        
        print(f"Found {len(template_files)} template files")
        
        # Split dataframe
        template_rows = df.iloc[:template_count].copy()
        ai_rows = df.iloc[template_count:].copy()
        
        success_count = 0
        error_count = 0
        
        # Process TEMPLATE articles
        for index, row in template_rows.iterrows():
            try:
                success = generate_template_article(row, template_files, output_path)
                if success:
                    success_count += 1
                else:
                    error_count += 1
            except Exception as e:
                print(f"Error processing template row {index}: {e}")
                error_count += 1
        
        # Process AI articles
        for index, row in ai_rows.iterrows():
            try:
                success = generate_ai_article(row, output_path)
                if success:
                    success_count += 1
                else:
                    error_count += 1
            except Exception as e:
                print(f"Error processing AI row {index}: {e}")
                error_count += 1
        
        # Clean up uploaded file
        try:
            os.remove(upload_path)
        except:
            pass
        
        # Show results
        if success_count > 0:
            flash(f'✅ Successfully generated {success_count} articles!')
            if error_count > 0:
                flash(f'⚠️ {error_count} articles failed to generate')
        else:
            flash('❌ No articles were generated successfully')
        
        return redirect('/custom_template_content')
        
    except Exception as e:
        print(f"Error in WMR article generation: {e}")
        flash(f'❌ Error: {str(e)}')
        return redirect('/custom_template_content')


def generate_template_article(row, template_files, output_path):
    """Generate article using random template"""
    try:
        keyword = str(row['KEYWORD']).strip()
        promobuy = str(row['PROMOBUY']).strip()
        samplecopy = str(row['SAMPLECOPY']).strip()
        
        # Strip "Market" from keyword to avoid duplication (case-insensitive)
        # This handles cases like "Portable Dishwasher Market" -> "Portable Dishwasher"
        if keyword.lower().endswith(' market'):
            clean_keyword = keyword[:-7]  # Remove last 7 characters (" market")
        elif keyword.lower().endswith('market'):
            clean_keyword = keyword[:-6]  # Remove last 6 characters ("market")
        else:
            clean_keyword = keyword
        
        print(f"Generating template article for: {keyword} -> cleaned: {clean_keyword}")
        
        # Randomly select template
        template_file = random.choice(template_files)
        print(f"Using template: {os.path.basename(template_file)}")
        
        # Read template document
        try:
            doc = Document(template_file)
        except Exception as e:
            print(f"Error reading template {template_file}: {e}")
            return False
        
        # Get additional placeholders from OpenAI (using cleaned keyword)
        segments = get_openai_placeholder("SEGMENTS", clean_keyword)
        applications = get_openai_placeholder("APPLICATIONS", clean_keyword)
        keyplayers = get_openai_placeholder("KEYPLAYERS", clean_keyword)
        
        if not segments or not applications or not keyplayers:
            print(f"Failed to get OpenAI placeholders for {clean_keyword}")
            return False
        
        # Replace placeholders in document (using CLEANED keyword)
        replacements = {
            'KEYWORD': clean_keyword,  # ⭐ FIXED: Use clean_keyword instead of keyword
            'PROMOBUY': promobuy,
            'SAMPLECOPY': samplecopy,
            'SEGMENTS': segments,
            'APPLICATIONS': applications,
            'KEYPLAYERS': keyplayers
        }
        
        print(f"✅ Using replacements: KEYWORD -> '{clean_keyword}'")
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Replace in tables
        for table in doc.tables:
            for row_cells in table.rows:
                for cell in row_cells.cells:
                    for placeholder, value in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
        
        # Generate output filename (using cleaned keyword for consistency)
        today = datetime.now().strftime("%Y_%m_%d")
        clean_keyword_filename = clean_keyword.replace(" ", "_").replace("-", "_").lower()
        output_filename = f"WMR_{clean_keyword_filename}_{today}.docx"
        output_file_path = os.path.join(output_path, output_filename)
        
        # Save document
        doc.save(output_file_path)
        print(f"Template article saved: {output_filename}")
        return True
        
    except Exception as e:
        print(f"Error in template article generation: {e}")
        return False
def generate_ai_article(row, output_path):
    """Generate article using OpenAI API"""
    try:
        keyword = str(row['KEYWORD']).strip()
        promobuy = str(row['PROMOBUY']).strip()
        samplecopy = str(row['SAMPLECOPY']).strip()
        
         # Strip "Market" from keyword to avoid duplication (case-insensitive)
        if keyword.lower().endswith(' market'):
            clean_keyword = keyword[:-7]  # Remove last 7 characters (" market")
        elif keyword.lower().endswith('market'):
            clean_keyword = keyword[:-6]  # Remove last 6 characters ("market")
        elif keyword.lower().endswith('Market'):
            clean_keyword = keyword[:-6]
        elif keyword.lower().endswith(' Market'):
            clean_keyword = keyword[:-7]
        else:
            clean_keyword = keyword
        
        print(f"Generating AI article for: {keyword} -> cleaned: {clean_keyword}")
        
        # Prepare the AI prompt with PROPER LINK FORMATTING INSTRUCTIONS
        ai_prompt = f"""
       
We are coherent market insights and going to publish this data on third part PR website, thus do not mention any competitors name (market intelligence firm, market research, consulting and data and analytics companies), and take ownership for market share, size, growth rate as this research done by our own research team however any real instances and stats you can use reference such as company or organisation or others as mentioned in detail explanation. "IMPORTANT: Do not mention competing market research companies (like other market research firms, consulting companies, or data analytics companies that compete with us). However, DO mention actual companies that operate in the {clean_keyword} market - these are the real players/manufacturers/service providers in this specific industry that should be listed in the Leading Companies section."

CRITICAL LINK FORMATTING INSTRUCTIONS:
- Do NOT use brackets for links like [text](url) 
- Use ONLY this format: ➔ Link Text: URL 
- First CTA must be: ➔ Get the Sample Copy of the Research Report: {samplecopy}
- Second CTA must be: ➔ Get Instant Access! Purchase Research Report and Receive up to 70% Discount: {promobuy}

Write the blog covering the below pointers: Start content with market size and overview paragraph (do not add heading and do not use bullet for this section), ➤Strategic Actionable Insights for the Market, ➔ add first CTA link here, ➤Market Taxonomy and Regional Coverage of Report, ➤Leading Companies of the Market, ➔ add Second CTA link here, ➤Key Growth Drivers Fueling Market Expansion, ➤Key Reasons for Buying the (insert market name here) Report ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ❓ Frequently Asked Questions, All these pointers should act as a heading to respective paragraphs, do not miss any bullet for above given. Consider the following points while generating content: Flow of information, all given secondary keywords must be covered, and there should be a heading for each paragraph or bullet pointers. I need this content to be very presentable format thus make sure there should be space after and before of each heading or CTA links (Strategic Actionable Insights for the Market, first CTA link, Market Taxonomy and Regional coverage of Report, Leading Companies of the Market, Second CTA link, Key Growth Drivers Fueling Market Expansion, Key Reasons for Buying the (insert market name here) Report, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, and Frequently Asked Questions), this will increase the readability. Cover content in in bullet pointers whenever possible each paragraph should be short. Our readers are already experts in the field, so always try to generate content that provides unique insights and value addition for experts. Thus, while generating each piece of content, it should be data-backed with actual instances and stats from recent years 2025 and 2024, covering maximum stats that increase the authenticity and are enough to support or make decisions based upon reading this blog. Do not add generic content that is already known to readers, such as definitions, advantages, disadvantages, or other generic content. Flow of information should be as below. Start the blog with writing 5 to 7 sentence ling paragraph start content with talking about Industry in 2 to 3 sentences and should be matches with heading of the blog. followed by 2 to 4 sentence on market size and CAGR formatted as The Global (Insert Market Name) Market size is estimated to be valued at USD (Identify and Insert Market Value for 2025) in 2025 (market size) and is expected to reach USD (Identify and Insert Market Value for 2032) by 2032 (market forecast), exhibiting a compound annual growth rate (CAGR) of (Identify and Insert CAGR)% from 2025 to 2032. Do not add heading and do not use bullet for this section. (Make sure all values such as market size, CAGR, key players should be carefully identified with research approach). Then Strategic Actionable Insights for the Market: In Strategic Actionable Insights for the Market, cover 3 to 4 Strategic Actionable Insights for the Market in 10 to 12 sentences each insights comprises of 2 t 3 sentences covering supporting data for each insights. Insights could be essential quantitative indicators that lead to market size. For example, supply-side indicators (ex. Production Capacity, Pricing, and Exports. but not necessary to add if not matches to given market), demand-side indicators (ex. Pricing, Imports, Various Use Cases across industries. but not necessary to add if not matches to given market), micro-indicators, nano-size indicators—focus more on the quantitative aspects. Each Strategic Actionable Insights for the Market must have two sentence stats or actual instance examples from the recent year to support each point given in Strategic Actionable Insights for the Market, so that each given point look complete and meaningful. Then First CTA link. Next part is Market Taxonomy and Regional coverage of Report where enlist the all subsegment under each segment categories and fragment region into given format. Identify Comprehensive Market Taxonomy of the Report: • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. ➤ Regional and Country Analysis: • North America: U.S. and Canada • Latin America: Brazil, Argentina, Mexico, and Rest of Latin America • Europe: Germany, U.K., Spain, France, Italy, Benelux, Denmark, Norway, Sweden, Russia, and Rest of Europe • Asia Pacific: China, Taiwan, India, Japan, South Korea, Indonesia, Malaysia, Philippines, Singapore, Australia, and Rest of Asia Pacific • Middle East & Africa: Bahrain, Kuwait, Oman, Qatar, Saudi Arabia, United Arab Emirates, Israel, South Africa, North Africa, Central Africa, and Rest of MEA. "Then Leading Companies of the Market: Identify and Enlist 12 to 20 highly relevant Leading Companies of the Market for the given market. FORMAT REQUIREMENT: List each company name on a separate line with bullet points (• Company Name). Do NOT put all company names in one sentence.". "IMPORTANT FOR LEADING COMPANIES: List only real, actual company names (like Apple, Samsung, Microsoft, etc.) - do NOT generate fake website domains like 'company.com' or 'website.com'. Use only legitimate business names that actually operate in this market." Furthermore, add 2-3 statements on competitive strategies adopted by a few Leading Companies of the Market, mentioning actual strategies and entities involved along with the actual outcome. Then Add Second CTA link. Key Reasons for Buying the (insert market name here) Report, and its exact content as shared in data. Key Growth Drivers Fueling Market Expansion: Growth factor heading and short paragraph (3-4 Key Growth Drivers Fueling Market Expansion covered under 10 to 12 sentences) with supporting stats or examples from the recent year in the content, each factors should be covered in two to three sentences thus entire Key Growth Drivers Fueling Market Expansion content will be covered in 10 to 12 sentences long. No sub bullet is needed in Growth Factor. Then Emerging Trends and Market Shift: Market Trend heading and short paragraphs with supporting stats or examples from the recent year in the content (No bullet needed for as opportunity are written in paragraph format). Then High-Impact Market Opportunities by Segment and Region: Provide 3 to 4 High-Impact Market Opportunities by Segment and Region, 2-3 opportunities based upon segment and one opportunity based upon region in a paragraph format. Each opportunity described in two to three sentences and supported by actual data. For each opportunity, identify a single segment and provide a short description of the opportunity within that segment. Similarly, highlight one region, along with a brief description of the opportunity in each regional market (No bullet needed for as opportunity are written in paragraph format). Make sure to not to mention our organization name or relevant terms anywhere in the output such as coherent market insights or our analyst team or our research team Given Market Name and Data: {clean_keyword}. 

 ➤Key Reasons for Buying the (insert market name here) Report: • Comprehensive analysis of the changing competitive landscape • Assists in decision-making processes for the businesses along with detailed strategic planning methodologies • The report offers forecast data and an assessment of the (insert market name here) • Helps in understanding the key product segments and their estimated growth rate • In-depth analysis of market drivers, restraints, trends, and opportunities • Comprehensive regional analysis of the (insert market name here) • Extensive profiling of the key stakeholders of the business sphere • Detailed analysis of the factors influencing the growth of the (insert market name here). From an SEO perspective, we need to cover all given keywords from the list below. However, they should appear naturally so that the content flow looks natural for the reader. Keyword List: market share, market size, market research, market insights, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, market challenges, market growth, market forecast, market companies, market players, market analysis, market drivers, market restraints, market scope, market dynamics, market segments, market report, market growth strategies, market revenue, industry size, industry share, industry trends, and business growth, furthermore - Market size and market report, market revenue, market share, trends keywords are mandatory to be added twice in content. In addition to the above requirement, in 5 places, add the actual market name along with the above keywords so that long-tail keywords will be generated. These long-tail keywords are market name + size, market name + report, market name + revenue, market name + share, market name + trends. Make sure all given keywords are naturally fit, do not try to infuse forcefully, flow of information should be natural and meaningful, furthermore make sure spelling and structure of sentences from generated output are grammatically correct. Furthermore, based on the market name, create a set of Frequently Asked Questions that are highly relevant and customized to the specific market. The sample Frequently Asked Questions below are for understanding purposes only. For the given market, questions can be completely replaced. However, please tailor the actual questions to the market name and the insights provided in the report: 1. Who are the dominant players in the (Market Name) market? 2. What will be the size of the (Market Name) market in the coming years? 3. Which end users industry has the largest growth opportunity? 4. How will market development trends evolve over the next five years? 5. What is the nature of the competitive landscape and challenges in the (Market Name) market? 6. What go-to-market strategies are commonly adopted in the (Market Name) market? Make sure to answer to all FAQs. In the case of country-level markets, please exclude the word 'Global' and High-Impact Market Opportunities by Segment and Region where other regions are mentioned. Make sure to add catchy bullet in generated output. I have shared the reference bullet with you. Make sure to add this bullet. For heading use these bullet- ), ➤Strategic Actionable Insights for the Market, ➔ add first CTA link here ➤ Market Taxonomy and Regional coverage of Report, ➤ Leading Companies of the Market, ➔ Inserted Second CTA link, ➤Key Reasons for Buying the (insert market name here) Report, ➤Key Growth Drivers Fueling Market Expansion, ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ➤ Frequently Asked Questions. Make sure do not miss any bullet including CTA bullet which is ➔. For subpointers under main headings use bullets which is in reference as provided- Strategic Actionable Insights for the Market ●, Market Taxonomy and Regional coverage of Report●, Leading Companies of the Market●. Make sure to use these bullets for given subpointers. Ensure proper bullet formatting so that each point is marked with a single bullet only, and avoid placing two bullets adjacent to each other.
"""
        
        # Make OpenAI API call
        try:
            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    
                    {"role": "user", "content": ai_prompt}
                ],
                max_tokens=3000,
                temperature=0.7
            )
            
            article_content = response.choices[0].message.content
            
        except Exception as e:
            print(f"OpenAI API error for {keyword}: {e}")
            return False
        
        # Create new document and add content
        doc = Document()
        
        # Split content into paragraphs and add to document
        paragraphs = article_content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Generate output filename
        today = datetime.now().strftime("%Y_%m_%d")
        clean_keyword_file = keyword.replace(" ", "_").replace("-", "_").lower()
        output_filename = f"WMR_AI_{clean_keyword_file}_{today}.docx"
        output_file_path = os.path.join(output_path, output_filename)
        
        # Save document
        doc.save(output_file_path)
        print(f"WMR AI article saved: {output_filename}")
        return True
        
    except Exception as e:
        print(f"Error in AI article generation: {e}")
        return False
    
def get_openai_placeholder(placeholder_type, keyword):
    """Get placeholder values from OpenAI API with bullet points"""
    try:
        if placeholder_type == "SEGMENTS":
            prompt = f"List 5-8 main market segments for {keyword} market. Provide only the segment names separated by commas."
        elif placeholder_type == "APPLICATIONS":
            prompt = f"List 5-8 main applications or use cases for {keyword}. Provide only the application names separated by commas."
        elif placeholder_type == "KEYPLAYERS":
            prompt = f"List 8-12 key companies/players in the {keyword} market. Provide only company names separated by commas."
        else:
            return None
        
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a market research expert. Provide concise, accurate information."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=200,
            temperature=0.5
        )
        
        result = response.choices[0].message.content.strip()
        print(f"Got {placeholder_type} for {keyword}: {result[:50]}...")
        
        # Convert comma-separated to bullet points format
        # Split by comma, strip whitespace, ADD BULLETS, and join with newlines
        items = [item.strip() for item in result.split(',')]
        formatted_result = '\n'.join([f"• {item}" for item in items])  # ⭐ FIXED: Added bullet points
        
        print(f"Formatted {placeholder_type} with bullets:\n{formatted_result}")
        
        # Add small delay to avoid rate limiting
        time.sleep(0.5)
        
        return formatted_result
        
    except Exception as e:
        print(f"Error getting {placeholder_type} for {keyword}: {e}")
        return None

# You'll also need to add this route to handle the GET request for the form
@app.route('/custom_template_content')
def custom_template_content():
    """Show WMR template content generation form"""
    return render_template('custom_template_content.html')


# Helper function to check allowed file types (if not already defined)
def allowed_file(filename):
    ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

## CUSTOM NEW NEW NEW
@app.route('/custom_cmi_gen_articles', methods=['POST'])
def custom_cmi_gen_articles():
    """Generate CMI articles with AI and Template split based on user percentages and categories"""
    try:
        print("=== CMI ARTICLE GENERATION STARTED ===")
        
        # Get form parameters
        ai_percentage = int(request.form.get('ai_percentage', 50))
        template_percentage = int(request.form.get('template_percentage', 50))
        
        # Validate percentages
        if ai_percentage + template_percentage != 100:
            flash('❌ Percentages must add up to 100%!')
            return redirect('/custom_cmi_content')
        
        print(f"Generation Split: {ai_percentage}% AI, {template_percentage}% Template")
        
        # Handle file upload
        cta_file = request.files.get('cta_file')
        if not cta_file or cta_file.filename == '':
            flash('❌ Please select a CTA Excel file!')
            return redirect('/custom_cmi_content')
        
        # Save uploaded file
        filename = secure_filename(cta_file.filename)
        upload_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        cta_file.save(upload_path)
        
        # Read Excel file
        try:
            if filename.endswith('.csv'):
                df = pd.read_csv(upload_path,encoding='utf-8')
            else:
                df = pd.read_excel(upload_path, engine='openpyxl')
            print(f"Excel loaded: {len(df)} rows")
        except Exception as e:
            flash(f'❌ Error reading Excel file: {str(e)}')
            return redirect('/custom_cmi_content')
        
        # Validate required columns
        required_columns = ['KEYWORD', 'PROMOBUY', 'SAMPLECOPY', 'Category']
        missing_columns = [col for col in required_columns if col not in df.columns]
        if missing_columns:
            flash(f'❌ Missing required columns: {", ".join(missing_columns)}')
            return redirect('/custom_cmi_content')
        
        # Calculate split counts (overall, regardless of category)
        total_rows = len(df)
        ai_count = int((ai_percentage / 100) * total_rows)
        template_count = total_rows - ai_count
        
        print(f"Processing: {ai_count} AI articles, {template_count} template articles")
        
        # Setup base directories
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        cmi_base_path = os.path.join(desktop_path, "RPA", "Custom CMI")
        
        # Check if base CMI directory exists
        if not os.path.exists(cmi_base_path):
            flash('❌ CMI base directory not found: Desktop/RPA/Custom CMI')
            return redirect('/custom_cmi_content')
        
        # Split dataframe - First X% rows for AI, remaining for Template
        ai_rows = df.iloc[:ai_count].copy()
        template_rows = df.iloc[ai_count:].copy()
        
        success_count = 0
        error_count = 0
        
        # Process AI articles (first X% rows)
        print(f"Processing {len(ai_rows)} AI articles...")
        for index, row in ai_rows.iterrows():
            try:
                success = generate_cmi_ai_article(row, cmi_base_path)
                if success:
                    success_count += 1
                else:
                    error_count += 1
            except Exception as e:
                print(f"Error processing AI row {index}: {e}")
                error_count += 1
        
        # Process TEMPLATE articles (remaining rows)
        print(f"Processing {len(template_rows)} template articles...")
        for index, row in template_rows.iterrows():
            try:
                success = generate_cmi_template_article(row, cmi_base_path)
                if success:
                    success_count += 1
                else:
                    error_count += 1
            except Exception as e:
                print(f"Error processing template row {index}: {e}")
                error_count += 1
        
        # Clean up uploaded file
        try:
            os.remove(upload_path)
        except:
            pass
        
        # Show results
        if success_count > 0:
            flash(f'✅ Successfully generated {success_count} articles!')
            if error_count > 0:
                flash(f'⚠️ {error_count} articles failed to generate')
        else:
            flash('❌ No articles were generated successfully')
        
        return redirect('/custom_cmi_content')
        
    except Exception as e:
        print(f"Error in CMI article generation: {e}")
        flash(f'❌ Error: {str(e)}')
        return redirect('/custom_cmi_content')


def generate_cmi_template_article(row, cmi_base_path):
    """Generate CMI article using random template based on category"""
    try:
        keyword = str(row['KEYWORD']).strip()
        promobuy = str(row['PROMOBUY']).strip()
        samplecopy = str(row['SAMPLECOPY']).strip()
        category = str(row['Category']).strip().upper()
        
        print(f"Generating CMI template article for: {keyword} (Category: {category})")
        
        # Validate category
        valid_categories = ['HC', 'ICT', 'CMFE']
        if category not in valid_categories:
            print(f"❌ Invalid category '{category}' for keyword '{keyword}'. Skipping.")
            return False
        
        # Setup category-specific paths
        category_path = os.path.join(cmi_base_path, category)
        output_path = os.path.join(category_path, "Generated")
        
        # Check if category directory exists
        if not os.path.exists(category_path):
            print(f"❌ Category directory not found: {category_path}")
            return False
        
        # Create output directory if it doesn't exist
        if not os.path.exists(output_path):
            os.makedirs(output_path)
            print(f"Created output directory: {output_path}")
        
        # Get available template files for this category
        template_files = []
        for file in os.listdir(category_path):
            if file.endswith(('.doc', '.docx')) and not file.startswith('~'):
                template_files.append(os.path.join(category_path, file))
        
        if not template_files:
            print(f"❌ No template files found in {category_path}")
            return False
        
        # Randomly select template
        template_file = random.choice(template_files)
        print(f"Using template: {os.path.basename(template_file)} from {category}")
        
        # Read template document
        try:
            doc = Document(template_file)
        except Exception as e:
            print(f"Error reading template {template_file}: {e}")
            return False
        
        # Get additional placeholders from OpenAI
        segments = get_openai_placeholder("SEGMENTS", keyword)
        applications = get_openai_placeholder("APPLICATIONS", keyword)
        keyplayers = get_openai_placeholder("KEYPLAYERS", keyword)
        
        if not segments or not applications or not keyplayers:
            print(f"Failed to get OpenAI placeholders for {keyword}")
            return False
        
        # Replace placeholders in document
        replacements = {
            'KEYWORD': keyword,
            'PROMOBUY': promobuy,
            'SAMPLECOPY': samplecopy,
            'SEGMENTS': segments,
            'APPLICATIONS': applications,
            'KEYPLAYERS': keyplayers,
            'BUYNOW': promobuy
        }
        
        # Replace in paragraphs
        for paragraph in doc.paragraphs:
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
        
        # Replace in tables
        for table in doc.tables:
            for row_cells in table.rows:
                for cell in row_cells.cells:
                    for placeholder, value in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
        
        # Generate output filename
        today = datetime.now().strftime("%Y_%m_%d")
        clean_keyword = keyword.replace(" ", "_").replace("-", "_").lower()
        output_filename = f"CMI_{clean_keyword}_{today}.docx"
        output_file_path = os.path.join(output_path, output_filename)
        
        # Save document
        doc.save(output_file_path)
        print(f"CMI template article saved: {output_filename} in {category}")
        return True
        
    except Exception as e:
        print(f"Error in CMI template article generation: {e}")
        return False

def generate_cmi_ai_article(row, cmi_base_path):
    """Generate CMI article using OpenAI API based on category"""
    try:
        keyword = str(row['KEYWORD']).strip()
        promobuy = str(row['PROMOBUY']).strip()
        samplecopy = str(row['SAMPLECOPY']).strip()
        category = str(row['Category']).strip().upper()

        # ADD THIS BLOCK HERE:
        # Strip "Market" from keyword to avoid duplication (case-insensitive)
        if keyword.lower().endswith(' market'):
            clean_keyword = keyword[:-7]  # Remove last 7 characters (" market")
        elif keyword.lower().endswith('market'):
            clean_keyword = keyword[:-6]  # Remove last 6 characters ("market")
        elif keyword.lower().endswith('Market'):
            clean_keyword = keyword[:-6]
        elif keyword.lower().endswith(' Market'):
            clean_keyword = keyword[:-7]
        else:
            clean_keyword = keyword
        
        print(f"Generating CMI AI article for: {keyword} (Category: {category})")
        
        # Validate category
        valid_categories = ['HC', 'ICT', 'CMFE']
        if category not in valid_categories:
            print(f"❌ Invalid category '{category}' for keyword '{keyword}'. Skipping.")
            return False
        
        # Setup category-specific paths
        category_path = os.path.join(cmi_base_path, category)
        output_path = os.path.join(category_path, "Generated")
        
        # Check if category directory exists
        if not os.path.exists(category_path):
            print(f"❌ Category directory not found: {category_path}")
            return False
        
        # Create output directory if it doesn't exist
        if not os.path.exists(output_path):
            os.makedirs(output_path)
            print(f"Created output directory: {output_path}")
        
        # Prepare the AI prompt with PROPER LINK FORMATTING INSTRUCTIONS
        ai_prompt = f"""
              
We are coherent market insights and going to publish this data on third part PR website, thus do not mention any competitors name (market intelligence firm, market research, consulting and data and analytics companies), and take ownership for market share, size, growth rate as this research done by our own research team however any real instances and stats you can use reference such as company or organisation or others as mentioned in detail explanation. "IMPORTANT: Do not mention competing market research companies (like other market research firms, consulting companies, or data analytics companies that compete with us). However, DO mention actual companies that operate in the {clean_keyword} market - these are the real players/manufacturers/service providers in this specific industry that should be listed in the Leading Companies section."

CRITICAL LINK FORMATTING INSTRUCTIONS:
- Do NOT use brackets for links like [text](url) 
- Use ONLY this format: ➔ Link Text: URL
- First CTA must be: ➔ Get the Sample Copy of the Research Report: {samplecopy}
- Second CTA must be: ➔ Get Instant Access! Purchase Research Report and Receive up to 70% Discount: {promobuy}
- The second CTA with promobuy must always be the last link in the article

Write the blog covering the below pointers: Start content with market size and overview paragraph (do not add heading and do not use bullet for this section), ➤Strategic Actionable Insights for the Market, ➔ add first CTA link here, ➤Market Taxonomy and Regional Coverage of Report, ➤Leading Companies of the Market, ➔ add Second CTA link here, ➤Key Growth Drivers Fueling Market Expansion, ➤Key Reasons for Buying the (insert market name here) Report ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ❓ Frequently Asked Questions, All these pointers should act as a heading to respective paragraphs, do not miss any bullet for above given. Consider the following points while generating content: Flow of information, all given secondary keywords must be covered, and there should be a heading for each paragraph or bullet pointers. I need this content to be very presentable format thus make sure there should be space after and before of each heading or CTA links (Strategic Actionable Insights for the Market, first CTA link, Market Taxonomy and Regional coverage of Report, Leading Companies of the Market, Second CTA link, Key Growth Drivers Fueling Market Expansion, Key Reasons for Buying the (insert market name here) Report, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, and Frequently Asked Questions), this will increase the readability. Cover content in in bullet pointers whenever possible each paragraph should be short. Our readers are already experts in the field, so always try to generate content that provides unique insights and value addition for experts. Thus, while generating each piece of content, it should be data-backed with actual instances and stats from recent years 2025 and 2024, covering maximum stats that increase the authenticity and are enough to support or make decisions based upon reading this blog. Do not add generic content that is already known to readers, such as definitions, advantages, disadvantages, or other generic content. Flow of information should be as below. Start the blog with writing 5 to 7 sentence ling paragraph start content with talking about Industry in 2 to 3 sentences and should be matches with heading of the blog. followed by 2 to 4 sentence on market size and CAGR formatted as The Global (Insert Market Name) Market size is estimated to be valued at USD (Identify and Insert Market Value for 2025) in 2025 (market size) and is expected to reach USD (Identify and Insert Market Value for 2032) by 2032 (market forecast), exhibiting a compound annual growth rate (CAGR) of (Identify and Insert CAGR)% from 2025 to 2032. Do not add heading and do not use bullet for this section. (Make sure all values such as market size, CAGR, key players should be carefully identified with research approach). Then Strategic Actionable Insights for the Market: In Strategic Actionable Insights for the Market, cover 3 to 4 Strategic Actionable Insights for the Market in 10 to 12 sentences each insights comprises of 2 t 3 sentences covering supporting data for each insights. Insights could be essential quantitative indicators that lead to market size. For example, supply-side indicators (ex. Production Capacity, Pricing, and Exports. but not necessary to add if not matches to given market), demand-side indicators (ex. Pricing, Imports, Various Use Cases across industries. but not necessary to add if not matches to given market), micro-indicators, nano-size indicators—focus more on the quantitative aspects. Each Strategic Actionable Insights for the Market must have two sentence stats or actual instance examples from the recent year to support each point given in Strategic Actionable Insights for the Market, so that each given point look complete and meaningful. Then First CTA link. Next part is Market Taxonomy and Regional coverage of Report where enlist the all subsegment under each segment categories and fragment region into given format. Identify Comprehensive Market Taxonomy of the Report: • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. • By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. ➤ Regional and Country Analysis: • North America: U.S. and Canada • Latin America: Brazil, Argentina, Mexico, and Rest of Latin America • Europe: Germany, U.K., Spain, France, Italy, Benelux, Denmark, Norway, Sweden, Russia, and Rest of Europe • Asia Pacific: China, Taiwan, India, Japan, South Korea, Indonesia, Malaysia, Philippines, Singapore, Australia, and Rest of Asia Pacific • Middle East & Africa: Bahrain, Kuwait, Oman, Qatar, Saudi Arabia, United Arab Emirates, Israel, South Africa, North Africa, Central Africa, and Rest of MEA. "Then Leading Companies of the Market: Identify and Enlist 12 to 20 highly relevant Leading Companies of the Market for the given market. FORMAT REQUIREMENT: List each company name on a separate line with bullet points (• Company Name). Do NOT put all company names in one sentence.""IMPORTANT FOR LEADING COMPANIES: List only real, actual company names (like Apple, Samsung, Microsoft, etc.) - do NOT generate fake website domains like 'company.com' or 'website.com'. Use only legitimate business names that actually operate in this market." Furthermore, add 2-3 statements on competitive strategies adopted by a few Leading Companies of the Market, mentioning actual strategies and entities involved along with the actual outcome. Then Add Second CTA link. Key Reasons for Buying the (insert market name here) Report, and its exact content as shared in data. Key Growth Drivers Fueling Market Expansion: Growth factor heading and short paragraph (3-4 Key Growth Drivers Fueling Market Expansion covered under 10 to 12 sentences) with supporting stats or examples from the recent year in the content, each factors should be covered in two to three sentences thus entire Key Growth Drivers Fueling Market Expansion content will be covered in 10 to 12 sentences long. No sub bullet is needed in Growth Factor. Then Emerging Trends and Market Shift: Market Trend heading and short paragraphs with supporting stats or examples from the recent year in the content (No bullet needed for as opportunity are written in paragraph format). Then High-Impact Market Opportunities by Segment and Region: Provide 3 to 4 High-Impact Market Opportunities by Segment and Region, 2-3 opportunities based upon segment and one opportunity based upon region in a paragraph format. Each opportunity described in two to three sentences and supported by actual data. For each opportunity, identify a single segment and provide a short description of the opportunity within that segment. Similarly, highlight one region, along with a brief description of the opportunity in each regional market (No bullet needed for as opportunity are written in paragraph format). Make sure to not to mention our organization name or relevant terms anywhere in the output such as coherent market insights or our analyst team or our research team Given Market Name and Data: {keyword}. 

 ➤Key Reasons for Buying the (insert market name here) Report: • Comprehensive analysis of the changing competitive landscape • Assists in decision-making processes for the businesses along with detailed strategic planning methodologies • The report offers forecast data and an assessment of the (insert market name here) • Helps in understanding the key product segments and their estimated growth rate • In-depth analysis of market drivers, restraints, trends, and opportunities • Comprehensive regional analysis of the (insert market name here) • Extensive profiling of the key stakeholders of the business sphere • Detailed analysis of the factors influencing the growth of the (insert market name here). From an SEO perspective, we need to cover all given keywords from the list below. However, they should appear naturally so that the content flow looks natural for the reader. Keyword List: market share, market size, market research, market insights, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, market challenges, market growth, market forecast, market companies, market players, market analysis, market drivers, market restraints, market scope, market dynamics, market segments, market report, market growth strategies, market revenue, industry size, industry share, industry trends, and business growth, furthermore - Market size and market report, market revenue, market share, trends keywords are mandatory to be added twice in content. In addition to the above requirement, in 5 places, add the actual market name along with the above keywords so that long-tail keywords will be generated. These long-tail keywords are market name + size, market name + report, market name + revenue, market name + share, market name + trends. Make sure all given keywords are naturally fit, do not try to infuse forcefully, flow of information should be natural and meaningful, furthermore make sure spelling and structure of sentences from generated output are grammatically correct. Furthermore, based on the market name, create a set of Frequently Asked Questions that are highly relevant and customized to the specific market. The sample Frequently Asked Questions below are for understanding purposes only. For the given market, questions can be completely replaced. However, please tailor the actual questions to the market name and the insights provided in the report: 1. Who are the dominant players in the (Market Name) market? 2. What will be the size of the (Market Name) market in the coming years? 3. Which end users industry has the largest growth opportunity? 4. How will market development trends evolve over the next five years? 5. What is the nature of the competitive landscape and challenges in the (Market Name) market? 6. What go-to-market strategies are commonly adopted in the (Market Name) market? Make sure to answer to all FAQs. In the case of country-level markets, please exclude the word 'Global' and High-Impact Market Opportunities by Segment and Region where other regions are mentioned. Make sure to add catchy bullet in generated output. I have shared the reference bullet with you. Make sure to add this bullet. For heading use these bullet- ), ➤Strategic Actionable Insights for the Market, ➔ add first CTA link here ➤ Market Taxonomy and Regional coverage of Report, ➤ Leading Companies of the Market, ➔ Inserted Second CTA link, ➤Key Reasons for Buying the (insert market name here) Report, ➤Key Growth Drivers Fueling Market Expansion, ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ➤Frequently Asked Questions. Make sure do not miss any bullet including CTA bullet which is ➔. For subpointers under main headings use bullets which is in reference as provided- Strategic Actionable Insights for the Market ●, Market Taxonomy and Regional coverage of Report●, Leading Companies of the Market●. Make sure to use these bullets for given subpointers. Ensure proper bullet formatting so that each point is marked with a single bullet only, and avoid placing two bullets adjacent to each other.
"""
        
        # Make OpenAI API call
        try:
            response = openai.chat.completions.create(
                model="gpt-4o-mini",
                messages=[
                    
                    {"role": "user", "content": ai_prompt}
                ],
                max_tokens=3000,
                temperature=0.7
            )
            
            article_content = response.choices[0].message.content
            
        except Exception as e:
            print(f"OpenAI API error for {keyword}: {e}")
            return False
        
        # Create new document and add content
        doc = Document()
        
        # Split content into paragraphs and add to document
        paragraphs = article_content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Generate output filename
        today = datetime.now().strftime("%Y_%m_%d")
        clean_keyword = keyword.replace(" ", "_").replace("-", "_").lower()
        output_filename = f"CMI_AI_{clean_keyword}_{today}.docx"
        output_file_path = os.path.join(output_path, output_filename)
        
        # Save document
        doc.save(output_file_path)
        print(f"CMI AI article saved: {output_filename} in {category}")
        return True
        
    except Exception as e:
        print(f"Error in CMI AI article generation: {e}")
        return False
    
# Route to show the CMI form (you'll need to create custom_ai_content.html)
@app.route('/custom_cmi_content')
def custom_cmi_content():
    """Show CMI template content generation form"""
    return render_template('custom_cmi_content.html')



if __name__ == '__main__':
    import webbrowser
    webbrowser.open('http://127.0.0.1:5000/')
    app.run(debug=True, host='0.0.0.0', port=5000)
